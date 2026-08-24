import os
import io
import shutil
import tempfile
import concurrent.futures
from datetime import timedelta, datetime
from functools import wraps
from uuid import uuid4

import requests
from dotenv import load_dotenv
from flask import (
    Flask,
    render_template,
    request,
    session,
    redirect,
    url_for,
    jsonify,
    flash,
    make_response,
    send_file,
)
from flask_wtf.csrf import CSRFProtect
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import google.generativeai as genai
from PyPDF2 import PdfReader

# Load environment variables from .env if present
load_dotenv()

from auth_db import (
    init_db,
    create_user,
    get_user_by_login,
    get_user_by_id,
    get_user_by_email,
    update_last_login,
    create_password_reset_token,
    get_valid_reset_token,
    update_user_password,
    save_user_api_keys,
    get_user_api_keys,
    delete_user_api_key,
    save_user_resume,
    get_user_resume,
    get_user_upload_dir,
    get_canonical_main_resume_path,
    get_canonical_ats_resume_path,
    mask_api_key,
    mask_api_key_bullet,
    get_key_fingerprint,
    record_api_key_history,
    update_key_history_status,
    get_api_key_history,
    log_api_usage,
    get_gemini_usage_summary,
    get_user_search_cooldown,
    set_search_in_progress,
    set_user_search_cooldown,
    clear_user_search_cooldown,
    create_remember_token,
    get_user_by_remember_token,
    delete_remember_token,
    delete_all_user_remember_tokens,
    generate_job_id,
    save_job,
    remove_saved_job,
    get_saved_jobs,
    get_saved_job_ids,
    is_job_saved,
    mark_job_applied,
    mark_job_not_applied,
    get_applied_jobs,
    get_applied_job_ids,
    is_job_applied,
    save_job_search_results,
    get_current_job_search,
    has_current_job_search,
    delete_user_job_searches,
    save_ats_analysis,
    get_latest_ats_analysis,
    delete_user_ats_analyses,
    set_topic_completion,
    get_user_completed_topic_ids,
    get_course_progress_stats,
    get_all_courses_progress_for_user,
)
from storage_manager import (
    upload_resume_file,
    download_resume_file,
    delete_resume_file,
    get_resume_bytes,
    resume_exists,
    get_storage_path,
    temp_resume_context,
    get_resume_text_content,
    is_supabase_storage_configured,
)
from ats_engine import (
    analyze_resume_ats,
    extract_resume_document,
    parse_resume_structure,
)
from api_limit_manager import (
    fetch_serpapi_account_info,
    check_serpapi_preflight,
    classify_gemini_error,
    validate_gemini_key,
    DEFAULT_GEMINI_MODEL,
)
from auth_utils import (
    validate_password,
    validate_username,
    validate_email,
    generate_secure_reset_token,
    hash_reset_token,
    generate_remember_token,
    hash_remember_token,
    send_password_reset_email,
    check_rate_limit,
    record_failed_attempt,
    reset_failed_attempts,
    is_safe_url,
)
from skill_extractor import (
    extract_skills,
    map_skills_to_roles,
    get_missing_skills,
    calculate_role_matches,
)
from job_matcher import (
    extract_experience_info,
    extract_salary_info,
    extract_posted_time,
)
from course_data import (
    get_all_courses,
    get_course_by_id,
    get_course_topics,
    get_topic_by_id,
    extract_youtube_video_id,
)
from job_matcher import (
    extract_openings_info,
    calculate_job_match_score,
    calculate_local_match_score,
    normalize_and_filter_jobs,
)
from ai_assistant import (
    process_chat_message,
    extract_attachment_text,
)

# --- FLASK APPLICATION CONFIGURATION ---
app = Flask(__name__)

# Secret key configuration from environment variable
app.secret_key = os.environ.get(
    "SECRET_KEY", "skillbridge_dev_secret_key_change_in_production_f789a2b"
)

# Session cookie security and persistence settings
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=30)
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_SECURE"] = (
    os.environ.get("SESSION_COOKIE_SECURE", "False").lower() in ("true", "1")
)

# Enable CSRF Protection
csrf = CSRFProtect(app)

# Upload folder configuration
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Initialize SQLite database tables automatically
init_db(app)

SNAKE_GAMES = {}


# --- JINJA TEMPLATE FILTERS ---

@app.template_filter("format_date")
def format_date_filter(val):
    """Formats ISO datetime string to clean human-readable date (e.g. 12 May 2025)."""
    if not val:
        return "Not available"
    try:
        if isinstance(val, str):
            # Parse ISO string
            cleaned_val = val.replace("Z", "+00:00")
            dt = datetime.fromisoformat(cleaned_val)
        elif isinstance(val, datetime):
            dt = val
        else:
            return str(val)
        return dt.strftime("%d %b %Y")
    except Exception:
        return str(val)


@app.template_filter("format_size")
def format_size_filter(size_bytes):
    """Formats integer byte size to human readable (e.g. 512 KB or 1.4 MB)."""
    if not size_bytes or not isinstance(size_bytes, (int, float)):
        return "0 KB"
    kb = size_bytes / 1024.0
    if kb < 1024:
        return f"{int(round(kb))} KB"
    mb = kb / 1024.0
    return f"{mb:.1f} MB"


@app.template_filter("mask_bullets")
def mask_bullets_filter(val):
    """Replaces asterisk masking characters with clean visual bullet points."""
    if not val:
        return ""
    return str(val).replace("*", "•")




# --- USER DATA & AUTHENTICATION HELPERS ---

@app.before_request
def restore_persistent_session():
    """
    Checks if a returning user has a valid persistent remember token cookie.
    Restores the authenticated session seamlessly across browser restarts.
    """
    if session.get("authenticated") and session.get("user_id"):
        return

    remember_token = request.cookies.get("sb_remember")
    if not remember_token:
        return

    token_hash = hash_remember_token(remember_token)
    user, _ = get_user_by_remember_token(token_hash)
    if user and user.get("is_active"):
        session.permanent = True
        session["user_id"] = user["id"]
        session["username"] = user["username"]
        session["authenticated"] = True
        update_last_login(user["id"])

        # Automatically load user's persistent API keys from DB
        user_keys = get_user_api_keys(user["id"], decrypted=True)
        if user_keys.get("serpapi_key"):
            session["serpapi_key"] = user_keys["serpapi_key"]
        if user_keys.get("gemini_api_key"):
            session["gemini_api_key"] = user_keys["gemini_api_key"]

        # Automatically load user's persistent resume from DB
        user_resume = get_user_resume(user["id"], current_only=True)
        if user_resume:
            session["current_resume_path"] = user_resume["file_path"]
            if user_resume.get("extracted_data") and "missing_skills" in user_resume["extracted_data"]:
                session["missing_skills_data"] = user_resume["extracted_data"]["missing_skills"]


def get_current_user_id():
    """Returns current authenticated user ID or None."""
    if session.get("authenticated") and session.get("user_id"):
        return session.get("user_id")
    return None


def get_current_user():
    """Returns current authenticated user database record or None."""
    user_id = get_current_user_id()
    if not user_id:
        return None
    return get_user_by_id(user_id)


def get_current_user_api_keys(decrypted=True):
    """Returns current authenticated user's saved API keys from DB."""
    user_id = get_current_user_id()
    if not user_id:
        return {
            "serpapi_key": None,
            "gemini_api_key": None,
            "serpapi_masked": "",
            "gemini_masked": "",
            "has_serpapi": False,
            "has_gemini": False,
        }
    return get_user_api_keys(user_id, decrypted=decrypted)


def get_current_user_resume(current_only=True):
    """Returns current authenticated user's saved resume record from DB."""
    user_id = get_current_user_id()
    if not user_id:
        return None
    return get_user_resume(user_id, current_only=current_only)


# --- AUTHENTICATION GUARD DECORATOR ---
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get("authenticated") or not session.get("user_id"):
            next_url = request.full_path if request.method == "GET" else request.path
            if next_url and next_url.endswith("?"):
                next_url = next_url[:-1]
            return redirect(url_for("login_route", next=next_url))
        return f(*args, **kwargs)
    return decorated_function


# --- AUTHENTICATION ROUTES ---

@app.route("/register", methods=["GET", "POST"])
def register_route():
    # If user is already authenticated, redirect them based on their state
    if session.get("authenticated") and session.get("user_id"):
        user_id = session.get("user_id")
        user_keys = get_user_api_keys(user_id)
        user_resume = get_user_resume(user_id, current_only=True)
        if not user_keys.get("has_serpapi"):
            return redirect(url_for("setup"))
        if not user_resume:
            return redirect(url_for("upload_resume_page"))
        return redirect(url_for("jobs_home"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")

        # 1. Validate Username
        is_user_valid, user_err = validate_username(username)
        if not is_user_valid:
            flash(user_err, "error")
            return render_template("register.html", username=username, email=email)

        # 2. Validate Email
        is_email_valid, email_err = validate_email(email)
        if not is_email_valid:
            flash(email_err, "error")
            return render_template("register.html", username=username, email=email)

        # 3. Validate Password
        is_pass_valid, pass_err = validate_password(password)
        if not is_pass_valid:
            flash(pass_err, "error")
            return render_template("register.html", username=username, email=email)

        # 4. Hash password and store user in SQLite
        password_hash = generate_password_hash(password)
        try:
            create_user(username, email, password_hash)
            flash("Account created successfully. Please log in.", "success")
            return redirect(url_for("login_route"))
        except ValueError as e:
            if str(e) == "USERNAME_EXISTS":
                flash("Username is already in use.", "error")
            elif str(e) == "EMAIL_EXISTS":
                flash("An account with this email already exists.", "error")
            else:
                flash("An error occurred during registration. Please try again.", "error")
            return render_template("register.html", username=username, email=email)
        except Exception:
            flash("An unexpected error occurred. Please try again later.", "error")
            return render_template("register.html", username=username, email=email)

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login_route():
    # If user is already authenticated, redirect them based on their state
    if session.get("authenticated") and session.get("user_id"):
        user_id = session.get("user_id")
        user_keys = get_user_api_keys(user_id)
        user_resume = get_user_resume(user_id, current_only=True)
        if not user_keys.get("has_serpapi"):
            return redirect(url_for("setup"))
        if not user_resume:
            return redirect(url_for("upload_resume_page"))
        return redirect(url_for("jobs_home"))

    if request.method == "POST":
        login_id = request.form.get("login_id", "").strip()
        password = request.form.get("password", "")
        client_ip = request.remote_addr or "unknown"

        # Rate limiting check
        is_allowed, remaining_sec = check_rate_limit(client_ip)
        if not is_allowed:
            flash(f"Too many failed login attempts. Please try again in {remaining_sec} seconds.", "error")
            return render_template("login.html", login_id=login_id)

        if not login_id or not password:
            record_failed_attempt(client_ip)
            flash("Invalid username or password.", "error")
            return render_template("login.html", login_id=login_id)

        user = get_user_by_login(login_id)

        # Verify password securely
        if (
            not user
            or not user["is_active"]
            or not check_password_hash(user["password_hash"], password)
        ):
            record_failed_attempt(client_ip)
            flash("Invalid username or password.", "error")
            return render_template("login.html", login_id=login_id)

        # Successful Login
        reset_failed_attempts(client_ip)
        
        # Clear previous session state
        session.clear()

        # Set persistent authenticated user session
        session.permanent = True
        session["user_id"] = user["id"]
        session["username"] = user["username"]
        session["authenticated"] = True

        # Update last login timestamp in DB
        update_last_login(user["id"])

        # Automatically load user's persistent API keys from DB
        user_keys = get_user_api_keys(user["id"], decrypted=True)
        if user_keys.get("serpapi_key"):
            session["serpapi_key"] = user_keys["serpapi_key"]
        if user_keys.get("gemini_api_key"):
            session["gemini_api_key"] = user_keys["gemini_api_key"]

        # Automatically load user's persistent resume from DB
        user_resume = get_user_resume(user["id"], current_only=True)
        if user_resume:
            session["current_resume_path"] = user_resume["file_path"]
            if user_resume.get("extracted_data") and "missing_skills" in user_resume["extracted_data"]:
                session["missing_skills_data"] = user_resume["extracted_data"]["missing_skills"]

        # Create persistent remember token in DB
        raw_token, token_hash, expires_at = generate_remember_token(days=30)
        create_remember_token(
            user_id=user["id"],
            token_hash=token_hash,
            expires_at=expires_at,
            user_agent=str(request.user_agent.string if request.user_agent else "unknown"),
            ip_address=client_ip,
        )

        # Determine destination URL
        next_param = request.args.get("next") or request.form.get("next")
        if next_param and is_safe_url(next_param):
            target_url = next_param
        elif not user_keys.get("has_serpapi"):
            target_url = url_for("setup")
        elif not user_resume:
            target_url = url_for("upload_resume_page")
        else:
            target_url = url_for("jobs_home")

        response = make_response(redirect(target_url))
        is_secure = app.config.get("SESSION_COOKIE_SECURE", False)
        response.set_cookie(
            "sb_remember",
            raw_token,
            max_age=30 * 24 * 60 * 60,
            httponly=True,
            samesite="Lax",
            secure=is_secure,
        )
        return response

    return render_template("login.html")



@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password_route():
    dev_reset_url = None
    smtp_status = None
    if request.method == "POST":
        email = request.form.get("email", "").strip()

        # Non-enumerating feedback message
        flash("If an account exists for that email address, a password reset link has been sent.", "info")

        if email:
            user = get_user_by_email(email)
            if user and user["is_active"]:
                raw_token, token_hash, expires_at = generate_secure_reset_token()
                create_password_reset_token(user["id"], token_hash, expires_at)
                result = send_password_reset_email(user["email"], user["username"], raw_token)
                
                status = result.get("status")
                provider = result.get("provider")
                smtp_status = status
                if status == "SENT":
                    flash("A password reset email has been sent to your inbox.", "success")
                elif status in ("NO_PASSWORD", "AUTH_ERROR", "CONNECT_ERROR"):
                    if provider == "smtp" or not provider:
                        dev_reset_url = result.get("reset_url")
                        if status == "AUTH_ERROR":
                            flash("Gmail SMTP authentication failed. Please verify the Gmail App Password in .env.", "error")
                        elif status == "CONNECT_ERROR":
                            flash("Could not connect to Gmail SMTP server. Please check your network connection.", "error")

        return render_template("forgot_password.html", email="", dev_reset_url=dev_reset_url, smtp_status=smtp_status)

    return render_template("forgot_password.html")


@app.route("/reset-password/<token>", methods=["GET", "POST"])
def reset_password_route(token):
    token_hash = hash_reset_token(token)
    token_row, status = get_valid_reset_token(token_hash)

    if status != "VALID" or not token_row:
        return render_template("reset_password.html", is_valid_token=False, token=token)

    if request.method == "POST":
        password = request.form.get("password", "")
        confirm_password = request.form.get("confirm_password", "")

        if password != confirm_password:
            flash("Passwords do not match.", "error")
            return render_template("reset_password.html", is_valid_token=True, token=token)

        is_valid_pass, pass_err = validate_password(password)
        if not is_valid_pass:
            flash(pass_err, "error")
            return render_template("reset_password.html", is_valid_token=True, token=token)

        # Update password in database and invalidate tokens (preserves all API keys & resumes)
        new_password_hash = generate_password_hash(password)
        update_user_password(token_row["user_id"], new_password_hash)

        flash("Password updated successfully. Please log in with your new password.", "success")
        return redirect(url_for("login_route"))

    return render_template("reset_password.html", is_valid_token=True, token=token)


@app.route("/logout")
def logout():
    remember_token = request.cookies.get("sb_remember")
    if remember_token:
        token_hash = hash_remember_token(remember_token)
        delete_remember_token(token_hash)
    session.clear()
    flash("You have been logged out.", "info")
    resp = make_response(redirect(url_for("login_route")))
    resp.delete_cookie("sb_remember")
    return resp


# --- RESUME VALIDATION & EXTRACTION HELPERS ---

def validate_resume_document(file_storage_or_bytes, filename: str, is_profile: bool = True):
    """
    Validates file format, size limit (<= 5MB), readability, and structure.
    Returns (is_valid: bool, error_message: str, content_bytes: bytes).
    """
    if not filename:
        return False, "No file selected. Please select a resume file.", b""

    ext = os.path.splitext(filename)[1].lower()
    allowed_exts = [".pdf"] if is_profile else [".pdf", ".docx", ".doc"]
    if ext not in allowed_exts:
        msg = "Invalid file format. Only PDF files are supported." if is_profile else "Please upload a PDF or DOCX file."
        return False, msg, b""

    if hasattr(file_storage_or_bytes, "read"):
        content_bytes = file_storage_or_bytes.read()
    elif isinstance(file_storage_or_bytes, bytes):
        content_bytes = file_storage_or_bytes
    else:
        return False, "Invalid file object.", b""

    if len(content_bytes) == 0:
        return False, "The uploaded file is empty. Please upload a valid document.", b""

    if len(content_bytes) > 5 * 1024 * 1024:
        return False, "File size exceeds 5MB limit. Please upload a smaller document.", b""

    # Validate PDF readability & structure
    if ext == ".pdf":
        try:
            reader = PdfReader(io.BytesIO(content_bytes))
            if len(reader.pages) == 0:
                return False, "The uploaded PDF has no readable pages.", b""
            # Verify we can access pages without corruption
            _ = [p.extract_text() or "" for p in reader.pages]
        except Exception as e:
            return False, f"The uploaded PDF is invalid or corrupted: {str(e)}", b""
    elif ext in (".docx", ".doc"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content_bytes))
            if not doc.paragraphs and not doc.tables:
                return False, "The uploaded document has no readable content.", b""
        except Exception:
            # If docx library cannot parse (e.g. legacy binary .doc), verify basic content presence
            if len(content_bytes) < 32:
                return False, "The uploaded document is too short or invalid.", b""

    return True, "", content_bytes


def get_pdf_text(pdf_path, user_id=None):
    if not pdf_path:
        return ""
    if os.path.exists(pdf_path) and os.path.isfile(pdf_path):
        try:
            reader = PdfReader(pdf_path)
            text = " ".join(page.extract_text() or "" for page in reader.pages)
            return text.strip()
        except Exception:
            return ""
    # Safe storage abstraction extraction (Supabase or relative path)
    try:
        return get_resume_text_content(user_id=user_id, storage_path=pdf_path)
    except Exception:
        return ""


def get_gemini_model():
    """
    Returns configured application Gemini generative model.
    """
    return genai.GenerativeModel(DEFAULT_GEMINI_MODEL)


# --- AGENT 1: COVER LETTER WRITER ---
def agent_write_cover_letter(resume_text, job_title, company, location, api_key):
    genai.configure(api_key=api_key)
    model = get_gemini_model()
    prompt = f"""
    Act as a Professional Career Coach. Write a Cover Letter.
    
    JOB DETAILS:
    Role: {job_title}
    Company: {company}
    Location: {location}
    
    MY RESUME:
    {resume_text[:3000]}... (truncated)
    
    INSTRUCTIONS:
    1. Keep it under 200 words.
    2. Specifically mention 2 skills from my resume that fit this job.
    3. Professional and confident tone.
    4. Output ONLY the body of the letter (no placeholders like [Your Name]).
    """
    response = model.generate_content(prompt)
    return response.text


# --- AGENT 2: GAP ANALYSIS COACH (Uses Tools) ---
def search_tutorial(skill, serp_key):
    """Tool: Searches for the best free tutorial for a specific skill."""
    if not skill or not serp_key:
        return None
    
    url = "https://serpapi.com/search.json"
    params = {
        "engine": "google",
        "q": f"best free {skill} tutorial youtube 2025", 
        "api_key": serp_key,
        "num": 1
    }
    try:
        response = requests.get(url, params=params, timeout=5)
        data = response.json()
        if "organic_results" in data and len(data["organic_results"]) > 0:
            result = data["organic_results"][0]
            return {"title": result.get("title"), "link": result.get("link")}
    except Exception as e:
        print(f"Search failed for {skill}: {e}")
    return None


def agent_gap_analysis(role, missing_skills, serp_key, gemini_key):
    genai.configure(api_key=gemini_key)
    
    top_skills = missing_skills[:3]
    resources = {}
    if serp_key:
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future_to_skill = {executor.submit(search_tutorial, skill, serp_key): skill for skill in top_skills}
            for future in concurrent.futures.as_completed(future_to_skill):
                skill = future_to_skill[future]
                try:
                    resources[skill] = future.result()
                except Exception:
                    resources[skill] = None
    
    resource_text = ""
    for skill in top_skills:
        res = resources.get(skill)
        if res:
            resource_text += f"- For {skill}: Found tutorial '{res['title']}' at {res['link']}\n"
        else:
            resource_text += f"- For {skill}: Please recommend a popular free resource (like Coursera, YouTube, or Documentation).\n"

    model = get_gemini_model()

    prompt = f"""
    You are a Technical Mentor, your name is "RJM-agent" introduce your self in single line in professional way. Create a "Gap Analysis Learning Plan" for a user wanting to be a {role}.
    
    MISSING SKILLS: {', '.join(top_skills)}
    
    FOUND RESOURCES (Prioritize these links):
    {resource_text}
    
    INSTRUCTIONS:
    1. Create a concise 2-week plan.
    2. For each skill, explain WHY it is critical.
    3. Provide the specific link if I gave one, otherwise suggest a general one.
    4. Use HTML formatting: <b>Bold</b> for emphasis, <br> for line breaks.
    5. Do NOT use Markdown (like ** or ##), use HTML tags only.
    """
    response = model.generate_content(prompt)
    return response.text


# --- MARKET DATA & JOB SEARCH ---
def fetch_market_count(role, api_key):
    url = "https://serpapi.com/search.json"
    try:
        params = {"engine": "google", "q": f"{role} jobs in India", "location": "India", "api_key": api_key}
        response = requests.get(url, params=params, timeout=8)
        data = response.json()
        search_info = data.get("search_information", {})
        total_results = search_info.get("total_results_formatted")
        if not total_results:
             if search_info.get("total_results"):
                 total_results = f"{int(search_info.get('total_results')):,}+"
        return {"role": role, "count": total_results if total_results else "High Demand"}
    except Exception:
        return {"role": role, "count": "Data Unavailable"}


def get_market_insights(api_key):
    trending_roles = ["Full Stack Developer", "Data Scientist", "Python Developer", "DevOps Engineer", "Cyber Security Engineer", "Java Developer", "Cloud Architect", "Frontend Developer", "Machine Learning Engineer", "UI/UX Designer"]
    insights = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
        future_to_role = {executor.submit(fetch_market_count, role, api_key): role for role in trending_roles}
        for future in concurrent.futures.as_completed(future_to_role):
            insights.append(future.result())
    insights.sort(key=lambda x: trending_roles.index(x["role"]) if x["role"] in trending_roles else 999)
    return insights


def get_fallback_curated_jobs(job_roles=None, resume_data=None, gemini_api_key=None):
    """
    Curated high-quality, fresher-compatible jobs tailored to matched roles and skills.
    Used when live searches return 0 or during test runs.
    Calculates individual match scores based on the candidate's actual resume.
    """
    raw_curated = [
        {
            "title": "Software Engineer",
            "company_name": "Microsoft",
            "company_brand": "microsoft",
            "is_verified": True,
            "location": "Hyderabad, India",
            "salary": "₹ 18 - 28 LPA",
            "experience": "0-2 Yrs",
            "job_type": "Full-time",
            "posted_at": "Posted 2h ago",
            "apply_link": "https://careers.microsoft.com",
            "thumbnail": None,
            "description": "Microsoft is seeking entry-level Software Engineers (0-2 years) to develop cloud and web services. Skills required: Python, Data Structures, Algorithms, SQL, Git, problem solving, and OOP.",
            "searched_role": "Software Engineer",
        },
        {
            "title": "Backend Developer",
            "company_name": "Google",
            "company_brand": "google",
            "is_verified": True,
            "location": "Bangalore, India",
            "salary": "₹ 16 - 24 LPA",
            "experience": "Freshers",
            "job_type": "Full-time",
            "posted_at": "Posted 5h ago",
            "apply_link": "https://careers.google.com",
            "thumbnail": None,
            "description": "Google India is hiring Backend Developers for freshers and early-career engineers. Requirements: Python, Flask, Django, REST APIs, SQL, PostgreSQL, Docker, and system design fundamentals.",
            "searched_role": "Backend Developer",
        },
        {
            "title": "Jr. Full Stack Developer",
            "company_name": "Swiggy",
            "company_brand": "swiggy",
            "is_verified": True,
            "location": "Hyderabad, India",
            "salary": "Not mentioned",
            "experience": "No experience required",
            "job_type": "Full-time",
            "posted_at": "Posted 6h ago",
            "apply_link": "https://careers.swiggy.com",
            "thumbnail": None,
            "description": "Swiggy Engineering is looking for Junior Full Stack Developers with strong skills in JavaScript, React, Node.js, HTML, CSS, SQL, MongoDB, and Git. Freshers and college graduates welcome.",
            "searched_role": "Full Stack Developer",
        },
        {
            "title": "Python Developer",
            "company_name": "Flipkart",
            "company_brand": "flipkart",
            "is_verified": True,
            "location": "Bangalore, India",
            "salary": "₹ 14 - 20 LPA",
            "experience": "0-1 Yrs",
            "job_type": "Full-time",
            "posted_at": "Posted 1d ago",
            "apply_link": "https://flipkartcareers.com",
            "thumbnail": None,
            "description": "Flipkart is hiring Python Developers with 0-1 years of experience. Responsibilities include building scalable services with Python, SQL, REST APIs, OOP, and data structures.",
            "searched_role": "Python Developer",
        },
        {
            "title": "Associate Cloud & DevOps Engineer",
            "company_name": "Amazon",
            "company_brand": "amazon",
            "is_verified": True,
            "location": "Hyderabad, India",
            "salary": "Not mentioned",
            "experience": "Entry level",
            "job_type": "Full-time",
            "posted_at": "Posted 2d ago",
            "apply_link": "https://amazon.jobs",
            "thumbnail": None,
            "description": "Amazon is seeking Associate Cloud and DevOps Engineers for entry level candidates. Proficiencies: Linux, Docker, AWS fundamentals, Python, Bash scripting, Git, and networking.",
            "searched_role": "DevOps Engineer",
        },
        {
            "title": "Software Developer",
            "company_name": "Zomato",
            "company_brand": "zomato",
            "is_verified": True,
            "location": "Gurgaon, India",
            "salary": "Not mentioned",
            "experience": "",
            "job_type": "Full-time",
            "posted_at": "Posted 3d ago",
            "apply_link": "https://zomato.com/careers",
            "thumbnail": None,
            "description": "Zomato is hiring software developers to build consumer-facing applications. Strong foundation in computer science, coding proficiency in Python or JavaScript, SQL, and web technologies.",
            "searched_role": "Software Developer",
        },
    ]

    r_data = resume_data if isinstance(resume_data, dict) else {
        "skills": resume_data if isinstance(resume_data, list) else ["Python", "SQL", "Flask", "React", "Git"],
        "roles": job_roles or ["Full Stack Developer", "Backend Developer", "Python Developer"],
        "resume_text": "",
    }

    curated = []
    for j in raw_curated:
        job_copy = dict(j)
        if not job_copy.get("job_id"):
            job_copy["job_id"] = generate_job_id(
                job_copy.get("company_name"),
                job_copy.get("title"),
                job_copy.get("location"),
                job_copy.get("apply_link") or ""
            )
        # Calculate INDIVIDUAL match score based on candidate resume
        score, matched_skills = calculate_job_match_score(r_data, job_copy, gemini_api_key)
        job_copy["match_percent"] = score
        job_copy["match_score"] = score
        job_copy["matching_skills"] = matched_skills
        job_copy["openings"] = extract_openings_info(job_copy)
        curated.append(job_copy)

    curated.sort(key=lambda j: j.get("match_percent", 0), reverse=True)
    return curated


def fetch_jobs_multi_location(job_roles, api_key, resume_data=None, gemini_api_key=None):
    """
    Performs real SerpAPI job searches across multiple Indian tech hubs.
    Deduplicates, extracts experience & salary, filters out experienced-only roles,
    and calculates individual match percentages using candidate resume and Gemini / local fallback.
    """
    if not api_key:
        return get_fallback_curated_jobs(job_roles, resume_data, gemini_api_key)

    url = "https://serpapi.com/search.json"
    locations = [
        "Hyderabad, India",
        "Bangalore, India",
        "Chennai, India",
        "Pune, India",
        "Mumbai, India",
        "Delhi, India",
        "Remote",
        "India",
    ]
    roles_to_query = (job_roles[:3] if job_roles else ["Full Stack Developer", "Backend Developer", "Python Developer"])

    raw_results = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=6) as executor:
        future_map = {}
        for role in roles_to_query:
            for loc in locations[:3]:
                params = {
                    "engine": "google_jobs",
                    "q": role,
                    "location": loc,
                    "hl": "en",
                    "api_key": api_key,
                }
                future = executor.submit(requests.get, url, params=params, timeout=8)
                future_map[future] = (role, loc)

        for future in concurrent.futures.as_completed(future_map):
            role, loc = future_map[future]
            try:
                resp = future.result()
                if resp.status_code == 200:
                    data = resp.json()
                    jobs_res = data.get("jobs_results", [])
                    for j in jobs_res:
                        j["searched_role"] = role
                        j["searched_location"] = loc
                        raw_results.append(j)
            except Exception as e:
                print(f"SerpAPI query failed for {role} in {loc}: {e}")

    if not raw_results:
        return get_fallback_curated_jobs(job_roles, resume_data, gemini_api_key)

    r_data = resume_data if isinstance(resume_data, dict) else {
        "skills": resume_data if isinstance(resume_data, list) else [],
        "roles": job_roles or [],
        "resume_text": "",
    }

    # Run the comprehensive 10-step processing pipeline
    processed = normalize_and_filter_jobs(raw_results, r_data, gemini_api_key)

    if not processed:
        # If all live jobs were filtered out (e.g. 5+ years only), return curated fresher jobs
        return get_fallback_curated_jobs(job_roles, resume_data, gemini_api_key)

    return processed


def fetch_jobs(job_roles, api_key, resume_data=None, gemini_api_key=None):
    """Backward compatibility wrapper for fetch_jobs."""
    return fetch_jobs_multi_location(job_roles, api_key, resume_data, gemini_api_key)


# --- APPLICATION CORE ROUTES ---

def get_snake_id():
    if "snake_game_id" not in session:
        session["snake_game_id"] = str(uuid4())
    return session["snake_game_id"]


@app.route("/")
def home():
    """
    Intelligent Root Route:
    - CASE 2: Not authenticated -> /login
    - CASE 3: Authenticated but missing API keys -> /setup
    - CASE 4: Authenticated with API keys but missing resume -> /upload-resume
    - CASE 1 & 5: Authenticated with API keys and resume -> /jobs-home
    """
    if not session.get("authenticated") or not session.get("user_id"):
        return redirect(url_for("login_route"))
    
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)

    if not user_keys.get("has_serpapi"):
        return redirect(url_for("setup"))

    # Ensure session keys are populated from DB
    if "serpapi_key" not in session and user_keys.get("serpapi_key"):
        session["serpapi_key"] = user_keys["serpapi_key"]
    if "gemini_api_key" not in session and user_keys.get("gemini_api_key"):
        session["gemini_api_key"] = user_keys["gemini_api_key"]

    user_resume = get_user_resume(user_id, current_only=True)
    if not user_resume:
        return redirect(url_for("upload_resume_page"))

    if "current_resume_path" not in session:
        session["current_resume_path"] = user_resume["file_path"]
        if user_resume.get("extracted_data") and "missing_skills" in user_resume["extracted_data"]:
            session["missing_skills_data"] = user_resume["extracted_data"]["missing_skills"]

    return redirect(url_for("jobs_home"))


@app.route("/jobs-home")
@login_required
def jobs_home():
    """
    Main authenticated Jobs Home / Career Home page.
    Displays hero card, feature cards, statistics, and bottom navigation.
    """
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    if not user_keys.get("has_serpapi"):
        flash("Please configure your SerpAPI key before proceeding.", "warning")
        return redirect(url_for("setup"))

    user_resume = get_user_resume(user_id, current_only=True)
    if not user_resume:
        flash("Please upload your resume to access your Career Home.", "info")
        return redirect(url_for("upload_resume_page"))

    # Ensure session data is synced
    if "serpapi_key" not in session and user_keys.get("serpapi_key"):
        session["serpapi_key"] = user_keys["serpapi_key"]
    if "gemini_api_key" not in session and user_keys.get("gemini_api_key"):
        session["gemini_api_key"] = user_keys["gemini_api_key"]
    if "current_resume_path" not in session:
        session["current_resume_path"] = user_resume["file_path"]
        if user_resume.get("extracted_data") and "missing_skills" in user_resume["extracted_data"]:
            session["missing_skills_data"] = user_resume["extracted_data"]["missing_skills"]

    user = get_current_user()
    return render_template(
        "jobs_home.html",
        user=user,
        resume=user_resume,
        keys=user_keys,
    )


@app.route("/resume")
@app.route("/upload-resume")
@login_required
def upload_resume_page():
    """
    Resume upload / ATS Resume Analyzer Home page.
    """
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    if not user_keys.get("has_serpapi"):
        flash("Please configure your API keys before uploading your resume.", "warning")
        return redirect(url_for("setup"))

    user_resume = get_user_resume(user_id, current_only=True)
    return render_template("index.html", existing_resume=user_resume)


@app.route("/ats-score")
@login_required
def ats_score_page():
    """
    Authoritative ATS Score & Deep Resume Analysis page.
    Displays latest persistent ATS analysis for the authenticated user.
    Survives refresh, navigation, and application restart without re-invoking APIs.
    """
    user_id = session.get("user_id")
    ats_data = get_latest_ats_analysis(user_id)
    return render_template("ats_score.html", ats_data=ats_data)


@app.route("/download-ats-resume")
@login_required
def download_ats_resume():
    """
    Downloads the currently stored ATS resume for the authenticated user only.
    Enforces strict session user ownership. Never exposes server filesystem paths.
    Supports both Supabase Storage (cloud) and local uploads/ fallback.
    """
    user_id = session.get("user_id")
    ats_data = get_latest_ats_analysis(user_id)
    if not ats_data or not ats_data.get("file_path"):
        flash("No ATS resume found for download. Please upload a resume first.", "warning")
        return redirect(url_for("upload_resume_page"))

    file_path = ats_data["file_path"]
    download_name = ats_data.get("filename") or "ATS_Resume.pdf"
    ext = os.path.splitext(download_name)[1].lower() or ".pdf"
    default_mimetype = "application/pdf" if ext == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # 1. Try downloading from storage abstraction (Supabase or local)
    content_bytes, _, detected_mimetype = download_resume_file(
        user_id=user_id,
        resume_type="ats",
        storage_path=file_path,
    )
    if content_bytes:
        return send_file(
            io.BytesIO(content_bytes),
            as_attachment=True,
            download_name=download_name,
            mimetype=detected_mimetype or default_mimetype,
        )

    # 2. Local filesystem direct fallback
    if not os.path.isabs(file_path):
        cand_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), file_path)
        if os.path.exists(cand_path):
            file_path = cand_path

    if os.path.exists(file_path) and os.path.isfile(file_path):
        return send_file(
            file_path,
            as_attachment=True,
            download_name=download_name,
            mimetype=default_mimetype,
        )

    flash("Stored ATS resume file could not be found on server.", "error")
    return redirect(url_for("ats_score_page"))


@app.route("/snake")
@login_required
def snake():
    return render_template("snake.html")


@app.route("/setup", methods=["GET", "POST"])
@login_required
def setup():
    user_id = session.get("user_id")

    if request.method == "POST":
        api_key = request.form.get("api_key", "").strip()
        gemini_key = request.form.get("gemini_key", "").strip()

        # Save to database (preserves any existing key if not changed)
        save_user_api_keys(
            user_id=user_id,
            serpapi_key=api_key if api_key else None,
            gemini_api_key=gemini_key if gemini_key else None,
        )

        # Reload fresh keys from database into session
        user_keys = get_user_api_keys(user_id, decrypted=True)
        if user_keys.get("serpapi_key"):
            session["serpapi_key"] = user_keys["serpapi_key"]
        if user_keys.get("gemini_api_key"):
            session["gemini_api_key"] = user_keys["gemini_api_key"]

        flash("API keys saved successfully.", "success")

        # Intelligently route to Jobs Home if resume already uploaded, else to Resume Upload
        user_resume = get_user_resume(user_id, current_only=True)
        if user_resume:
            return redirect(url_for("jobs_home"))
        return redirect(url_for("upload_resume_page"))

    # GET: Load current user's saved keys from database
    user_keys = get_user_api_keys(user_id, decrypted=True)
    if user_keys.get("serpapi_key"):
        session["serpapi_key"] = user_keys["serpapi_key"]
    if user_keys.get("gemini_api_key"):
        session["gemini_api_key"] = user_keys["gemini_api_key"]

    return render_template("setup.html", keys=user_keys)


@app.route("/api/snake/new", methods=["POST"])
@login_required
def snake_new():
    game_id = get_snake_id()
    return jsonify({"status": "ready", "game_id": game_id})


@app.route("/upload", methods=["POST"])
@login_required
def upload_resume():
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    
    if not user_keys.get("has_serpapi"):
        flash("Please set up your SerpAPI key before uploading your resume.", "warning")
        return redirect(url_for("setup"))

    serpapi_key = user_keys.get("serpapi_key")
    session["serpapi_key"] = serpapi_key

    if "resume" not in request.files:
        flash("No file was uploaded. Please select a resume file.", "error")
        return redirect(url_for("upload_resume_page"))
    
    resume = request.files["resume"]
    is_valid, err_msg, content_bytes = validate_resume_document(
        resume, resume.filename, is_profile=False
    )
    if not is_valid:
        flash(err_msg, "error")
        return redirect(url_for("upload_resume_page"))

    ext = os.path.splitext(resume.filename)[1].lower()

    # Upload ATS Resume via universal storage manager (Supabase Storage or Local Fallback)
    try:
        upload_res = upload_resume_file(
            user_id=user_id,
            file_data=content_bytes,
            original_filename=resume.filename,
            resume_type="ats",
        )
    except Exception as e:
        print(f"[ATS Resume Upload Error]: {e}")
        flash("Could not save ATS resume file. Please try again.", "error")
        return redirect(url_for("upload_resume_page"))

    stored_ats_filename = upload_res["stored_filename"]
    ats_file_path = upload_res["storage_path"]
    file_size = upload_res["file_size"]

    # Deep Resume Extraction & Deterministic ATS Analysis using temporary processing context
    gemini_key = user_keys.get("gemini_api_key") or session.get("gemini_api_key")
    try:
        with temp_resume_context(file_bytes=content_bytes, filename=resume.filename) as temp_ats_path:
            ats_result = analyze_resume_ats(
                file_path=temp_ats_path,
                original_filename=resume.filename,
                gemini_api_key=gemini_key,
            )

        # Atomically save authoritative ATS analysis in database (safely replaces old analysis)
        save_ats_analysis(
            user_id=user_id,
            filename=resume.filename,
            stored_filename=stored_ats_filename,
            file_path=ats_file_path,
            file_size=file_size,
            file_type=ext.lstrip("."),
            analysis_data=ats_result,
        )
    except Exception as e:
        print(f"[ATS Pipeline Error]: {e}")
        flash("Resume analysis could not be completed. Please try again.", "error")
        return redirect(url_for("upload_resume_page"))

    # Check if user had an existing main profile resume
    user_resume = get_user_resume(user_id, current_only=True)
    if not user_resume:
        # Onboarding: also upload and initialize baseline profile resume record
        main_upload = upload_resume_file(
            user_id=user_id,
            file_data=content_bytes,
            original_filename=resume.filename,
            resume_type="main",
        )
        main_file_path = main_upload["storage_path"]
        main_stored_filename = main_upload["stored_filename"]

        skills = ats_result.get("detected_skills", [])
        if not skills:
            with temp_resume_context(file_bytes=content_bytes, filename=resume.filename) as temp_proc_path:
                skills = extract_skills(temp_proc_path) if ext == ".pdf" else ["Python", "SQL"]
        job_roles = map_skills_to_roles(skills)
        missing_skills = get_missing_skills(skills, job_roles)
        save_user_resume(
            user_id=user_id,
            original_filename=resume.filename,
            stored_filename=main_stored_filename,
            file_path=main_file_path,
            file_size=file_size,
            file_type=ext.lstrip("."),
            extracted_data={
                "skills": skills,
                "roles": job_roles,
                "missing_skills": missing_skills,
                "resume_text_preview": "",
            },
            processing_status="completed",
        )
        session["current_resume_path"] = main_file_path
        session["missing_skills_data"] = missing_skills

        # If upload was from onboarding (not explicit ATS flow), route to jobs-home
        if request.form.get("flow") != "ats":
            flash("Resume uploaded and saved successfully!", "success")
            return redirect(url_for("jobs_home"))

    flash("Resume analyzed successfully! Check your ATS score and breakdown below.", "success")
    return redirect(url_for("ats_score_page"))


@app.route("/analyze-and-find-jobs", methods=["GET", "POST"])
@login_required
def analyze_and_find_jobs():
    """
    Explicit Job Search & Resume Intelligence analysis trigger.
    ONLY triggered when user explicitly clicks "Analyze Resume & Find Jobs".
    Performs a fresh live API search (SerpAPI + Gemini), atomically saves the complete
    normalized result set to SQLite database for the authenticated user, and redirects to /jobs.
    If the search fails or is blocked, preserves previous working results without data loss.
    """
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    if not user_keys.get("has_serpapi"):
        flash("Please configure your SerpAPI key before searching jobs.", "warning")
        return redirect(url_for("setup"))
    if not user_keys.get("has_gemini"):
        flash("Please configure your Gemini API key before proceeding with AI analysis.", "warning")
        return redirect(url_for("setup"))

    user_resume = get_user_resume(user_id, current_only=True)
    if not user_resume:
        flash("Please upload your resume first.", "info")
        return redirect(url_for("upload_resume_page"))

    # Ensure session keys & resume path are synchronized from database
    serpapi_key = user_keys.get("serpapi_key")
    gemini_key = user_keys.get("gemini_api_key")
    if serpapi_key:
        session["serpapi_key"] = serpapi_key
    if gemini_key:
        session["gemini_api_key"] = gemini_key

    resume_path = user_resume.get("file_path")
    if resume_path:
        session["current_resume_path"] = resume_path

    # Authoritative Pre-Flight Check: Search-in-progress, 60s Cooldown, and Account Limits
    allowed, reason, msg, rem_sec, action_url = check_serpapi_preflight(user_id, serpapi_key, check_account_api=True)
    if not allowed:
        if reason == "APPLICATION_COOLDOWN":
            flash(f"Next search available in: 00:{rem_sec:02d}", "warning")
            if has_current_job_search(user_id):
                return redirect(url_for("jobs_route"))
            return redirect(url_for("jobs_home"))
        elif reason == "SEARCH_IN_PROGRESS":
            flash(msg, "info")
            if has_current_job_search(user_id):
                return redirect(url_for("jobs_route"))
            return redirect(url_for("jobs_home"))
        elif reason == "MONTHLY_LIMIT_REACHED":
            flash(msg, "error")
            target_url = url_for("jobs_route") if has_current_job_search(user_id) else url_for("jobs_home")
            return redirect(target_url + "?api_error=serpapi_quota")
        elif reason == "HOURLY_LIMIT_REACHED":
            flash(msg, "error")
            target_url = url_for("jobs_route") if has_current_job_search(user_id) else url_for("jobs_home")
            return redirect(target_url + "?api_error=serpapi_hourly")
        elif reason == "INVALID_KEY":
            flash(msg, "error")
            target_url = url_for("jobs_route") if has_current_job_search(user_id) else url_for("jobs_home")
            return redirect(target_url + "?api_error=serpapi_invalid")
        else:
            flash(msg or "SerpAPI search is currently unavailable.", "error")
            if has_current_job_search(user_id):
                return redirect(url_for("jobs_route"))
            return redirect(url_for("jobs_home"))

    # Set Search-in-progress flag to protect against simultaneous requests
    set_search_in_progress(user_id, in_progress=True)

    try:
        # Extract or retrieve skills from stored resume
        skills = []
        if resume_path:
            with temp_resume_context(user_id=user_id, storage_path=resume_path) as proc_path:
                if proc_path and os.path.exists(proc_path):
                    skills = extract_skills(proc_path)

        if not skills and user_resume.get("extracted_data"):
            skills = user_resume["extracted_data"].get("skills", [])

        if not skills:
            skills = [
                "Python", "SQL", "JavaScript", "HTML", "CSS", "React", "Node.js", "Git",
                "MongoDB", "Express.js", "Flask", "REST API", "Docker", "PostgreSQL",
                "Linux", "Problem Solving", "Data Structures", "OOP", "Communication", "Teamwork"
            ]

        job_roles = map_skills_to_roles(skills)
        missing_skills = get_missing_skills(skills, job_roles)
        role_matches = calculate_role_matches(skills, job_roles)

        # Cache missing skills in session for the learning plan generator
        session["missing_skills_data"] = missing_skills

        # Build complete candidate resume data for per-job semantic & local matching
        resume_text = get_pdf_text(resume_path, user_id=user_id) if resume_path else ""
        if not resume_text and user_resume and user_resume.get("extracted_data"):
            resume_text = user_resume["extracted_data"].get("resume_text_preview", "")

        resume_data = {
            "skills": skills,
            "roles": job_roles,
            "missing_skills": missing_skills,
            "resume_text": resume_text,
        }

        # Fetch live jobs from SerpAPI across multiple locations with fresher filtering & individual scoring
        jobs = fetch_jobs_multi_location(job_roles, serpapi_key, resume_data, gemini_key)
        market_insights = get_market_insights(serpapi_key)

        # Ensure deterministic job_id and complete normalized fields on all jobs before persisting
        for j in jobs:
            if not j.get("job_id"):
                j["job_id"] = generate_job_id(
                    j.get("company_name") or j.get("company"),
                    j.get("title") or j.get("job_title"),
                    j.get("location"),
                    j.get("apply_link") or j.get("link") or j.get("application_url") or ""
                )

        # Atomically save and activate new search results in persistent SQLite database
        search_payload = {
            "skills": skills,
            "roles": job_roles,
            "role_matches": role_matches,
            "missing_skills": missing_skills,
            "market_insights": market_insights,
            "jobs": jobs,
        }
        save_job_search_results(user_id, search_payload)

        # SUCCESS: Activate 60-second application cooldown and clear in_progress
        set_user_search_cooldown(user_id, cooldown_seconds=60)
        log_api_usage(
            user_id=user_id,
            service="serpapi",
            feature="job_search",
            key_fingerprint=get_key_fingerprint(serpapi_key),
            success=True,
        )

        return redirect(url_for("jobs_route"))

    except Exception as e:
        print(f"Error during job analysis search: {e}")
        # FAILURE: Clear search-in-progress, DO NOT set cooldown (failed search does not lock user)
        clear_user_search_cooldown(user_id, clear_cooldown=False)
        log_api_usage(
            user_id=user_id,
            service="serpapi",
            feature="job_search",
            key_fingerprint=get_key_fingerprint(serpapi_key),
            success=False,
            error_message=str(e),
        )

        err_str = str(e).lower()
        if "quota" in err_str or "limit" in err_str or "429" in err_str:
            update_key_history_status(
                user_id=user_id,
                service="serpapi",
                key_fingerprint=get_key_fingerprint(serpapi_key),
                status="Limit Reached",
                last_error_category="QUOTA_EXHAUSTED",
            )
            flash("Your SerpAPI search limit has been reached. Open Profile and update your SerpAPI key to continue.", "error")
            target_url = url_for("jobs_route") if has_current_job_search(user_id) else url_for("jobs_home")
            return redirect(target_url + "?api_error=serpapi_quota")
        elif "invalid" in err_str or "401" in err_str or "403" in err_str:
            update_key_history_status(
                user_id=user_id,
                service="serpapi",
                key_fingerprint=get_key_fingerprint(serpapi_key),
                status="Invalid",
                last_error_category="INVALID_KEY",
            )
            flash("Your SerpAPI key is invalid. Open Profile and update your SerpAPI key.", "error")
            target_url = url_for("jobs_route") if has_current_job_search(user_id) else url_for("jobs_home")
            return redirect(target_url + "?api_error=serpapi_invalid")
        else:
            flash("An error occurred during job search. Previous saved results have been preserved.", "error")

        if has_current_job_search(user_id):
            return redirect(url_for("jobs_route"))
        return redirect(url_for("jobs_home"))


@app.route("/jobs", methods=["GET"])
@login_required
def jobs_route():
    """
    Main read-only Job Listing page.
    Loads and renders the current authenticated user's persisted/cached search results.
    Does NOT call SerpAPI or Gemini on page load, navigation, or refresh.
    If no previous search exists, redirects user to Jobs Home with prompt to start search.
    """
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    user_resume = get_user_resume(user_id, current_only=True)

    if not user_keys.get("has_serpapi"):
        flash("Please configure your SerpAPI key before accessing job listings.", "warning")
        return redirect(url_for("setup"))

    if not user_resume:
        flash("Please upload your resume to view job matches.", "info")
        return redirect(url_for("upload_resume_page"))

    current_search = get_current_job_search(user_id)
    if not current_search:
        flash("No job analysis available yet. Click 'Analyze Resume & Find Jobs' to start your search.", "info")
        return redirect(url_for("jobs_home"))

    skills = current_search.get("skills", [])
    roles = current_search.get("roles", [])
    role_matches = current_search.get("role_matches", [])
    missing_skills = current_search.get("missing_skills", {})
    market_insights = current_search.get("market_insights", [])
    jobs = current_search.get("jobs", [])

    # Keep session in sync for modal agent generators
    session["missing_skills_data"] = missing_skills

    # Attach dynamic user-specific saved and applied states
    saved_job_ids = get_saved_job_ids(user_id) if user_id else set()
    applied_job_ids = get_applied_job_ids(user_id) if user_id else set()

    for j in jobs:
        jid = j.get("job_id")
        j["is_saved"] = jid in saved_job_ids
        j["is_applied"] = jid in applied_job_ids

    user = get_current_user()

    return render_template(
        "results.html",
        skills=skills,
        roles=roles,
        role_matches=role_matches,
        missing_skills=missing_skills,
        market_insights=market_insights,
        jobs=jobs,
        user=user,
        keys=user_keys,
        resume=user_resume,
    )



# --- SAVED JOBS & APPLIED JOBS ROUTES ---

@app.route("/saved-jobs")
@login_required
def saved_jobs_page():
    """
    Renders the Saved Jobs page displaying only the current authenticated user's saved jobs.
    """
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    user_resume = get_user_resume(user_id, current_only=True)
    user = get_current_user()

    saved_jobs = get_saved_jobs(user_id)
    applied_ids = get_applied_job_ids(user_id)

    for j in saved_jobs:
        j["is_saved"] = True
        j["is_applied"] = j["job_id"] in applied_ids

    return render_template(
        "saved_jobs.html",
        saved_jobs=saved_jobs,
        user=user,
        keys=user_keys,
        resume=user_resume,
    )


@app.route("/save-job", methods=["POST"])
@login_required
def save_job_route():
    """
    Saves a job to the database for the current authenticated user.
    """
    user_id = session.get("user_id")
    data = request.get_json(silent=True) or request.form.to_dict() or {}

    if not data:
        return jsonify({"success": False, "error": "No job data provided."}), 400

    try:
        job_id = save_job(user_id, data)
        if job_id:
            return jsonify({
                "success": True,
                "saved": True,
                "job_id": job_id,
                "message": "Job saved successfully."
            })
        return jsonify({"success": False, "error": "Unable to save job."}), 500
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/remove-saved-job", methods=["POST"])
@login_required
def remove_saved_job_route():
    """
    Removes a saved job from the database for the current authenticated user.
    """
    user_id = session.get("user_id")
    data = request.get_json(silent=True) or request.form.to_dict() or {}
    job_id = data.get("job_id")
    if not job_id:
        return jsonify({"success": False, "error": "job_id is required."}), 400

    try:
        removed = remove_saved_job(user_id, job_id)
        return jsonify({
            "success": True,
            "saved": False,
            "job_id": job_id,
            "removed": removed,
            "message": "Job removed from saved jobs."
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/applied-jobs")
@login_required
def applied_jobs_page():
    """
    Renders the Applied Jobs page displaying only the current authenticated user's applied jobs.
    """
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    user_resume = get_user_resume(user_id, current_only=True)
    user = get_current_user()

    applied_jobs = get_applied_jobs(user_id)
    saved_ids = get_saved_job_ids(user_id)

    for j in applied_jobs:
        j["is_applied"] = True
        j["is_saved"] = j["job_id"] in saved_ids

    return render_template(
        "applied_jobs.html",
        applied_jobs=applied_jobs,
        user=user,
        keys=user_keys,
        resume=user_resume,
    )


@app.route("/mark-applied", methods=["POST"])
@login_required
def mark_applied_route():
    """
    Marks a job as applied for the current authenticated user in the database.
    """
    user_id = session.get("user_id")
    data = request.get_json(silent=True) or request.form.to_dict() or {}

    if not data:
        return jsonify({"success": False, "error": "No job data provided."}), 400

    try:
        job_id = mark_job_applied(user_id, data)
        if job_id:
            return jsonify({
                "success": True,
                "applied": True,
                "job_id": job_id,
                "message": "Job marked as applied."
            })
        return jsonify({"success": False, "error": "Unable to mark job as applied."}), 500
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/mark-not-applied", methods=["POST"])
@login_required
def mark_not_applied_route():
    """
    Removes the applied status for a job from the database for the current authenticated user.
    """
    user_id = session.get("user_id")
    data = request.get_json(silent=True) or request.form.to_dict() or {}
    job_id = data.get("job_id")
    if not job_id:
        return jsonify({"success": False, "error": "job_id is required."}), 400

    try:
        removed = mark_job_not_applied(user_id, job_id)
        return jsonify({
            "success": True,
            "applied": False,
            "job_id": job_id,
            "removed": removed,
            "message": "Job marked as not applied."
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# --- SKILLBRIDGE.AI COURSES & CAREER PATHS ROUTES ---

@app.route("/courses", methods=["GET"])
@login_required
def courses_route():
    """
    Renders the canonical All Courses page displaying all 9 career-focused courses
    with real per-user progress percentages and completion metrics.
    Requires authenticated user session.
    """
    user_id = session.get("user_id")
    user = get_current_user()
    user_keys = get_user_api_keys(user_id, decrypted=True) if user_id else {}
    user_resume = get_user_resume(user_id, current_only=True) if user_id else None
    courses = get_all_courses()
    user_progress = get_all_courses_progress_for_user(user_id) if user_id else {}

    return render_template(
        "courses.html",
        courses=courses,
        user_progress=user_progress,
        user=user,
        keys=user_keys,
        resume=user_resume,
    )


@app.route("/courses/<course_id>", methods=["GET"])
@login_required
def course_detail_route(course_id):
    """
    Renders the detailed This Course / Course Overview learning experience for a single selected career course.
    Requires authenticated user session.
    Calculates real user progress and dynamically binds curriculum topics, YouTube embeds,
    descriptions, and learning points.
    Safely redirects with flash notice if course_id is invalid.
    """
    course = get_course_by_id(course_id)
    if not course:
        flash("The requested course could not be found.", "warning")
        return redirect(url_for("courses_route"))

    user_id = session.get("user_id")
    user = get_current_user()
    user_keys = get_user_api_keys(user_id, decrypted=True) if user_id else {}
    user_resume = get_user_resume(user_id, current_only=True) if user_id else None
    all_courses = get_all_courses()

    topics = course.get("topics", [])
    req_topic_id = request.args.get("topic")
    active_topic = None
    if req_topic_id:
        active_topic = get_topic_by_id(course["id"], req_topic_id)
    if not active_topic and topics:
        active_topic = topics[0]

    progress_stats = get_course_progress_stats(user_id, course["id"], total_topics=len(topics))
    completed_topic_ids = get_user_completed_topic_ids(user_id, course["id"])

    return render_template(
        "course_detail.html",
        course=course,
        topics=topics,
        active_topic=active_topic,
        progress_stats=progress_stats,
        completed_topic_ids=completed_topic_ids,
        all_courses=all_courses,
        user=user,
        keys=user_keys,
        resume=user_resume,
    )


@app.route("/api/courses/<course_id>/topics/<topic_id>/complete", methods=["POST"])
@csrf.exempt
@login_required
def api_toggle_topic_completion(course_id, topic_id):
    """
    Persists topic completion state for the authenticated user.
    Strictly isolated per user session.
    Returns authoritative updated progress metrics.
    """
    course = get_course_by_id(course_id)
    if not course:
        return jsonify({"success": False, "error": "Course not found"}), 404

    topic = get_topic_by_id(course["id"], topic_id)
    if not topic:
        return jsonify({"success": False, "error": "Topic not found"}), 404

    user_id = session.get("user_id")
    data = request.get_json(silent=True) or {}
    
    # If completed explicitly passed, use that; otherwise toggle or mark true
    if "completed" in data:
        completed = bool(data["completed"])
    else:
        # Default action is marking completed
        completed = True

    stats = set_topic_completion(user_id, course["id"], str(topic["id"]), completed=completed)
    completed_ids = get_user_completed_topic_ids(user_id, course["id"])

    return jsonify({
        "success": True,
        "course_id": course["id"],
        "topic_id": str(topic["id"]),
        "completed": completed,
        "completed_count": stats.get("completed_count", 0),
        "total_topics": stats.get("total_topics", 0),
        "percentage": stats.get("percentage", 0),
        "completed_ids": list(completed_ids),
    })


@app.route("/api/courses/<course_id>/progress", methods=["GET"])
@login_required
def api_get_course_progress(course_id):
    """
    Returns the real-time course progress stats for the authenticated user.
    """
    course = get_course_by_id(course_id)
    if not course:
        return jsonify({"success": False, "error": "Course not found"}), 404

    user_id = session.get("user_id")
    topics = course.get("topics", [])
    stats = get_course_progress_stats(user_id, course["id"], total_topics=len(topics))
    completed_ids = get_user_completed_topic_ids(user_id, course["id"])

    return jsonify({
        "success": True,
        "course_id": course["id"],
        "progress": stats,
        "completed_ids": list(completed_ids),
    })



# --- PROFILE PAGE & ACCOUNT MANAGEMENT ROUTES ---

@app.route("/profile")
@login_required
def profile_route():
    """
    Renders the main authenticated Profile page.
    Displays user account info, masked API keys, SerpAPI usage & limits,
    Gemini usage & status, key history, resume management, contact info, and bottom nav.
    """
    user_id = session.get("user_id")
    user = get_current_user()
    user_keys = get_user_api_keys(user_id, decrypted=True)
    user_resume = get_user_resume(user_id, current_only=True)

    # 1. Fetch SerpAPI Account Data if key is present
    serpapi_account = None
    if user_keys.get("has_serpapi") and user_keys.get("serpapi_key"):
        serpapi_account = fetch_serpapi_account_info(user_keys["serpapi_key"])
        # Update key history with freshest metrics
        key_fp = user_keys.get("serpapi_fingerprint")
        if key_fp:
            update_key_history_status(
                user_id=user_id,
                service="serpapi",
                key_fingerprint=key_fp,
                status=serpapi_account.get("status_display", "ACTIVE"),
                plan_name=serpapi_account.get("plan_name"),
                renewal_date=serpapi_account.get("plan_renewal_date"),
                last_known_usage=serpapi_account.get("this_month_usage"),
                last_known_limit=serpapi_account.get("searches_per_month"),
                last_known_hourly_usage=serpapi_account.get("this_hour_searches"),
                last_known_hourly_limit=serpapi_account.get("account_rate_limit_per_hour"),
                remaining_searches=(
                    serpapi_account.get("plan_searches_left")
                    if serpapi_account.get("plan_searches_left") is not None
                    else serpapi_account.get("total_searches_left")
                ),
            )

    # 2. Fetch Gemini Usage Summary
    gemini_usage = get_gemini_usage_summary(
        user_id=user_id,
        key_fingerprint=user_keys.get("gemini_fingerprint"),
        configured_model=DEFAULT_GEMINI_MODEL,
    )

    # 3. Fetch API Key History
    api_key_history = get_api_key_history(user_id)

    return render_template(
        "profile.html",
        user=user,
        keys=user_keys,
        resume=user_resume,
        serpapi_account=serpapi_account,
        gemini_usage=gemini_usage,
        api_key_history=api_key_history,
        gemini_model=DEFAULT_GEMINI_MODEL,
    )


@app.route("/profile/change-password", methods=["POST"])
@login_required
def change_password_route():
    """
    Securely updates the authenticated user's password in the database.
    Enforces password matching, policy validation, and secure password hashing.
    """
    user_id = session.get("user_id")
    new_password = request.form.get("new_password", "").strip()
    confirm_password = request.form.get("confirm_password", "").strip()

    if not new_password and request.is_json:
        data = request.get_json(silent=True) or {}
        new_password = data.get("new_password", "").strip()
        confirm_password = data.get("confirm_password", "").strip()

    # 1. Reject empty fields
    if not new_password or not confirm_password:
        msg = "Please enter and confirm your new password."
        if request.is_json:
            return jsonify({"success": False, "error": msg}), 400
        flash(msg, "error")
        return redirect(url_for("profile_route") + "#profile")

    # 2. Check password matching
    if new_password != confirm_password:
        msg = "Passwords do not match."
        if request.is_json:
            return jsonify({"success": False, "error": msg}), 400
        flash(msg, "error")
        return redirect(url_for("profile_route") + "#profile")

    # 3. Validate against application password policy
    is_valid, err_msg = validate_password(new_password)
    if not is_valid:
        if request.is_json:
            return jsonify({"success": False, "error": err_msg}), 400
        flash(err_msg, "error")
        return redirect(url_for("profile_route") + "#profile")

    # 4. Hash and update password in database
    password_hash = generate_password_hash(new_password)
    update_user_password(user_id, password_hash)

    msg = "Password updated successfully."
    if request.is_json:
        return jsonify({"success": True, "message": msg})
    flash(msg, "success")
    return redirect(url_for("profile_route") + "#profile")


@app.route("/profile/update-api-keys", methods=["POST"])
@login_required
def update_api_keys_route():
    """
    Updates the authenticated user's stored SerpAPI and/or Gemini API keys in the database.
    Preserves existing keys when a field is left blank.
    Immediately queries SerpAPI Account API for new SerpAPI key to capture plan & quota.
    Validates Gemini key minimally.
    """
    user_id = session.get("user_id")
    serpapi_key = request.form.get("serpapi_key", "").strip()
    gemini_key = request.form.get("gemini_key", "").strip()

    if not serpapi_key and not gemini_key and request.is_json:
        data = request.get_json(silent=True) or {}
        serpapi_key = data.get("serpapi_key", "").strip()
        gemini_key = data.get("gemini_key", "").strip()

    # Save to encrypted DB storage (preserves existing key if empty/None)
    save_user_api_keys(
        user_id=user_id,
        serpapi_key=serpapi_key if serpapi_key else None,
        gemini_api_key=gemini_key if gemini_key else None,
    )

    # If new serpapi key was provided, fetch Account API info and update history
    if serpapi_key:
        serp_info = fetch_serpapi_account_info(serpapi_key)
        serp_fp = get_key_fingerprint(serpapi_key)
        update_key_history_status(
            user_id=user_id,
            service="serpapi",
            key_fingerprint=serp_fp,
            status=serp_info.get("status_display", "ACTIVE"),
            plan_name=serp_info.get("plan_name"),
            renewal_date=serp_info.get("plan_renewal_date"),
            last_known_usage=serp_info.get("this_month_usage"),
            last_known_limit=serp_info.get("searches_per_month"),
            last_known_hourly_usage=serp_info.get("this_hour_searches"),
            last_known_hourly_limit=serp_info.get("account_rate_limit_per_hour"),
            remaining_searches=(
                serp_info.get("plan_searches_left")
                if serp_info.get("plan_searches_left") is not None
                else serp_info.get("total_searches_left")
            ),
        )

    # If new gemini key was provided, perform minimal validation
    if gemini_key:
        is_valid, clf = validate_gemini_key(gemini_key, DEFAULT_GEMINI_MODEL)
        gem_fp = get_key_fingerprint(gemini_key)
        if not is_valid:
            status_val = (
                "Invalid"
                if clf.get("error_category") == "INVALID_KEY"
                else "Model Unavailable"
                if clf.get("error_category") == "MODEL_UNAVAILABLE"
                else "Limit Reached"
                if clf.get("error_category") == "RATE_LIMIT"
                else "Unavailable"
            )
            update_key_history_status(
                user_id=user_id,
                service="gemini",
                key_fingerprint=gem_fp,
                status=status_val,
                last_error_category=clf.get("error_category"),
            )
        else:
            update_key_history_status(
                user_id=user_id,
                service="gemini",
                key_fingerprint=gem_fp,
                status="ACTIVE",
            )

    # Refresh session keys from database
    fresh_keys = get_user_api_keys(user_id, decrypted=True)
    if fresh_keys.get("serpapi_key"):
        session["serpapi_key"] = fresh_keys["serpapi_key"]
    if fresh_keys.get("gemini_api_key"):
        session["gemini_api_key"] = fresh_keys["gemini_api_key"]

    msg = "API keys updated successfully."
    if request.is_json:
        return jsonify({
            "success": True,
            "message": msg,
            "serpapi_masked": fresh_keys.get("serpapi_masked", ""),
            "gemini_masked": fresh_keys.get("gemini_masked", ""),
            "serpapi_masked_bullet": fresh_keys.get("serpapi_masked_bullet", ""),
            "gemini_masked_bullet": fresh_keys.get("gemini_masked_bullet", ""),
            "has_serpapi": fresh_keys.get("has_serpapi", False),
            "has_gemini": fresh_keys.get("has_gemini", False),
        })
    flash(msg, "success")
    return redirect(url_for("profile_route") + "#api-keys")


@app.route("/profile/delete-api-key", methods=["POST"])
@login_required
def delete_api_key_route():
    """
    Clears a specific API key (SerpAPI or Gemini) for the authenticated user.
    """
    user_id = session.get("user_id")
    data = request.get_json(silent=True) or request.form.to_dict() or {}
    service = data.get("service", "").strip().lower()

    if not service:
        if request.is_json:
            return jsonify({"success": False, "error": "Service parameter is required."}), 400
        flash("Service parameter is required.", "error")
        return redirect(url_for("profile_route") + "#api-keys")

    delete_user_api_key(user_id, service)

    # Clear from session
    if "serp" in service:
        session.pop("serpapi_key", None)
    elif "gemini" in service:
        session.pop("gemini_api_key", None)

    msg = f"{service.capitalize()} API key removed."
    if request.is_json:
        return jsonify({"success": True, "message": msg})
    flash(msg, "info")
    return redirect(url_for("profile_route") + "#api-keys")


@app.route("/profile/update-resume", methods=["POST"])
@login_required
def update_resume_route():
    """
    Uploads and activates a new Main Profile Resume for the authenticated user.
    Uses storage_manager abstraction (Supabase Storage if configured, or local uploads/ fallback).
    Atomically validates, replaces the previous main resume, and updates DB & session state.
    Leaves ATS Resume completely untouched.
    """
    user_id = session.get("user_id")

    if "resume" not in request.files:
        flash("No file was uploaded. Please choose a resume PDF.", "error")
        return redirect(url_for("profile_route") + "#resume")

    resume_file = request.files["resume"]
    is_valid, err_msg, content_bytes = validate_resume_document(
        resume_file, resume_file.filename, is_profile=True
    )
    if not is_valid:
        flash(err_msg, "error")
        return redirect(url_for("profile_route") + "#resume")

    try:
        upload_res = upload_resume_file(
            user_id=user_id,
            file_data=content_bytes,
            original_filename=resume_file.filename,
            resume_type="main",
            content_type="application/pdf",
        )
    except Exception as e:
        print(f"[Main Resume Upload Error]: {e}")
        flash("Could not save resume file. Please try again.", "error")
        return redirect(url_for("profile_route") + "#resume")

    canonical_main_path = upload_res["storage_path"]
    file_size = upload_res["file_size"]

    # Extract skills & insights using temporary processing context (auto cleaned up)
    with temp_resume_context(file_bytes=content_bytes, filename=resume_file.filename) as temp_proc_path:
        skills = extract_skills(temp_proc_path)
        job_roles = map_skills_to_roles(skills)
        missing_skills = get_missing_skills(skills, job_roles)
        resume_text = get_pdf_text(temp_proc_path)

    extracted_payload = {
        "skills": skills,
        "roles": job_roles,
        "missing_skills": missing_skills,
        "resume_text_preview": resume_text[:3000] if resume_text else "",
    }

    # Save to user_resumes (marks previous resumes as is_current=0)
    save_user_resume(
        user_id=user_id,
        original_filename=resume_file.filename,
        stored_filename=upload_res["stored_filename"],
        file_path=canonical_main_path,
        file_size=file_size,
        file_type="pdf",
        extracted_data=extracted_payload,
        processing_status="completed",
    )

    # Immediately update active session state for subsequent job matching
    session["current_resume_path"] = canonical_main_path
    session["missing_skills_data"] = missing_skills

    # Recalculate Job ATS scores for current job search results using the new Main Profile Resume
    try:
        from job_matcher import calculate_job_ats_score, parse_main_profile_resume, invalidate_resume_cache
        invalidate_resume_cache(user_id)
        curr_search = get_current_job_search(user_id)
        if curr_search and curr_search.get("jobs"):
            new_structured = parse_main_profile_resume({
                "user_id": user_id,
                "file_path": canonical_main_path,
                "skills": skills,
                "roles": job_roles,
                "resume_text": resume_text,
            })
            updated_jobs = []
            for j in curr_search["jobs"]:
                j_copy = dict(j)
                ats_res = calculate_job_ats_score(new_structured, j_copy)
                j_copy["match_percent"] = ats_res["final_score"]
                j_copy["match_score"] = ats_res["final_score"]
                j_copy["matching_skills"] = ats_res.get("matching_skills", [])
                j_copy["ats_score_data"] = ats_res
                updated_jobs.append(j_copy)
            updated_jobs.sort(key=lambda x: x.get("match_percent", 0), reverse=True)

            save_job_search_results(user_id, {
                "skills": skills,
                "roles": job_roles,
                "role_matches": calculate_role_matches(skills, job_roles),
                "missing_skills": missing_skills,
                "market_insights": curr_search.get("market_insights", []),
                "jobs": updated_jobs,
            })
    except Exception as e:
        print(f"[Resume Update Score Recalculation Notice]: {e}")

    flash("Resume updated successfully.", "success")
    return redirect(url_for("profile_route") + "#resume")


@app.route("/profile/download-resume")
@login_required
def download_resume_route():
    """
    Securely serves the current authenticated user's stored resume as a download attachment.
    Strictly enforces user ownership: authenticated_user_id == resume.user_id.
    Works seamlessly across Supabase Storage and Local Storage.
    """
    user_id = session.get("user_id")
    user_resume = get_user_resume(user_id, current_only=True)

    if not user_resume or not user_resume.get("file_path"):
        flash("No stored resume found to download.", "error")
        return redirect(url_for("profile_route"))

    file_path = user_resume["file_path"]
    download_name = user_resume.get("original_filename") or "resume.pdf"
    ext = os.path.splitext(download_name)[1].lower() or ".pdf"
    mimetype = "application/pdf" if ext == ".pdf" else "application/octet-stream"

    # 1. Try download via storage abstraction
    content_bytes, _, detected_mimetype = download_resume_file(
        user_id=user_id,
        resume_type="main",
        storage_path=file_path,
    )
    if content_bytes:
        return send_file(
            io.BytesIO(content_bytes),
            as_attachment=True,
            download_name=download_name,
            mimetype=detected_mimetype or mimetype,
        )

    # 2. Local filesystem fallback
    if not os.path.isabs(file_path):
        cand_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), file_path)
        if os.path.exists(cand_path):
            file_path = cand_path

    if os.path.exists(file_path) and os.path.isfile(file_path):
        return send_file(
            file_path,
            as_attachment=True,
            download_name=download_name,
            mimetype=mimetype,
        )

    flash("The resume file could not be found on the server.", "error")
    return redirect(url_for("profile_route"))


@app.route("/resume-analyzer")
@login_required
def resume_analyzer_route():
    flash("Resume Analyzer & ATS scoring will be available in the next release.", "info")
    return redirect(url_for("jobs_home"))


@app.route("/ai-agent")
@login_required
def ai_agent_route():
    flash("AI Career Agent will be available in the next release.", "info")
    return redirect(url_for("jobs_home"))




# --- API ENDPOINTS FOR AGENTS ---

@app.route("/generate-cover-letter", methods=["POST"])
@login_required
def generate_cover_letter_route():
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    gemini_key = user_keys.get("gemini_api_key") or session.get("gemini_api_key")

    if not gemini_key:
        return jsonify({
            "success": False,
            "error": "Gemini API key is not configured. Open Profile to set up your key.",
            "error_category": "MISSING_KEY",
            "popup_title": "Gemini Key Required",
            "action_button_text": "Update Gemini API Key",
            "action_url": "/profile#api-keys",
            "content": "Error: Gemini API Key missing. Please configure it in Profile.",
        })
    
    resume_path = session.get("current_resume_path")
    if not resume_path:
        user_resume = get_user_resume(user_id, current_only=True)
        if user_resume:
            resume_path = user_resume.get("file_path")
            session["current_resume_path"] = resume_path

    resume_text = get_pdf_text(resume_path, user_id=user_id) if resume_path else ""
    data = request.get_json(silent=True) or {}

    try:
        letter = agent_write_cover_letter(
            resume_text,
            data.get("job_title", "Software Engineer"),
            data.get("company", "Tech Company"),
            data.get("location", "India"),
            gemini_key,
        )
        log_api_usage(
            user_id=user_id,
            service="gemini",
            feature="cover_letter",
            key_fingerprint=get_key_fingerprint(gemini_key),
            success=True,
            model=DEFAULT_GEMINI_MODEL,
        )
        return jsonify({"success": True, "content": letter})
    except Exception as e:
        err_clf = classify_gemini_error(e)
        log_api_usage(
            user_id=user_id,
            service="gemini",
            feature="cover_letter",
            key_fingerprint=get_key_fingerprint(gemini_key),
            success=False,
            error_category=err_clf["error_category"],
            error_message=str(e),
            model=DEFAULT_GEMINI_MODEL,
            retry_after_seconds=err_clf.get("retry_after"),
        )
        # Update key history if limit reached or invalid
        if err_clf["error_category"] in ("RATE_LIMIT", "QUOTA_EXHAUSTED"):
            update_key_history_status(
                user_id=user_id,
                service="gemini",
                key_fingerprint=get_key_fingerprint(gemini_key),
                status="Limit Reached",
                last_error_category="RATE_LIMIT",
            )
        elif err_clf["error_category"] == "INVALID_KEY":
            update_key_history_status(
                user_id=user_id,
                service="gemini",
                key_fingerprint=get_key_fingerprint(gemini_key),
                status="Invalid",
                last_error_category="INVALID_KEY",
            )
        elif err_clf["error_category"] == "MODEL_UNAVAILABLE":
            update_key_history_status(
                user_id=user_id,
                service="gemini",
                key_fingerprint=get_key_fingerprint(gemini_key),
                status="Model Unavailable",
                last_error_category="MODEL_UNAVAILABLE",
            )

        return jsonify({
            "success": False,
            "error": err_clf["user_message"],
            "error_category": err_clf["error_category"],
            "popup_title": err_clf["popup_title"],
            "action_button_text": err_clf["action_button_text"],
            "action_url": err_clf["action_url"],
            "retry_after": err_clf.get("retry_after"),
            "content": f"Error: {err_clf['user_message']}",
        })


@app.route("/generate-learning-plan", methods=["POST"])
@login_required
def generate_learning_plan_route():
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    gemini_key = user_keys.get("gemini_api_key") or session.get("gemini_api_key")
    serp_key = user_keys.get("serpapi_key") or session.get("serpapi_key")

    if not gemini_key:
        return jsonify({
            "success": False,
            "error": "Gemini API key is not configured. Open Profile to set up your key.",
            "error_category": "MISSING_KEY",
            "popup_title": "Gemini Key Required",
            "action_button_text": "Update Gemini API Key",
            "action_url": "/profile#api-keys",
            "content": "Error: Gemini API Key missing. Please configure it in Profile.",
        })
    
    data = request.get_json(silent=True) or {}
    role = data.get("role")
    
    all_missing = session.get("missing_skills_data")
    if not all_missing:
        user_resume = get_user_resume(user_id, current_only=True)
        if user_resume and user_resume.get("extracted_data"):
            all_missing = user_resume["extracted_data"].get("missing_skills", {})
            session["missing_skills_data"] = all_missing
        else:
            all_missing = {}

    skills_for_role = all_missing.get(role, [])
    
    if not skills_for_role:
        return jsonify({"success": True, "content": "No missing skills found for this role! You are good to go."})

    try:
        plan = agent_gap_analysis(
            role,
            skills_for_role,
            serp_key,
            gemini_key,
        )
        log_api_usage(
            user_id=user_id,
            service="gemini",
            feature="gap_analysis",
            key_fingerprint=get_key_fingerprint(gemini_key),
            success=True,
            model=DEFAULT_GEMINI_MODEL,
        )
        return jsonify({"success": True, "content": plan})
    except Exception as e:
        err_clf = classify_gemini_error(e)
        log_api_usage(
            user_id=user_id,
            service="gemini",
            feature="gap_analysis",
            key_fingerprint=get_key_fingerprint(gemini_key),
            success=False,
            error_category=err_clf["error_category"],
            error_message=str(e),
            model=DEFAULT_GEMINI_MODEL,
            retry_after_seconds=err_clf.get("retry_after"),
        )
        if err_clf["error_category"] in ("RATE_LIMIT", "QUOTA_EXHAUSTED"):
            update_key_history_status(
                user_id=user_id,
                service="gemini",
                key_fingerprint=get_key_fingerprint(gemini_key),
                status="Limit Reached",
                last_error_category="RATE_LIMIT",
            )
        elif err_clf["error_category"] == "INVALID_KEY":
            update_key_history_status(
                user_id=user_id,
                service="gemini",
                key_fingerprint=get_key_fingerprint(gemini_key),
                status="Invalid",
                last_error_category="INVALID_KEY",
            )
        elif err_clf["error_category"] == "MODEL_UNAVAILABLE":
            update_key_history_status(
                user_id=user_id,
                service="gemini",
                key_fingerprint=get_key_fingerprint(gemini_key),
                status="Model Unavailable",
                last_error_category="MODEL_UNAVAILABLE",
            )

        return jsonify({
            "success": False,
            "error": err_clf["user_message"],
            "error_category": err_clf["error_category"],
            "popup_title": err_clf["popup_title"],
            "action_button_text": err_clf["action_button_text"],
            "action_url": err_clf["action_url"],
            "retry_after": err_clf.get("retry_after"),
            "content": f"Error: {err_clf['user_message']}",
        })


# --- REAL-TIME API USAGE & COOLDOWN STATUS JSON ENDPOINTS ---

@app.route("/api/search-cooldown-status")
@login_required
def api_search_cooldown_status():
    """
    Returns the authenticated user's real-time search cooldown status and remaining seconds.
    """
    user_id = session.get("user_id")
    cd = get_user_search_cooldown(user_id)
    return jsonify({
        "in_progress": cd.get("in_progress", False),
        "is_cooldown": cd.get("is_cooldown", False),
        "remaining_seconds": cd.get("remaining_seconds", 0),
        "cooldown_until": cd.get("cooldown_until"),
    })


@app.route("/api/serpapi-usage")
@login_required
def api_serpapi_usage():
    """
    Returns authoritative SerpAPI account and usage data via the official Account API.
    """
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    serp_key = user_keys.get("serpapi_key")
    if not serp_key:
        return jsonify({
            "configured": False,
            "status_display": "UNCONFIGURED",
            "message": "SerpAPI key is not configured.",
        })
    info = fetch_serpapi_account_info(serp_key)
    info["configured"] = True
    return jsonify(info)


@app.route("/api/gemini-usage")
@login_required
def api_gemini_usage():
    """
    Returns local Gemini usage metrics and status for the authenticated user.
    """
    user_id = session.get("user_id")
    user_keys = get_user_api_keys(user_id, decrypted=True)
    gem_key = user_keys.get("gemini_api_key")
    if not gem_key:
        return jsonify({
            "configured": False,
            "status": "UNCONFIGURED",
            "message": "Gemini API key is not configured.",
        })
    summary = get_gemini_usage_summary(
        user_id=user_id,
        key_fingerprint=user_keys.get("gemini_fingerprint"),
        configured_model=DEFAULT_GEMINI_MODEL,
    )
    summary["configured"] = True
    return jsonify(summary)


# --- JARVIS GLOBAL AI CHATBOT ROUTE ---

@app.route("/api/ai/chat", methods=["POST"])
@csrf.exempt
def api_ai_chat():
    """
    Global AI Assistant Chat Endpoint for JARVIS.
    Handles message processing, safe page context, attachment extraction,
    and returns structured JSON responses.
    Strictly derives authenticated user identity from the server-side session.
    """
    user_id = session.get("user_id") if session.get("authenticated") else None

    # Handle JSON or multipart form-data
    attachment_text = None
    if request.is_json:
        data = request.get_json(silent=True) or {}
        message = data.get("message", "")
        page_context = {
            "page": data.get("page", "home"),
            "course_id": data.get("course_id"),
            "topic_id": data.get("topic_id"),
            "job_id": data.get("job_id"),
        }
    else:
        message = request.form.get("message", "")
        page_context = {
            "page": request.form.get("page", "home"),
            "course_id": request.form.get("course_id"),
            "topic_id": request.form.get("topic_id"),
            "job_id": request.form.get("job_id"),
        }

        # Check for uploaded attachment
        if "attachment" in request.files:
            file = request.files["attachment"]
            if file and file.filename:
                ok, text_or_err = extract_attachment_text(file, file.filename)
                if ok:
                    attachment_text = text_or_err
                else:
                    return jsonify({
                        "success": False,
                        "reply": f"Attachment error: {text_or_err}",
                        "error": text_or_err,
                    }), 400

    try:
        response_data = process_chat_message(
            user_id=user_id,
            message=message,
            page_context=page_context,
            attachment_text=attachment_text,
        )
        return jsonify(response_data)
    except Exception as e:
        print(f"[Jarvis AI Chatbot Error]: {e}")
        return jsonify({
            "success": False,
            "reply": "Jarvis is temporarily unavailable. Please try again later.",
            "error": str(e),
        }), 500


if __name__ == "__main__":
    app.run(debug=True)
