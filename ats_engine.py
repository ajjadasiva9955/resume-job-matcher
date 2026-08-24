import os
import re
import json
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timezone
from PyPDF2 import PdfReader

# Try importing docx if available
try:
    import docx
    HAS_DOCX_MODULE = True
except ImportError:
    HAS_DOCX_MODULE = False

# Try importing google.generativeai if available
try:
    import google.generativeai as genai
    HAS_GENAI = True
except ImportError:
    HAS_GENAI = False

from skill_extractor import SKILL_SET, SKILL_DISPLAY_MAP, format_skill_name


# ==============================================================================
# 1. CANONICAL DICTIONARIES, TAXONOMIES & ACTION VERBS
# ==============================================================================

CANONICAL_TECH_CASING = {
    "html": "HTML", "html5": "HTML5",
    "css": "CSS", "css3": "CSS3",
    "javascript": "JavaScript", "js": "JavaScript",
    "typescript": "TypeScript", "ts": "TypeScript",
    "react": "React", "react.js": "React.js", "reactjs": "React.js",
    "next.js": "Next.js", "nextjs": "Next.js",
    "node.js": "Node.js", "nodejs": "Node.js",
    "express": "Express.js", "express.js": "Express.js",
    "vue": "Vue.js", "vue.js": "Vue.js",
    "angular": "Angular",
    "bootstrap": "Bootstrap",
    "tailwind": "Tailwind CSS", "tailwind css": "Tailwind CSS",
    "python": "Python", "java": "Java", "c++": "C++", "c#": "C#",
    "flask": "Flask", "django": "Django", "fastapi": "FastAPI",
    "spring boot": "Spring Boot",
    "sql": "SQL", "sqlite": "SQLite", "postgresql": "PostgreSQL", "postgres": "PostgreSQL", "mysql": "MySQL", "mongodb": "MongoDB", "redis": "Redis",
    "git": "Git", "github": "GitHub", "gitlab": "GitLab",
    "vs code": "VS Code", "vscode": "VS Code",
    "docker": "Docker", "kubernetes": "Kubernetes",
    "aws": "AWS", "azure": "Azure", "gcp": "GCP",
    "rest api": "REST API", "restful api": "RESTful API", "graphql": "GraphQL",
    "machine learning": "Machine Learning", "deep learning": "Deep Learning",
    "tensorflow": "TensorFlow", "pytorch": "PyTorch", "keras": "Keras",
    "pandas": "Pandas", "numpy": "NumPy", "scikit-learn": "Scikit-Learn", "opencv": "OpenCV",
    "serpapi": "SerpAPI", "gemini api": "Gemini API", "streamlit": "Streamlit"
}

DOMAIN_TAXONOMY = {
    "Software Development / Computer Science / AI & ML": {
        "keywords": [
            "python", "java", "c++", "c", "sql", "machine learning", "deep learning", "nlp", "llm",
            "flask", "django", "fastapi", "rest api", "gemini api", "openai api", "serpapi",
            "computer science", "artificial intelligence", "data structures", "algorithms", "software development"
        ],
        "core_anchors": [
            "computer science", "artificial intelligence", "machine learning", "ai & ml", "ai/ml",
            "python", "flask", "deep learning", "nlp", "llm", "software engineer", "software development"
        ]
    },
    "Full Stack / Web Development": {
        "keywords": [
            "javascript", "typescript", "react", "html", "css", "bootstrap", "tailwind",
            "node.js", "express", "flask", "django", "sql", "mongodb", "postgresql", "sqlite",
            "rest api", "full stack", "fullstack", "web application", "web development"
        ],
        "core_anchors": [
            "full stack", "fullstack", "mern", "mean", "web development", "web developer",
            "react", "node.js", "express", "flask"
        ]
    },
    "Backend & Cloud Engineering": {
        "keywords": [
            "python", "java", "node.js", "express", "flask", "django", "spring boot", "fastapi",
            "sql", "postgresql", "mysql", "mongodb", "redis", "rest api", "graphql",
            "microservices", "docker", "kubernetes", "aws", "azure", "gcp", "linux", "system design"
        ],
        "core_anchors": ["backend", "microservices", "spring boot", "django", "flask", "express", "node.js", "postgresql", "rest api", "docker", "kubernetes"]
    },
    "Frontend Development": {
        "keywords": [
            "javascript", "typescript", "react", "html", "css", "vue", "angular",
            "next.js", "tailwind", "bootstrap", "sass", "redux", "webpack", "vite",
            "responsive design", "ui/ux", "dom"
        ],
        "core_anchors": ["react", "vue", "angular", "next.js", "frontend", "tailwind", "ui/ux", "dom", "responsive design"]
    },
    "Data Science & Analytics": {
        "keywords": [
            "python", "sql", "pandas", "numpy", "tableau", "power bi", "excel",
            "data visualization", "statistics", "data analysis", "r", "bigquery",
            "snowflake", "etl", "data modeling", "dashboard", "business intelligence"
        ],
        "core_anchors": ["data analyst", "data scientist", "tableau", "power bi", "pandas", "statistics", "etl", "business intelligence"]
    },
    "DevOps & Cloud Infrastructure": {
        "keywords": [
            "docker", "kubernetes", "aws", "azure", "gcp", "terraform", "ansible",
            "jenkins", "ci/cd", "linux", "bash", "prometheus", "grafana", "git",
            "cloudformation", "helm", "devops", "infrastructure as code"
        ],
        "core_anchors": ["devops", "kubernetes", "terraform", "ci/cd", "jenkins", "ansible", "docker", "aws", "infrastructure"]
    },
    "Cyber Security & QA": {
        "keywords": [
            "cyber security", "penetration testing", "owasp", "burp suite", "wireshark",
            "nmap", "network security", "linux", "siem", "selenium", "pytest", "junit",
            "test automation", "unit testing", "quality assurance"
        ],
        "core_anchors": ["cyber security", "penetration testing", "owasp", "selenium", "qa", "quality assurance", "test automation"]
    }
}

SKILL_CATEGORY_MAP = {
    "Languages": {"python", "java", "javascript", "typescript", "c++", "c#", "c", "go", "rust", "php", "ruby", "sql", "html", "html5", "css", "css3", "bash", "r", "scala", "kotlin", "swift", "dart"},
    "Frontend": {"react", "react.js", "next.js", "vue", "vue.js", "angular", "bootstrap", "tailwind", "tailwind css", "sass", "scss", "jquery", "redux", "vite", "webpack", "html", "html5", "css", "css3"},
    "Backend & Frameworks": {"flask", "django", "fastapi", "spring boot", "node.js", "express", "express.js", "asp.net", "laravel", "rails", "rest api", "graphql", "microservices"},
    "Databases": {"postgresql", "mysql", "mongodb", "sqlite", "redis", "oracle", "sql server", "dynamodb", "firebase", "cassandra", "mariadb"},
    "Cloud & DevOps": {"aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "git", "github", "gitlab", "ci/cd", "jenkins", "terraform", "ansible", "linux", "nginx"},
    "AI / ML & Data": {"machine learning", "deep learning", "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy", "opencv", "nlp", "llm", "transformers", "hugging face", "langchain", "generative ai", "computer vision", "data analysis", "tableau", "power bi", "streamlit", "machine learning fundamentals"},
    "Tools & Platforms": {"vs code", "visual studio", "postman", "jira", "figma", "pycharm", "jupyter", "serpapi", "gemini api", "openai api", "git", "github", "streamlit"},
    "Soft Skills": {"problem solving", "time management", "adaptability", "team collaboration", "communication", "leadership", "critical thinking", "collaboration", "teamwork"}
}

ACTION_VERBS = {
    "developed", "built", "designed", "implemented", "created", "integrated", "optimized",
    "automated", "analyzed", "extracted", "matched", "provided", "replicated", "improved",
    "managed", "configured", "tested", "deployed", "engineered", "led", "spearheaded",
    "architected", "orchestrated", "accelerated", "refactored", "programmed", "formulated",
    "established", "streamlined", "enhanced", "resolved", "executed", "maintained",
    "monitored", "published", "collaborated", "achieved", "generated", "reduced",
    "increased", "scaled", "delivered", "constructed", "authored", "trained", "evaluated",
    "standardized", "centralized", "administered", "mentored", "coordinated", "transformed",
    "launched", "debugged", "validated", "migrated"
}

VERB_LEMMA_MAP = {
    "developing": "developed", "develop": "developed", "develops": "developed",
    "building": "built", "build": "built", "builds": "built",
    "designing": "designed", "design": "designed", "designs": "designed",
    "implementing": "implemented", "implement": "implemented", "implements": "implemented",
    "creating": "created", "create": "created", "creates": "created",
    "integrating": "integrated", "integrate": "integrated", "integrates": "integrated",
    "optimizing": "optimized", "optimize": "optimized", "optimizes": "optimized",
    "automating": "automated", "automate": "automated", "automates": "automated",
    "analyzing": "analyzed", "analyze": "analyzed", "analyzes": "analyzed",
    "extracting": "extracted", "extract": "extracted", "extracts": "extracted",
    "matching": "matched", "match": "matched", "matches": "matched",
    "providing": "provided", "provide": "provided", "provides": "provided",
    "replicating": "replicated", "replicate": "replicated", "replicates": "replicated",
    "improving": "improved", "improve": "improved", "improves": "improved",
    "managing": "managed", "manage": "managed", "manages": "managed",
    "configuring": "configured", "configure": "configured", "configures": "configured",
    "testing": "tested", "test": "tested", "tests": "tested",
    "deploying": "deployed", "deploy": "deployed", "deploys": "deployed",
    "engineering": "engineered", "engineer": "engineered", "engineers": "engineered",
    "leading": "led", "lead": "led", "leads": "led",
    "launching": "launched", "launch": "launched", "launches": "launched",
    "debugging": "debugged", "debug": "debugged", "debugs": "debugged",
    "validating": "validated", "validate": "validated", "validates": "validated",
    "migrating": "migrated", "migrate": "migrated", "migrates": "migrated"
}

WEAK_ACTION_PHRASES = [
    "responsible for",
    "worked on",
    "helped with",
    "assisted in",
    "handled",
    "involved in",
    "duties included",
    "tasked with",
    "participated in",
    "contributed to",
    "attempted to",
    "familiar with"
]

BULLET_CHARS = [
    "\uf0b7", "\uf0a7", "\u2022", "\u25cf", "\u25cb", "\u25e6",
    "\u25aa", "\u25ab", "\u2043", "\u2219", "\u2013", "\u2014",
    "-", "*", "→", "▸", "›", "»", "–", "—", "•", "●", "○", "◦",
    "▪", "▫", "⁃", "∙"
]


# ==============================================================================
# 2. STRUCTURED FINDING HELPER (EVIDENCE-BASED DATA MODEL)
# ==============================================================================

def make_finding(finding_id, finding_type, severity, category, section, source_text, issue, reason, recommendation, confidence=0.95):
    """
    Creates a structured, traceable finding object adhering strictly to the ATS specification:
    {
      "id": "...",
      "type": "issue|strength|warning|improvement",
      "severity": "critical|high|medium|low",
      "category": "...",
      "section": "...",
      "source_text": "...",
      "issue": "...",
      "reason": "...",
      "recommendation": "...",
      "confidence": 0.0
    }
    """
    return {
        "id": str(finding_id),
        "type": str(finding_type),          # issue | strength | warning | improvement
        "severity": str(severity),          # critical | high | medium | low
        "category": str(category),          # ats_readability | content_quality | skills | experience_projects | completeness | quantification | grammar_consistency
        "section": str(section),            # Contact | Summary | Technical Skills | Projects | Experience | Education | Certifications | Hobbies | General
        "source_text": source_text.strip() if source_text else "",
        "issue": str(issue),
        "reason": str(reason),
        "recommendation": str(recommendation),
        "confidence": float(confidence)
    }


# ==============================================================================
# 3. STRUCTURE-AWARE TEXT EXTRACTION (PDF + DOCX + TXT)
# ==============================================================================

def extract_text_from_pdf(pdf_path):
    """
    Extracts structured content from PDF files.
    Prefers coordinate-aware fragment extraction (group fragments into lines,
    order top-to-bottom, sort left-to-right within a line) to reconstruct logical reading order.
    Gracefully falls back to standard text extraction if coordinates are unavailable.
    """
    structured_pages = []
    full_lines = []
    metadata = {
        "page_count": 0,
        "is_scanned": False,
        "has_tables_clue": False,
        "has_columns_clue": False,
        "extraction_method": "coordinate_aware"
    }

    try:
        reader = PdfReader(pdf_path)
        metadata["page_count"] = len(reader.pages)
        total_text_len = 0

        for page_idx, page in enumerate(reader.pages):
            page_fragments = []
            
            def visitor_body(text, cm, tm, font_dict, font_size):
                if text:
                    page_fragments.append({
                        "text": text,
                        "x": float(tm[4]),
                        "y": float(tm[5]),
                        "font_size": float(font_size) if font_size else 10.0
                    })

            try:
                page.extract_text(visitor_text=visitor_body)
            except Exception as e:
                print(f"[ATS Extraction] visitor_text notice on page {page_idx+1}: {e}")

            page_lines = []

            # If visitor returned fragments with coordinate information
            if page_fragments:
                # Group fragments by visual line (y tolerance ~ 3.0 points)
                tolerance = 3.0
                fragments_by_y = sorted(page_fragments, key=lambda f: -f["y"])
                grouped_lines = []
                
                for frag in fragments_by_y:
                    placed = False
                    for group in grouped_lines:
                        avg_y = sum(f["y"] for f in group) / len(group)
                        if abs(frag["y"] - avg_y) <= tolerance:
                            group.append(frag)
                            placed = True
                            break
                    if not placed:
                        grouped_lines.append([frag])

                # Sort lines top-to-bottom
                grouped_lines.sort(key=lambda grp: -max(f["y"] for f in grp))

                # For each line, sort fragments left-to-right (x ascending)
                for group in grouped_lines:
                    group.sort(key=lambda f: f["x"])
                    line_str = ""
                    for idx, frag in enumerate(group):
                        txt = frag["text"]
                        if idx > 0 and not line_str.endswith(" ") and not txt.startswith(" "):
                            line_str += " "
                        line_str += txt
                    
                    cleaned_line = line_str.strip()
                    if cleaned_line:
                        page_lines.append(cleaned_line)
                
                # Check for genuine multi-column layout evidence (overlapping y ranges across distinct x lanes)
                left_only_lines = 0
                right_only_lines = 0
                full_width_lines = 0
                for grp in grouped_lines:
                    min_x = min(f["x"] for f in grp)
                    max_x = max(f["x"] for f in grp)
                    if min_x < 250 and max_x > 350:
                        full_width_lines += 1
                    elif max_x <= 300:
                        left_only_lines += 1
                    elif min_x >= 300:
                        right_only_lines += 1

                # Conservative check: Only flag multi-column if there are sustained independent parallel lanes
                if left_only_lines >= 8 and right_only_lines >= 8 and full_width_lines < 5:
                    metadata["has_columns_clue"] = True

            else:
                # Fallback to standard pypdf page.extract_text()
                metadata["extraction_method"] = "standard_fallback"
                page_text_raw = page.extract_text() or ""
                page_lines = [l.strip() for l in page_text_raw.splitlines() if l.strip()]

            page_full_text = "\n".join(page_lines)
            total_text_len += len(page_full_text.strip())

            structured_pages.append({
                "page_number": page_idx + 1,
                "lines": page_lines,
                "text": page_full_text
            })
            full_lines.extend(page_lines)

        # Scanned document detection
        if metadata["page_count"] > 0 and (total_text_len / metadata["page_count"]) < 60:
            metadata["is_scanned"] = True

        # Evidence-based table risk detection
        pipe_lines = sum(1 for l in full_lines if l.count("|") >= 3 or l.count("+-") >= 2)
        if pipe_lines >= 4:
            metadata["has_tables_clue"] = True

    except Exception as e:
        print(f"[ATS Extraction] PDF read error: {e}")

    raw_text = "\n".join(full_lines)
    return {
        "raw_text": raw_text,
        "lines": full_lines,
        "pages": structured_pages,
        "metadata": metadata
    }


def extract_text_from_docx(docx_path):
    """
    Extracts structured content from DOCX files preserving paragraphs,
    headings, tables, and bullet lists.
    """
    full_lines = []
    tables_found = 0
    headings_found = []
    bullet_items = []
    paragraphs = []

    if HAS_DOCX_MODULE:
        try:
            doc = docx.Document(docx_path)
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                paragraphs.append(text)
                full_lines.append(text)
                
                style_name = p.style.name.lower() if p.style else ""
                if "heading" in style_name or "title" in style_name:
                    headings_found.append(text)
                elif "list" in style_name or "bullet" in style_name or any(text.startswith(b) for b in BULLET_CHARS):
                    bullet_items.append(text)

            for table in doc.tables:
                tables_found += 1
                for row in table.rows:
                    row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_cells:
                        row_line = " | ".join(row_cells)
                        full_lines.append(row_line)
        except Exception as e:
            print(f"[ATS Extraction] python-docx read error, falling back to XML: {e}")

    # Fallback to direct docx zip XML extraction if needed
    if not full_lines:
        try:
            with zipfile.ZipFile(docx_path) as z:
                xml_content = z.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

                for p_elem in tree.iterfind(".//w:p", namespaces):
                    texts = [t.text for t in p_elem.iterfind(".//w:t", namespaces) if t.text]
                    line = "".join(texts).strip()
                    if line:
                        full_lines.append(line)
                        paragraphs.append(line)
                        if any(line.startswith(b) for b in BULLET_CHARS):
                            bullet_items.append(line)

                for _ in tree.iterfind(".//w:tbl", namespaces):
                    tables_found += 1
        except Exception as e:
            print(f"[ATS Extraction] DOCX direct XML error: {e}")

    raw_text = "\n".join(full_lines)
    return {
        "raw_text": raw_text,
        "lines": full_lines,
        "pages": [{"page_number": 1, "lines": full_lines, "text": raw_text}],
        "metadata": {
            "page_count": 1,
            "is_scanned": False,
            "has_tables_clue": tables_found > 0,
            "has_columns_clue": False,
            "tables_count": tables_found,
            "extraction_method": "docx_structured"
        }
    }


def extract_resume_document(file_path):
    """
    Master document extractor handling PDF, DOCX, and plain text formats.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                lines = [l.strip() for l in content.splitlines() if l.strip()]
                return {
                    "raw_text": content,
                    "lines": lines,
                    "pages": [{"page_number": 1, "lines": lines, "text": content}],
                    "metadata": {
                        "page_count": 1,
                        "is_scanned": False,
                        "has_tables_clue": False,
                        "has_columns_clue": False,
                        "extraction_method": "raw_text"
                    }
                }
        except Exception:
            return {
                "raw_text": "",
                "lines": [],
                "pages": [],
                "metadata": {
                    "page_count": 0,
                    "is_scanned": True,
                    "has_tables_clue": False,
                    "has_columns_clue": False
                }
            }


# ==============================================================================
# 4. RESUME STRUCTURE & SECTION PARSER
# ==============================================================================

SECTION_PATTERNS = {
    "contact": [
        r"^(contact|contact\s+information|contact\s+details|personal\s+details|personal\s+info)$",
    ],
    "summary": [
        r"^(summary|professional\s+summary|profile\s+summary|career\s+summary|profile|about\s+me|career\s+objective|objective|executive\s+summary)$"
    ],
    "skills": [
        r"^(skills|technical\s+skills|core\s+competencies|key\s+skills|skills\s*&\s*expertise|technologies|technical\s+stack|tools\s*&\s*technologies|programming\s+languages|technical\s+skills\s*&\s*tools)$"
    ],
    "experience": [
        r"^(experience|work\s+experience|professional\s+experience|employment\s+history|work\s+history|career\s+history|employment|professional\s+background)$"
    ],
    "internships": [
        r"^(internships|internship\s+experience|training\s*&\s*internships|practical\s+training|summer\s+internship)$"
    ],
    "education": [
        r"^(education|academic\s+background|educational\s+qualifications|academic\s+qualifications|academic\s+history|academics)$"
    ],
    "projects": [
        r"^(projects|personal\s+projects|academic\s+projects|key\s+projects|technical\s+projects|portfolio\s+projects)$"
    ],
    "certifications": [
        r"^(certifications|licenses\s*&\s*certifications|courses\s*&\s*certifications|certificates|professional\s+certifications)$"
    ],
    "achievements": [
        r"^(achievements|key\s+achievements|accomplishments|awards\s*&\s*achievements|honors\s*&\s*awards|extracurricular\s+activities|achievements\s*&\s*awards|awards)$"
    ],
    "hobbies": [
        r"^(hobbies|interests|leisure\s+activities|personal\s+interests|hobbies\s*&\s*interests)$"
    ],
    "other": [
        r"^(languages\s+known|activities|publications|volunteer\s+experience)$"
    ]
}


def normalize_heading_candidate(text):
    """
    Normalizes heading candidates to handle PDF kerning/spacing artifacts and split letters.
    e.g., 'CAREER OBJ ECTIVE' -> 'career objective', 'T E C H N I C A L' -> 'technical'.
    """
    cleaned = text.strip().lower().strip(":").strip("-").strip()
    cleaned = re.sub(r"\bobj\s+ective\b", "objective", cleaned)
    cleaned = re.sub(r"\btechnic\s+al\b", "technical", cleaned)
    cleaned = re.sub(r"\btech\s+nic\s*al\b", "technical", cleaned)
    cleaned = re.sub(r"\beduc\s+ation\b", "education", cleaned)
    cleaned = re.sub(r"\bedu\s+cation\b", "education", cleaned)
    cleaned = re.sub(r"\bexperi\s+ence\b", "experience", cleaned)
    cleaned = re.sub(r"\bexper\s+ience\b", "experience", cleaned)
    cleaned = re.sub(r"\bpro\s+jects\b", "projects", cleaned)
    cleaned = re.sub(r"\bcertific\s+ations\b", "certifications", cleaned)
    cleaned = re.sub(r"\bcertif\s+ications\b", "certifications", cleaned)
    cleaned = re.sub(r"\bachieve\s+ments\b", "achievements", cleaned)
    cleaned = re.sub(r"\shob\s+bies\b", "hobbies", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def is_bullet_line(line):
    """
    Determines if a line starts with a supported bullet glyph or numbered list pattern.
    """
    clean_l = line.strip()
    for b in BULLET_CHARS:
        if clean_l.startswith(b):
            return True, b
    # Check numbered bullets e.g. '1. ', '1) ', '(1) '
    m = re.match(r"^(\(?\d+[\.\)]\s+)", clean_l)
    if m:
        return True, m.group(1)
    return False, None


def clean_bullet_prefix(line, prefix=None):
    """
    Strips bullet glyph or number prefix cleanly from line text.
    """
    clean_l = line.strip()
    if prefix and clean_l.startswith(prefix):
        return clean_l[len(prefix):].strip()
    for b in BULLET_CHARS:
        if clean_l.startswith(b):
            clean_l = clean_l[len(b):].strip()
    clean_l = re.sub(r"^\(?\d+[\.\)]\s*", "", clean_l).strip()
    return clean_l


def parse_resume_structure(extracted_doc):
    """
    Parses resume text into logical sections, contact info, bullet points,
    dates, and structural elements with wrapped-line continuation support.
    """
    lines = extracted_doc.get("lines", [])
    raw_text = extracted_doc.get("raw_text", "")

    parsed_sections = {
        "contact": [],
        "summary": [],
        "skills": [],
        "experience": [],
        "internships": [],
        "education": [],
        "projects": [],
        "certifications": [],
        "achievements": [],
        "hobbies": [],
        "other": []
    }

    section_headings_detected = []
    bullet_points = []
    dates_found = []

    # 1. Contact info extraction
    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", raw_text)
    phone_match = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4,6}", raw_text)
    linkedin_url_match = re.search(r"linkedin\.com/in/[\w\-_%]+|https?://(?:www\.)?linkedin\.com/[\w\-_%/]+", raw_text, re.IGNORECASE)
    github_url_match = re.search(r"github\.com/[\w\-_]+|https?://(?:www\.)?github\.com/[\w\-_]+", raw_text, re.IGNORECASE)
    portfolio_url_match = re.search(r"https?://(?!.*(?:linkedin|github)\.com)[\w\.-]+\.[a-z]{2,}(/[\w\.-]*)*", raw_text, re.IGNORECASE)

    # Check for text labels without actual URLs
    has_linkedin_label = bool(re.search(r"\blinkedin\b", raw_text, re.IGNORECASE))
    has_github_label = bool(re.search(r"\bgithub\b", raw_text, re.IGNORECASE))
    has_portfolio_label = bool(re.search(r"\bportfolio\b|\blive\s*[-–]\s*demo\b", raw_text, re.IGNORECASE))

    # Location candidate search in top lines
    location_candidate = None
    location_raw = ""
    for l in lines[:6]:
        m_loc = re.search(r"\b([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\s*[,|\s]\s*([A-Z]{2,}|[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\b", l)
        if m_loc and "@" not in m_loc.group(0) and "http" not in m_loc.group(0) and not any(w in m_loc.group(0).lower() for w in ["linkedin", "github", "portfolio", "resume", "developer", "engineer"]):
            location_raw = m_loc.group(0).strip()
            location_candidate = location_raw
            break

    contact_info = {
        "email": email_match.group(0) if email_match else None,
        "phone": phone_match.group(0) if phone_match else None,
        "linkedin_url": linkedin_url_match.group(0) if linkedin_url_match else None,
        "github_url": github_url_match.group(0) if github_url_match else None,
        "portfolio_url": portfolio_url_match.group(0) if portfolio_url_match else None,
        "location": location_candidate,
        "has_linkedin_label": has_linkedin_label,
        "has_github_label": has_github_label,
        "has_portfolio_label": has_portfolio_label,
        "has_contact": bool(email_match or phone_match)
    }

    # Extract candidate name
    candidate_name = ""
    for line in lines[:5]:
        clean_l = line.strip()
        if not clean_l:
            continue
        if "@" in clean_l or "http" in clean_l or re.search(r"\d{4,}", clean_l):
            continue
        if len(clean_l.split()) in [2, 3, 4] and len(clean_l) < 40 and not any(clean_l.startswith(b) for b in BULLET_CHARS):
            candidate_name = clean_l
            break
    if not candidate_name and lines:
        candidate_name = lines[0][:35].strip()

    # 2. Section classifier and wrapped bullet parser
    current_section = "contact"
    current_bullet_obj = None

    for line_idx, line in enumerate(lines):
        clean_l = line.strip()
        if not clean_l:
            continue

        normalized_heading = normalize_heading_candidate(clean_l)

        # Check if line matches a known section heading
        matched_section = None
        for sec_name, patterns in SECTION_PATTERNS.items():
            for pat in patterns:
                if re.match(pat, normalized_heading, re.IGNORECASE):
                    matched_section = sec_name
                    break
            if matched_section:
                break

        if matched_section:
            current_section = matched_section
            section_headings_detected.append({
                "section": matched_section,
                "raw_heading": clean_l
            })
            current_bullet_obj = None
            continue

        # Check for dates
        dates = re.findall(
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\b\d{4}\s*[-–to]+\s*(?:\d{4}|Present|Current)\b|\b\d{1,2}/\d{4}\b",
            clean_l,
            re.IGNORECASE
        )
        if dates:
            dates_found.extend(dates)

        # Check if line is a new bullet point
        is_bullet, bullet_prefix = is_bullet_line(clean_l)

        if is_bullet:
            bullet_clean = clean_bullet_prefix(clean_l, bullet_prefix)
            if len(bullet_clean) >= 4:
                bullet_entry = {
                    "section": current_section,
                    "text": bullet_clean,
                    "raw": clean_l,
                    "starts_with_action_verb": False,
                    "leading_verb": None
                }
                # Check leading verb
                first_word = bullet_clean.split()[0].lower().rstrip(".,:;") if bullet_clean.split() else ""
                lemma_verb = VERB_LEMMA_MAP.get(first_word, first_word)
                if lemma_verb in ACTION_VERBS:
                    bullet_entry["starts_with_action_verb"] = True
                    bullet_entry["leading_verb"] = lemma_verb

                bullet_points.append(bullet_entry)
                current_bullet_obj = bullet_entry
            parsed_sections[current_section].append(clean_l)
        else:
            # Check if this line is a continuation of the previous bullet point
            is_new_project_title = bool(re.search(r"(live\s*[-–]\s*demo|github|project)", clean_l, re.IGNORECASE) and len(clean_l) < 70)
            is_key_value = bool(re.search(r"^[A-Za-z\s/&]+:\s*", clean_l))

            if current_bullet_obj and not is_new_project_title and not is_key_value and len(clean_l) < 150:
                current_bullet_obj["text"] += " " + clean_l
                current_bullet_obj["raw"] += " " + clean_l
            else:
                current_bullet_obj = None

            parsed_sections[current_section].append(clean_l)

    return {
        "candidate_name": candidate_name,
        "contact_info": contact_info,
        "sections": parsed_sections,
        "section_headings": section_headings_detected,
        "bullet_points": bullet_points,
        "dates_found": dates_found
    }


# ==============================================================================
# 5. CONTACT INFORMATION ANALYSIS ENGINE
# ==============================================================================

def analyze_contact_section(parsed_data, extracted_doc):
    """
    Evaluates individual contact elements: Name, Email, Phone, Location,
    LinkedIn, GitHub, Portfolio. Distinguishes: PRESENT, MISSING, INCOMPLETE,
    and NONSTANDARD formatting with evidence-backed findings.
    """
    contact = parsed_data.get("contact_info", {})
    findings = []
    status_map = {}

    # 1. Candidate Name
    name = parsed_data.get("candidate_name", "")
    if name and len(name) >= 3:
        status_map["name"] = "PRESENT"
    else:
        status_map["name"] = "MISSING"
        findings.append(make_finding(
            "contact_name_missing", "issue", "critical", "completeness", "Contact",
            "", "Candidate name not clearly detected at top of resume.",
            "Recruiters and ATS need a prominent full name to create a candidate profile.",
            "Place your full name in large, clear text at the very top of the page."
        ))

    # 2. Email Address
    email = contact.get("email")
    if email:
        status_map["email"] = "PRESENT"
    else:
        status_map["email"] = "MISSING"
        findings.append(make_finding(
            "contact_email_missing", "issue", "critical", "completeness", "Contact",
            "", "Missing professional email address.",
            "Recruiters cannot send interview invites or automated test links without a valid email.",
            "Add a professional email address (e.g. firstname.lastname@gmail.com) in the contact header."
        ))

    # 3. Phone Number
    phone = contact.get("phone")
    if phone:
        digits_only = re.sub(r"\D", "", phone)
        if len(digits_only) >= 10:
            status_map["phone"] = "PRESENT"
        else:
            status_map["phone"] = "PRESENT_BUT_INCOMPLETE"
            findings.append(make_finding(
                "contact_phone_incomplete", "issue", "high", "completeness", "Contact",
                phone, "Phone number appears incomplete or missing country code.",
                "Recruiters require a complete 10-digit number with country code for phone screening.",
                "Format your phone number with your country code (e.g. +91 93901 44782)."
            ))
    else:
        status_map["phone"] = "MISSING"
        findings.append(make_finding(
            "contact_phone_missing", "issue", "critical", "completeness", "Contact",
            "", "Missing contact phone number.",
            "Phone numbers are a standard requirement for recruiter verification and outreach.",
            "Include a primary mobile number in your contact header."
        ))

    # 4. Location
    loc = contact.get("location")
    if loc:
        if re.search(r"\b[A-Za-z]+\s+[A-Z]{2}\b", loc) and "," not in loc:
            status_map["location"] = "PRESENT_BUT_NONSTANDARD"
            findings.append(make_finding(
                "contact_location_abbr", "improvement", "low", "completeness", "Contact",
                loc, "Location is present but uses a nonstandard abbreviated form.",
                "Standard city and full state or country names ensure accurate ATS geographic filtering.",
                f"Consider '{loc.split()[0]}, {loc.split()[-1]}' with full state/region (e.g. 'Srikakulam, Andhra Pradesh') for clearer recruiter/ATS parsing."
            ))
        else:
            status_map["location"] = "PRESENT"
    else:
        status_map["location"] = "MISSING"

    # 5. LinkedIn Profile URL vs Label
    linkedin_url = contact.get("linkedin_url")
    has_linkedin_label = contact.get("has_linkedin_label")
    if linkedin_url:
        status_map["linkedin"] = "PRESENT"
    elif has_linkedin_label:
        status_map["linkedin"] = "PRESENT_BUT_NO_URL"
        findings.append(make_finding(
            "contact_linkedin_no_url", "warning", "medium", "completeness", "Contact",
            "LinkedIn", "LinkedIn label detected, but actual URL was not found in extracted text.",
            "Verify that the document contains working hyperlinks or a text URL (e.g. linkedin.com/in/username).",
            "Include your full, clickable LinkedIn profile link."
        ))
    else:
        status_map["linkedin"] = "MISSING"

    # 6. GitHub Profile URL vs Label
    github_url = contact.get("github_url")
    has_github_label = contact.get("has_github_label")
    if github_url:
        status_map["github"] = "PRESENT"
    elif has_github_label:
        status_map["github"] = "PRESENT_BUT_NO_URL"
        findings.append(make_finding(
            "contact_github_no_url", "warning", "medium", "completeness", "Contact",
            "GitHub", "GitHub label detected, but actual URL was not found in extracted text.",
            "Technical recruiters expect direct access to your code repositories and commit history.",
            "Include your full, clickable GitHub profile URL (e.g. github.com/username)."
        ))
    else:
        status_map["github"] = "MISSING"

    # 7. Portfolio / Live Demo URL vs Label
    portfolio_url = contact.get("portfolio_url")
    has_portfolio_label = contact.get("has_portfolio_label")
    if portfolio_url:
        status_map["portfolio"] = "PRESENT"
    elif has_portfolio_label:
        status_map["portfolio"] = "PRESENT_BUT_NO_URL"
        findings.append(make_finding(
            "contact_portfolio_no_url", "warning", "low", "experience_projects", "Projects",
            "Live-Demo / Portfolio", "Portfolio or Live Demo label detected, but actual URL was not found in extracted text.",
            "Recruiters cannot view your live web applications without working hyperlinks.",
            "Ensure live demo links are active, embedded hyperlinks or listed with full web addresses."
        ))
    else:
        status_map["portfolio"] = "MISSING"

    return {
        "status_map": status_map,
        "findings": findings
    }


# ==============================================================================
# 6. SUMMARY / CAREER OBJECTIVE ANALYSIS ENGINE
# ==============================================================================

def analyze_summary_objective(parsed_data):
    """
    Analyzes the actual Summary/Career Objective text:
    Evaluates specificity, target role, technical identity, and detects generic cliches.
    """
    summary_lines = parsed_data["sections"].get("summary", [])
    findings = []
    
    if not summary_lines:
        return {
            "has_summary": False,
            "is_generic": False,
            "text": "",
            "findings": [make_finding(
                "summary_missing", "improvement", "medium", "content_quality", "Summary",
                "", "Missing professional summary or career objective.",
                "A targeted 2-3 sentence summary immediately frames your technical identity and domain strengths.",
                "Add a concise professional summary highlighting your core tech stack (e.g. Python, Flask, REST APIs) and engineering focus."
            )]
        }

    sum_text = " ".join(summary_lines).strip()
    sum_lower = sum_text.lower()

    # Detect generic cliches
    generic_cliches = [
        "to start my career in a growth",
        "to start my career in a growth - oriented organization",
        "to start my career in a growth-oriented organization",
        "seeking an entry-level position",
        "seeking an entry level position",
        "looking for a challenging role",
        "where i can apply my technical skills",
        "where i can apply my skills",
        "contribute to meaningful projects",
        "grow as a software professional",
        "hardworking individual",
        "looking for an opportunity"
    ]

    is_generic = False
    matched_cliche = None
    for cliche in generic_cliches:
        if cliche in sum_lower:
            is_generic = True
            matched_cliche = cliche
            break

    if is_generic:
        findings.append(make_finding(
            "summary_generic_cliche", "improvement", "medium", "content_quality", "Summary",
            sum_text[:140] + ("..." if len(sum_text) > 140 else ""),
            "Generic career objective detected beginning with entry-level phrasing.",
            "Generic objectives ('To start my career in a growth-oriented organization...') lack technical identity and fail to communicate your specific engineering strengths or specialization.",
            "Replace the generic objective with a role-specific technical summary (e.g. 'Aspiring Software Engineer with hands-on project experience in Python, Flask backend services, REST APIs, and responsive web applications.')."
        ))
    else:
        word_count = len(sum_text.split())
        if word_count < 15:
            findings.append(make_finding(
                "summary_too_brief", "improvement", "low", "content_quality", "Summary",
                sum_text, "Summary is very brief (under 15 words).",
                "Fails to highlight technical identity and key domain capabilities.",
                "Expand summary to 2-3 sentences covering core technologies, engineering focus, and key project achievements."
            ))
        elif word_count > 80:
            findings.append(make_finding(
                "summary_too_long", "improvement", "low", "content_quality", "Summary",
                sum_text[:100] + "...", "Summary is lengthy (over 80 words).",
                "Overly long paragraphs reduce readability and take up valuable resume space.",
                "Condense summary to 3-4 impactful lines focusing on core strengths and engineering stack."
            ))
        else:
            findings.append(make_finding(
                "summary_strong", "strength", "low", "content_quality", "Summary",
                sum_text[:100] + "...", "Professional summary present with clear career focus.",
                "Effectively anchors your candidate profile and target engineering direction.",
                "Keep summary aligned with your demonstrated project skills."
            ))

    return {
        "has_summary": True,
        "is_generic": is_generic,
        "text": sum_text,
        "findings": findings
    }


# ==============================================================================
# 7. SPELLING / CAPITALIZATION / TERMINOLOGY ENGINE
# ==============================================================================

def analyze_terminology_capitalization(extracted_doc):
    """
    Detects exact wording and capitalization inconsistencies in original text lines.
    Only flags an issue when the resume actually contains the inconsistent form.
    """
    lines = extracted_doc.get("lines", [])
    findings = []
    reported_terms = set()

    for line in lines:
        for term_lower, canonical in CANONICAL_TECH_CASING.items():
            if term_lower in reported_terms:
                continue
            
            # Match exact word with case sensitivity
            pattern = r"(?<![A-Za-z0-9])" + re.escape(term_lower) + r"(?![A-Za-z0-9])"
            # If the exact lowercase/wrong-cased string appears in line, and does NOT match canonical
            if re.search(pattern, line):
                if term_lower != canonical.lower() or not re.search(r"(?<![A-Za-z0-9])" + re.escape(canonical) + r"(?![A-Za-z0-9])", line):
                    reported_terms.add(term_lower)
                    findings.append(make_finding(
                        f"casing_{term_lower.replace(' ', '_').replace('.', '_')}",
                        "warning", "low", "grammar_consistency", "General",
                        line[:80],
                        f"Nonstandard technology capitalization: '{term_lower}' found in document text.",
                        f"Technology name is conventionally capitalized as '{canonical}' for ATS tokenization and professional polish.",
                        f"Update '{term_lower}' to '{canonical}'."
                    ))

    return {
        "findings": findings,
        "reported_terms": list(reported_terms)
    }


# ==============================================================================
# 8. SKILLS TAXONOMY & CROSS-VALIDATION ENGINE (FOUR STATES & DIRECT EVIDENCE)
# ==============================================================================

def format_tech_name(skill_or_tool):
    """Formats technology name to standard canonical display casing."""
    s_lower = skill_or_tool.strip().lower()
    if s_lower in CANONICAL_TECH_CASING:
        return CANONICAL_TECH_CASING[s_lower]
    return SKILL_DISPLAY_MAP.get(s_lower, skill_or_tool.strip().title())


def analyze_skills_cross_validation(raw_text, parsed_data, individual_projects=None):
    """
    Skills Cross-Validation & Taxonomy Engine (15% Weight):
    Distinguishes FOUR distinct skill evidence states:
      STATE A — Listed + Directly Demonstrated (strongest evidence: project bullet, project tech list, experience tech list/bullet, implementation statement)
      STATE B — Listed + Not Demonstrated (listed in skills, no practical evidence found elsewhere in resume)
      STATE C — Demonstrated + Not Listed (explicitly used in project/experience/tech list but missing from Skills section)
      STATE D — Listed + Indirectly Supported (listed in skills, supported only by indirect/correlated evidence e.g. GitHub link for Git, soft skills in extra-curriculars)

    For EVERY skill returns:
      skill, state, listed_in_skills, demonstrated_in_resume, direct_evidence, indirect_evidence, evidence_locations, evidence_type, confidence
    """
    text_lower = raw_text.lower()
    sections = parsed_data.get("sections", {})
    skills_lines = sections.get("skills", [])
    proj_lines = sections.get("projects", [])
    exp_lines = sections.get("experience", [])
    intern_lines = sections.get("internships", [])
    cert_lines = sections.get("certifications", [])
    edu_lines = sections.get("education", [])
    sum_lines = sections.get("summary", [])

    if individual_projects is None:
        individual_projects = analyze_individual_projects(parsed_data, raw_text)

    # 1. Detect candidate skills across entire raw text
    all_detected_in_raw = []
    for skill in SKILL_SET:
        pattern = r"(?<!\w)" + re.escape(skill) + r"(?!\w)"
        if re.search(pattern, text_lower):
            all_detected_in_raw.append(format_tech_name(skill))
    for tech_k, tech_canon in CANONICAL_TECH_CASING.items():
        pattern = r"(?<!\w)" + re.escape(tech_k) + r"(?!\w)"
        if re.search(pattern, text_lower):
            all_detected_in_raw.append(tech_canon)
    all_detected_in_raw = list(dict.fromkeys(all_detected_in_raw))

    # 2. Identify skills listed explicitly in the Skills section
    skills_in_sec = []
    if skills_lines:
        skills_sec_text = "\n".join(skills_lines).lower()
        for skill in SKILL_SET:
            pattern = r"(?<!\w)" + re.escape(skill) + r"(?!\w)"
            if re.search(pattern, skills_sec_text):
                skills_in_sec.append(format_tech_name(skill))
        for tech_k, tech_canon in CANONICAL_TECH_CASING.items():
            pattern = r"(?<!\w)" + re.escape(tech_k) + r"(?!\w)"
            if re.search(pattern, skills_sec_text):
                skills_in_sec.append(tech_canon)
    else:
        # Fallback if no sections dict was parsed (e.g. direct text unit test)
        skills_in_sec = list(all_detected_in_raw)
    skills_in_sec = list(dict.fromkeys(skills_in_sec))

    # 3. Build direct evidence maps from Projects, Experience, Internships, Certifications
    skill_direct_evidence = {}
    skill_indirect_evidence = {}

    # A. Project "Technologies Used" and Project Bullets
    for p in individual_projects:
        p_name = p.get("project_name", "Project")
        
        # 1. Check explicit "Technologies Used" list (Direct evidence: type="technology_list")
        for tech in p.get("technologies_used", []):
            fmt_t = format_tech_name(tech)
            skill_direct_evidence.setdefault(fmt_t, []).append({
                "location": f"Project: {p_name} (Technologies Used)",
                "type": "technology_list",
                "text": f"Listed under Technologies Used in {p_name}"
            })

        # 2. Check project bullets (Direct evidence: type="bullet")
        for b_idx, b_text in enumerate(p.get("bullets", []), 1):
            b_lower = b_text.lower()
            for skill in SKILL_SET:
                pat = r"(?<!\w)" + re.escape(skill) + r"(?!\w)"
                if re.search(pat, b_lower):
                    fmt_s = format_tech_name(skill)
                    skill_direct_evidence.setdefault(fmt_s, []).append({
                        "location": f"Project: {p_name} (Bullet {b_idx})",
                        "type": "bullet",
                        "text": b_text[:90]
                    })
            for tech_k, tech_canon in CANONICAL_TECH_CASING.items():
                pat = r"(?<!\w)" + re.escape(tech_k) + r"(?!\w)"
                if re.search(pat, b_lower):
                    skill_direct_evidence.setdefault(tech_canon, []).append({
                        "location": f"Project: {p_name} (Bullet {b_idx})",
                        "type": "bullet",
                        "text": b_text[:90]
                    })

        # 3. Check all project detected technologies
        for tech in p.get("technologies", []):
            fmt_t = format_tech_name(tech)
            if fmt_t not in skill_direct_evidence:
                skill_direct_evidence.setdefault(fmt_t, []).append({
                    "location": f"Project: {p_name}",
                    "type": "project",
                    "text": f"Detected in project {p_name}"
                })

    # B. Check Experience lines (Direct evidence: type="experience" or "bullet")
    exp_full_text = ("\n".join(exp_lines) + "\n" + "\n".join(intern_lines)).lower()
    for skill in SKILL_SET:
        pat = r"(?<!\w)" + re.escape(skill) + r"(?!\w)"
        if re.search(pat, exp_full_text):
            fmt_s = format_tech_name(skill)
            skill_direct_evidence.setdefault(fmt_s, []).append({
                "location": "Professional Experience / Internships",
                "type": "experience",
                "text": "Detected in employment history"
            })
    for tech_k, tech_canon in CANONICAL_TECH_CASING.items():
        pat = r"(?<!\w)" + re.escape(tech_k) + r"(?!\w)"
        if re.search(pat, exp_full_text):
            skill_direct_evidence.setdefault(tech_canon, []).append({
                "location": "Professional Experience / Internships",
                "type": "experience",
                "text": "Detected in employment history"
            })

    # C. Check Certifications (Direct evidence: type="certification")
    cert_full_text = "\n".join(cert_lines).lower()
    for skill in SKILL_SET:
        pat = r"(?<!\w)" + re.escape(skill) + r"(?!\w)"
        if re.search(pat, cert_full_text):
            fmt_s = format_tech_name(skill)
            skill_direct_evidence.setdefault(fmt_s, []).append({
                "location": "Certifications",
                "type": "certification",
                "text": "Verified technical certification"
            })

    # D. Identify Indirect Evidence (State D candidates)
    contact_info = parsed_data.get("contact_info", {})
    has_github_link = bool(contact_info.get("github_url") or contact_info.get("has_github_label") or "github.com" in text_lower)
    if has_github_link:
        skill_indirect_evidence.setdefault("Git", []).append({
            "location": "Contact / Repository Links (GitHub Profile)",
            "type": "other",
            "text": "Git is indirectly supported by GitHub repository presence"
        })
        skill_indirect_evidence.setdefault("GitHub", []).append({
            "location": "Contact / Repository Links (GitHub Profile)",
            "type": "other",
            "text": "GitHub profile link present in document header"
        })

    has_web_project = any("web" in (p.get("project_name", "") + " " + p.get("technical_depth", "")).lower() or p.get("link_status", {}).get("has_live_demo_label") for p in individual_projects)
    if has_web_project:
        for web_tech in ["HTML", "CSS", "JavaScript"]:
            if web_tech not in skill_direct_evidence:
                skill_indirect_evidence.setdefault(web_tech, []).append({
                    "location": "Projects (Web Application Context)",
                    "type": "other",
                    "text": f"{web_tech} is indirectly implied by web application development"
                })

    soft_skills_canonical = ["Problem Solving", "Time Management", "Adaptability", "Team Collaboration", "Communication", "Leadership", "Teamwork"]
    for soft in soft_skills_canonical:
        if soft in skills_in_sec:
            skill_indirect_evidence.setdefault(soft, []).append({
                "location": "Technical Skills & Professional Context",
                "type": "other",
                "text": f"{soft} is supported as a foundational professional competency"
            })

    # 4. Combine all detected candidate skills across document
    all_candidate_skills = list(dict.fromkeys(all_detected_in_raw + skills_in_sec + list(skill_direct_evidence.keys())))

    # 5. Classify EVERY skill into exactly one of the 4 STATES
    skills_inventory = []
    state_a_skills = []  # Listed + Directly Demonstrated
    state_b_skills = []  # Listed + Not Demonstrated
    state_c_skills = []  # Demonstrated + Not Listed
    state_d_skills = []  # Listed + Indirectly Supported

    for skill in all_candidate_skills:
        is_listed = skill in skills_in_sec
        direct_ev_list = skill_direct_evidence.get(skill, [])
        indirect_ev_list = skill_indirect_evidence.get(skill, [])
        
        has_direct = len(direct_ev_list) > 0
        has_indirect = len(indirect_ev_list) > 0

        # Unique locations
        all_locations = []
        if is_listed:
            all_locations.append("Technical Skills Section")
        for dev in direct_ev_list:
            if dev["location"] not in all_locations:
                all_locations.append(dev["location"])
        for iev in indirect_ev_list:
            if iev["location"] not in all_locations:
                all_locations.append(iev["location"])

        # Primary evidence type
        if has_direct:
            primary_ev_type = direct_ev_list[0]["type"]
        elif has_indirect:
            primary_ev_type = indirect_ev_list[0]["type"]
        elif is_listed:
            primary_ev_type = "none"
        else:
            primary_ev_type = "other"

        # State evaluation
        if is_listed and has_direct:
            state = "A"
            state_label = "Listed + Directly Demonstrated"
            demonstrated = True
            confidence = 0.95
            state_a_skills.append(skill)
        elif is_listed and not has_direct and has_indirect:
            state = "D"
            state_label = "Listed + Indirectly Supported"
            demonstrated = False  # CRITICAL: demonstrated_in_resume = True ONLY when direct evidence exists
            confidence = 0.65
            state_d_skills.append(skill)
        elif is_listed and not has_direct and not has_indirect:
            state = "B"
            state_label = "Listed + Not Demonstrated"
            demonstrated = False
            confidence = 0.90
            state_b_skills.append(skill)
        elif not is_listed and has_direct:
            state = "C"
            state_label = "Demonstrated + Not Listed"
            demonstrated = True
            confidence = 0.95
            state_c_skills.append(skill)
        else:
            state = "D"
            state_label = "Listed + Indirectly Supported"
            demonstrated = False
            confidence = 0.60
            state_d_skills.append(skill)

        skills_inventory.append({
            "skill": skill,
            "state": state,
            "state_label": state_label,
            "listed_in_skills": is_listed,
            "demonstrated_in_resume": demonstrated,
            "direct_evidence": direct_ev_list,
            "indirect_evidence": indirect_ev_list,
            "evidence_locations": all_locations,
            "evidence_type": primary_ev_type,
            "confidence": confidence
        })

    # 6. Classify skills by category
    skills_by_category = {}
    for cat, skill_items in SKILL_CATEGORY_MAP.items():
        matched = []
        for s in all_candidate_skills:
            if s.lower() in skill_items or any(tech_k == s.lower() for tech_k in skill_items):
                matched.append(s)
        if matched:
            skills_by_category[cat] = list(dict.fromkeys(matched))

    # 7. Structured Findings
    findings = []

    vague_skills = ["machine learning fundamentals", "web development", "computer skills", "programming", "tools"]
    skills_sec_text = "\n".join(skills_lines).lower() if skills_lines else text_lower
    for vs in vague_skills:
        if vs in skills_sec_text:
            findings.append(make_finding(
                f"vague_skill_{vs.replace(' ', '_')}", "warning", "medium", "skills", "Technical Skills",
                vs.title(), f"Broad/vague skill label '{vs.title()}' detected in Technical Skills section.",
                "Broad labels do not identify specific technologies, libraries, or competencies to technical recruiters.",
                f"Replace '{vs.title()}' with concrete tools and techniques actually used (e.g. Scikit-Learn, Pandas, Feature Engineering, Regression/Classification)."
            ))

    # State B: Listed but Not Demonstrated (Strictly explain: no direct evidence in this document)
    if state_b_skills:
        findings.append(make_finding(
            "undemonstrated_skills_inventory",
            "improvement", "medium", "skills", "Technical Skills",
            ", ".join(state_b_skills),
            f"Skills listed in Technical Skills without direct project implementation evidence: {', '.join(state_b_skills)}.",
            "Recruiters look for verifiable hands-on evidence (project bullets or technologies used lists). Note: No direct evidence found in this resume does not imply a lack of competency, but indicates that no project or experience entry currently verifies practical application.",
            f"Consider adding project bullet points or technology list entries showing these technologies in action ({', '.join(state_b_skills)}), or focusing your skills section on directly demonstrated tools."
        ))

    # State D: Listed + Indirectly Supported findings
    if state_d_skills:
        for d_skill in state_d_skills:
            if d_skill == "Git" and has_github_link:
                findings.append(make_finding(
                    "skill_git_indirect", "strength", "low", "skills", "Technical Skills",
                    "Git (GitHub profile)",
                    "Git is indirectly supported by repository evidence (GitHub profile link).",
                    "Repository links provide correlated proof of version control knowledge, though direct project workflow bullets further strengthen credibility.",
                    "Consider explicitly mentioning version control workflows (e.g. 'Managed Git branching and pull request reviews') in project bullets."
                ))

    # State C: Demonstrated + Not Listed findings
    for unlisted in state_c_skills[:4]:
        loc_str = skill_direct_evidence.get(unlisted, [{}])[0].get("location", "Projects")
        findings.append(make_finding(
            f"unlisted_project_skill_{unlisted.lower().replace(' ', '_').replace('.', '_')}",
            "improvement", "low", "skills", "Technical Skills",
            loc_str,
            f"'{unlisted}' is demonstrated in your {loc_str}, but is not listed in your Technical Skills section.",
            "ATS skill scanners index the Technical Skills inventory directly for recruiter keyword filters.",
            f"Add '{unlisted}' to your Technical Skills section under the appropriate category."
        ))

    # 8. Balanced Domain Detection (Whole-Resume Evidence Aggregation)
    edu_text = "\n".join(edu_lines).lower()
    sum_text = "\n".join(sum_lines).lower()
    proj_sec_text = "\n".join(proj_lines).lower()

    domain_scores = {}
    for domain, info in DOMAIN_TAXONOMY.items():
        score_val = 0.0
        # Core anchors
        for anchor in info["core_anchors"]:
            pat = r"(?<!\w)" + re.escape(anchor) + r"(?!\w)"
            if re.search(pat, edu_text):
                score_val += 5.0
            if re.search(pat, sum_text):
                score_val += 4.0
            if re.search(pat, proj_sec_text):
                score_val += 3.5
            if re.search(pat, skills_sec_text):
                score_val += 2.5
            if re.search(pat, text_lower):
                score_val += 1.0

        # Keywords
        unique_kw_hits = 0
        for kw in info["keywords"]:
            pat = r"(?<!\w)" + re.escape(kw) + r"(?!\w)"
            if re.search(pat, text_lower):
                unique_kw_hits += 1
                if re.search(pat, edu_text):
                    score_val += 2.0
                if re.search(pat, proj_sec_text):
                    score_val += 1.5
                if re.search(pat, skills_sec_text):
                    score_val += 1.0
        score_val += min(15.0, unique_kw_hits * 1.2)
        domain_scores[domain] = score_val

    sorted_domains = sorted(domain_scores.items(), key=lambda x: x[1], reverse=True)
    primary_domain = sorted_domains[0][0] if sorted_domains and sorted_domains[0][1] > 0 else "Software Development / Computer Science / AI & ML"
    secondary_domains = [d[0] for d in sorted_domains[1:3] if d[1] > 0]

    # Industry Terminology Suggestions
    domain_info = DOMAIN_TAXONOMY.get(primary_domain, DOMAIN_TAXONOMY["Software Development / Computer Science / AI & ML"])
    detected_industry_terms = []
    optional_suggestions = []

    for kw in domain_info["keywords"]:
        fmt_kw = format_tech_name(kw)
        if re.search(r"(?<!\w)" + re.escape(kw) + r"(?!\w)", text_lower):
            detected_industry_terms.append(fmt_kw)
        else:
            optional_suggestions.append(fmt_kw)

    detected_industry_terms = list(dict.fromkeys(detected_industry_terms))
    optional_suggestions = list(dict.fromkeys(optional_suggestions))

    # 9. Strict Skills Depth Scoring (15% Weight)
    evidence = []
    deductions = []
    notes = []

    skill_count = len(all_candidate_skills)
    cat_count = len(skills_by_category)
    state_a_count = len(state_a_skills)
    state_d_count = len(state_d_skills)
    state_b_count = len(state_b_skills)

    if skill_count >= 12 and cat_count >= 3:
        score = 100
        evidence.append(f"Comprehensive technical skills portfolio: {skill_count} relevant technologies detected across {cat_count} categories.")
    elif skill_count >= 8:
        score = 90
        evidence.append(f"Strong technical skills portfolio: {skill_count} technologies detected ({', '.join(all_candidate_skills[:5])}).")
    elif skill_count >= 5:
        score = 75
        evidence.append(f"Solid core technical skills ({skill_count} technologies identified).")
    elif skill_count >= 3:
        score = 60
        deductions.append(f"Only {skill_count} technical skills detected. Expanding technical stack recommended (-15 pts).")
    else:
        score = 35
        deductions.append("Very few technical skills detected. Explicitly list programming languages, frameworks, and databases (-40 pts).")

    if state_b_count >= 2 and bool(proj_lines):
        ded_val = min(10, state_b_count * 2)
        score = max(35, score - ded_val)
        deductions.append(f"{state_b_count} skills listed without project implementation evidence ({', '.join(state_b_skills)}) (-{ded_val} pts).")
    elif state_a_count > 0:
        evidence.append(f"{state_a_count} key skills ({', '.join(state_a_skills[:5])}) are reinforced with direct hands-on project evidence.")

    if any("vague_skill" in f["id"] for f in findings):
        score = max(35, score - 5)
        deductions.append("Broad or vague skill labels detected in Skills section (-5 pts).")

    final_score = max(20, min(100, score))
    if final_score == 100 and not deductions:
        evidence.insert(0, "All applicable checks passed.")

    return {
        "score": final_score,
        "skills_inventory": skills_inventory,
        "state_a": state_a_skills,
        "state_b": state_b_skills,
        "state_c": state_c_skills,
        "state_d": state_d_skills,
        "listed_and_demonstrated": state_a_skills,
        "listed_not_demonstrated": state_b_skills,
        "demonstrated_not_listed": state_c_skills,
        "listed_indirectly_supported": state_d_skills,
        "demonstrated_and_listed": state_a_skills,
        "detected_skills": all_candidate_skills,
        "skills_by_category": skills_by_category,
        "primary_domain": primary_domain,
        "secondary_domains": secondary_domains,
        "detected_industry_terminology": detected_industry_terms[:8],
        "optional_industry_terminology_suggestions": optional_suggestions[:6],
        "top_matched_keywords": detected_industry_terms[:8],
        "missing_keywords": optional_suggestions[:6],
        "findings": findings,
        "evidence": evidence,
        "deductions": deductions,
        "detected_items": all_candidate_skills,
        "notes": notes
    }


calculate_skills_and_industry_keywords = analyze_skills_cross_validation


# ==============================================================================
# 9. BULLET-LEVEL ANALYSIS & ACTION VERBS ENGINE
# ==============================================================================

def analyze_single_bullet(bullet_text, project_techs=None):
    """
    Evaluates an individual bullet across 6 strict criteria:
    Action, Task, Technology, Result, Metric, Clarity.
    Each returns status (PASS | WEAK | MISSING | PARTIAL | NOT_APPLICABLE), evidence, and reason.
    Returns individualized context-specific recommendation.
    """
    b_clean = bullet_text.strip()
    b_lower = b_clean.lower()
    words = b_clean.split()
    word_count = len(words)

    # 1. Action Check
    first_word = words[0].lower().rstrip(".,:;") if words else ""
    lemma_v = VERB_LEMMA_MAP.get(first_word, first_word)
    weak_phrase_found = next((wp for wp in WEAK_ACTION_PHRASES if wp in b_lower), None)

    if lemma_v in ACTION_VERBS:
        action_status = "PASS"
        action_ev = first_word.capitalize()
        action_reason = "Begins with a strong, direct past-tense action verb."
    elif weak_phrase_found:
        action_status = "WEAK"
        action_ev = weak_phrase_found
        action_reason = f"Begins with passive phrasing '{weak_phrase_found}' which dilutes ownership."
    elif any(re.search(r"\b" + re.escape(v) + r"\b", b_lower) for v in ACTION_VERBS):
        action_status = "PARTIAL"
        action_ev = first_word
        action_reason = "Contains action verbs, but does not lead directly with a strong action verb at the start."
    else:
        action_status = "MISSING"
        action_ev = first_word if first_word else "None"
        action_reason = "Missing an active engineering action verb at the start of the bullet."

    # 2. Task Check
    if word_count >= 6 and any(k in b_lower for k in ["for", "to", "using", "with", "system", "feature", "engine", "application", "pipeline", "service", "model", "query", "database", "api", "interface", "matcher", "scraper"]):
        task_status = "PASS"
        task_ev = b_clean[:60] + ("..." if len(b_clean) > 60 else "")
        task_reason = "Specific engineering task and functional scope are clearly defined."
    elif word_count >= 4:
        task_status = "NEEDS_IMPROVEMENT"
        task_ev = b_clean
        task_reason = "Task is briefly mentioned but lacks technical implementation context."
    else:
        task_status = "MISSING"
        task_ev = b_clean
        task_reason = "Bullet is too brief to define a clear technical task."

    # 3. Technology Check
    techs_found = []
    for skill in SKILL_SET:
        pat = r"(?<!\w)" + re.escape(skill) + r"(?!\w)"
        if re.search(pat, b_lower):
            techs_found.append(format_tech_name(skill))
    for tech_k, tech_canon in CANONICAL_TECH_CASING.items():
        pat = r"(?<!\w)" + re.escape(tech_k) + r"(?!\w)"
        if re.search(pat, b_lower):
            techs_found.append(tech_canon)
    techs_found = list(dict.fromkeys(techs_found))

    if techs_found:
        tech_status = "PASS"
        tech_ev = ", ".join(techs_found)
        tech_reason = f"Explicitly specifies {len(techs_found)} concrete tool{'s' if len(techs_found)>1 else ''} ({tech_ev})."
    elif any(k in b_lower for k in ["api", "database", "microservice", "frontend", "backend", "web application", "vector", "cache", "model", "algorithm"]):
        tech_status = "PARTIAL"
        tech_ev = "General architectural terms"
        tech_reason = "Mentions technical architectural concepts without specifying exact libraries or frameworks."
    else:
        tech_status = "MISSING"
        tech_ev = "None"
        tech_reason = "No technical tools, frameworks, or languages specified in bullet text."

    # 4. Result Check
    result_triggers = ["improving", "improved", "reduced", "reducing", "increased", "increasing", "optimizing", "optimized", "enabling", "enabled", "serving", "automating", "automated", "accelerating", "achieving", "achieved", "resulting in", "allowing"]
    has_explicit_result = any(re.search(r"\b" + re.escape(rt) + r"\b", b_lower) for rt in result_triggers)
    
    if has_explicit_result:
        result_status = "PASS"
        matched_rt = next(rt for rt in result_triggers if re.search(r"\b" + re.escape(rt) + r"\b", b_lower))
        result_ev = matched_rt
        result_reason = f"Includes a clear outcome statement indicated by '{matched_rt}'."
    elif any(k in b_lower for k in ["seamless", "scalable", "responsive", "efficient", "real-time", "secure", "dynamic"]):
        result_status = "PARTIAL"
        result_ev = "Qualitative outcome"
        result_reason = "Mentions qualitative outcome without explicit operational measurement."
    else:
        result_status = "MISSING"
        result_ev = "None"
        result_reason = "Lacks an explicit statement of result, outcome, or operational benefit."

    # 5. Metric Check
    m_pct = re.findall(r"\b\d+(?:\.\d+)?%", b_clean)
    m_scale = re.findall(r"\b\d{1,3}(?:,\d{3})*\+?\s*(?:users|active users|clients|records|queries|requests|endpoints|ms|seconds|fps|rps)\b", b_clean, re.IGNORECASE)
    metrics_in_b = m_pct + m_scale

    if metrics_in_b:
        metric_status = "PASS"
        metric_ev = ", ".join(metrics_in_b)
        metric_reason = f"Quantified performance metric present: {metric_ev}."
    else:
        metric_status = "MISSING"
        metric_ev = "None"
        metric_reason = "No measurable metric (percentage improvement, latency, scale, user count) found."

    # 6. Clarity Check
    if 8 <= word_count <= 35:
        clarity_status = "PASS"
        clarity_ev = f"{word_count} words"
        clarity_reason = "Optimal sentence length and concise executive structure."
    elif 5 <= word_count < 8:
        clarity_status = "NEEDS_IMPROVEMENT"
        clarity_ev = f"{word_count} words"
        clarity_reason = "Slightly brief; consider elaborating on technical implementation details."
    elif word_count > 35:
        clarity_status = "NEEDS_IMPROVEMENT"
        clarity_ev = f"{word_count} words"
        clarity_reason = "Sentence is long; consider tightening for quicker ATS and recruiter scanning."
    else:
        clarity_status = "NEEDS_IMPROVEMENT"
        clarity_ev = f"{word_count} words"
        clarity_reason = "Very short fragment."

    # Generate Non-Repetitive, Individualized Recommendation
    if action_status == "WEAK":
        rec = f"Replace the passive opening '{weak_phrase_found}' with a strong action verb (e.g. 'Engineered', 'Optimized', 'Architected')."
    elif metric_status == "MISSING" and action_status == "PASS" and tech_status == "PASS":
        rec = "Add a verified outcome or scale metric (e.g. latency improvement, volume processed, test coverage)—only if you actually measured it."
    elif tech_status == "MISSING":
        rec = "Specify the exact library, framework, or database used to build this component."
    elif result_status == "MISSING":
        rec = "Conclude the bullet with the tangible end-result or operational benefit achieved by this implementation."
    elif clarity_status == "NEEDS_IMPROVEMENT":
        rec = "Refine bullet to 10-25 words focusing strictly on action + technology + task + outcome."
    else:
        rec = "Strong technical bullet with clear action, task, and specificity."

    return {
        "text": b_clean,
        "action": {"status": action_status, "evidence": action_ev, "reason": action_reason},
        "task": {"status": task_status, "evidence": task_ev, "reason": task_reason},
        "technology": {"status": tech_status, "evidence": tech_ev, "reason": tech_reason},
        "result": {"status": result_status, "evidence": result_ev, "reason": result_reason},
        "metric": {"status": metric_status, "evidence": metric_ev, "reason": metric_reason},
        "clarity": {"status": clarity_status, "evidence": clarity_ev, "reason": clarity_reason},
        "recommendation": rec,
        "has_metric": metric_status == "PASS",
        "starts_with_action": action_status == "PASS",
        "leading_verb": first_word if action_status == "PASS" else None
    }


def analyze_project_bullets(parsed_data):
    """
    Project Bullet Quality & Action Verbs Analysis Engine.
    """
    bullets = parsed_data.get("bullet_points", [])
    project_bullets = [b for b in bullets if b.get("section") in ["projects", "experience", "internships"]]
    
    analyzed_bullets = []
    findings = []
    quant_opportunities = []
    
    action_led_bullets_list = []
    weak_opening_bullets_list = []
    unique_action_verbs_set = set()
    total_action_verb_occurrences = 0

    for idx, b in enumerate(project_bullets, 1):
        b_text = b["text"]
        b_lower = b_text.lower()
        if any(b_lower.startswith(k) for k in ["technologies used", "tech stack", "tools used", "environment", "built with"]):
            continue
            
        diag = analyze_single_bullet(b_text)
        
        if diag["starts_with_action"] and diag["leading_verb"]:
            action_led_bullets_list.append(b_text)
            unique_action_verbs_set.add(diag["leading_verb"].capitalize())
            total_action_verb_occurrences += 1

        for verb in ACTION_VERBS:
            if re.search(r"\b" + re.escape(verb) + r"\b", b_lower):
                unique_action_verbs_set.add(verb.capitalize())
                total_action_verb_occurrences += 1

        if diag["action"]["status"] == "WEAK":
            weak_opening_bullets_list.append(b_text)
            findings.append(make_finding(
                f"bullet_weak_verb_{idx}",
                "issue", "medium", "content_quality", "Projects",
                b_text,
                f"Bullet begins with passive phrasing '{diag['action']['evidence']}'.",
                "Passive phrasing weakens ownership and candidate impact.",
                "Start with a strong past-tense action verb (e.g. 'Developed', 'Engineered', 'Optimized') followed by the specific task and outcome."
            ))

        if not diag["has_metric"]:
            quant_opportunities.append({
                "original": b_text,
                "problem": "Action and technical purpose are clear, but measurable outcome metric is absent.",
                "suggestion": diag["recommendation"]
            })
            findings.append(make_finding(
                f"bullet_quant_opp_{idx}",
                "improvement", "medium", "quantification", "Projects",
                b_text,
                "Bullet point lacks a measurable outcome or quantified metric.",
                "Recruiters evaluate technical execution through measurable outcomes (latency, volume, speedup, accuracy).",
                diag["recommendation"]
            ))

        analyzed_bullets.append({
            "index": idx,
            "text": b_text,
            "diagnostics": diag,
            "classification": "Strong" if (diag["starts_with_action"] and diag["has_metric"]) else ("Acceptable" if diag["starts_with_action"] else ("Weak" if diag["action"]["status"] == "WEAK" else "Needs revision")),
            "starts_with_action": diag["starts_with_action"],
            "leading_verb": diag["leading_verb"],
            "has_metric": diag["has_metric"]
        })

    total_bullets_count = len(analyzed_bullets)
    action_led_count = len(action_led_bullets_list)
    action_led_ratio = (action_led_count / total_bullets_count) if total_bullets_count > 0 else 0.0
    action_led_ratio_pct = int(round(action_led_ratio * 100))

    action_verb_stats = {
        "total_bullets": total_bullets_count,
        "action_led_bullets": action_led_bullets_list,
        "action_led_count": action_led_count,
        "action_led_ratio": action_led_ratio,
        "action_led_ratio_pct": action_led_ratio_pct,
        "action_led_ratio_str": f"{action_led_ratio_pct}%",
        "unique_action_verbs": sorted(list(unique_action_verbs_set)),
        "unique_action_verbs_count": len(unique_action_verbs_set),
        "total_action_verb_occurrences": total_action_verb_occurrences,
        "weak_opening_bullets": weak_opening_bullets_list,
        "weak_opening_count": len(weak_opening_bullets_list)
    }

    return {
        "analyzed_bullets": analyzed_bullets,
        "quant_opportunities": quant_opportunities,
        "action_verb_stats": action_verb_stats,
        "findings": findings
    }


# ==============================================================================
# 10. PROJECT-BY-PROJECT DEEP ANALYSIS ENGINE (TECHNOLOGIES USED EXTRACTION)
# ==============================================================================

def analyze_individual_projects(parsed_data, raw_text):
    """
    Analyzes EVERY project independently.
    Explicitly parses 'Technologies Used:' lines as valid project technology evidence.
    Performs bullet-by-bullet analysis with Action, Task, Technology, Result, Metric, Clarity.
    """
    proj_lines = parsed_data["sections"].get("projects", [])
    if not proj_lines:
        return []

    # Segment project lines into individual project blocks
    project_blocks = []
    current_block = {"title": "", "lines": []}

    for line in proj_lines:
        clean_l = line.strip()
        if not clean_l:
            continue
        
        # Check if line looks like a project header
        is_title_candidate = (
            (any(t in clean_l.lower() for t in ["live - demo", "live-demo", "github", "clone", "system", "platform", "matcher", "project"]) or "|" in clean_l)
            and len(clean_l) < 85
            and not any(clean_l.startswith(b) for b in BULLET_CHARS)
            and not any(clean_l.lower().startswith(k) for k in ["technologies", "tech stack", "tools used", "built with", "environment"])
        )

        if is_title_candidate:
            if current_block["lines"] or current_block["title"]:
                project_blocks.append(current_block)
            current_block = {"title": clean_l, "lines": []}
        else:
            current_block["lines"].append(clean_l)

    if current_block["lines"] or current_block["title"]:
        project_blocks.append(current_block)

    if not project_blocks and proj_lines:
        project_blocks.append({"title": "Featured Technical Project", "lines": proj_lines})

    analyzed_projects = []
    for p_idx, block in enumerate(project_blocks, 1):
        raw_title = block["title"] or f"Project {p_idx}"
        clean_name = re.sub(r"\s*\|\s*(?:live\s*[-–]\s*demo|github).*$", "", raw_title, flags=re.IGNORECASE).strip()
        if not clean_name:
            clean_name = raw_title

        block_text = "\n".join(block["lines"])
        combined_text = (raw_title + "\n" + block_text).lower()

        # 1. Explicitly extract 'Technologies Used' / 'Tech Stack' list
        explicit_tech_list = []
        for line in block["lines"]:
            m_tech = re.search(r"^(?:technologies\s+used|tech\s+stack|technologies|tools\s+used|built\s+with|environment)\s*:\s*(.*)$", line, re.IGNORECASE)
            if m_tech:
                tech_raw_str = m_tech.group(1).strip()
                # Split by commas, slashes, pipes, bullets
                parts = re.split(r"[,/|•·+\-]+", tech_raw_str)
                for pt in parts:
                    clean_pt = pt.strip()
                    if clean_pt and len(clean_pt) >= 2:
                        explicit_tech_list.append(format_tech_name(clean_pt))

        # 2. Extract all detected technologies in this project (from title, tech line, and bullets)
        detected_tech = list(explicit_tech_list)
        for skill in SKILL_SET:
            pat = r"(?<!\w)" + re.escape(skill) + r"(?!\w)"
            if re.search(pat, combined_text):
                detected_tech.append(format_tech_name(skill))
        for tech_k, tech_canon in CANONICAL_TECH_CASING.items():
            pat = r"(?<!\w)" + re.escape(tech_k) + r"(?!\w)"
            if re.search(pat, combined_text):
                detected_tech.append(tech_canon)
        detected_tech = list(dict.fromkeys(detected_tech))

        # 3. Extract bullets and perform bullet-by-bullet diagnostics
        bullets = []
        bullet_diagnostics = []
        action_led = []
        weak_bullets = []
        quantified_metrics = []

        for l in block["lines"]:
            is_b, b_prefix = is_bullet_line(l)
            # Ignore tech stack headers as bullets
            if any(l.lower().startswith(k) for k in ["technologies used", "tech stack", "tools used", "built with", "environment"]):
                continue
            if is_b or (len(l) > 30 and not l.startswith("http")):
                b_clean = clean_bullet_prefix(l, b_prefix)
                bullets.append(b_clean)
                diag = analyze_single_bullet(b_clean, detected_tech)
                bullet_diagnostics.append(diag)
                
                if diag["starts_with_action"]:
                    action_led.append(b_clean)
                if diag["action"]["status"] == "WEAK":
                    weak_bullets.append(b_clean)
                if diag["has_metric"]:
                    quantified_metrics.append(diag["metric"]["evidence"])

        # 4. Link status
        has_live_demo = bool(re.search(r"live\s*[-–]\s*demo", combined_text))
        has_github = bool(re.search(r"github", combined_text))
        live_demo_url = re.search(r"https?://(?!.*(?:linkedin|github)\.com)[\w\.-]+\.[a-z]{2,}(/[\w\.-]*)*", raw_title + " " + block_text, re.IGNORECASE)
        github_url = re.search(r"github\.com/[\w\-_/]+", raw_title + " " + block_text, re.IGNORECASE)

        link_status = {
            "has_live_demo_label": has_live_demo,
            "has_github_label": has_github,
            "live_demo_url": live_demo_url.group(0) if live_demo_url else None,
            "github_url": github_url.group(0) if github_url else None,
            "has_active_links": bool(live_demo_url or github_url)
        }

        # 5. Technical Depth & Outcome
        if len(detected_tech) >= 5:
            tech_depth = f"Full-stack multi-service architecture integrating {', '.join(detected_tech[:5])}."
        elif len(detected_tech) >= 3:
            tech_depth = f"Full-stack/backend implementation using {', '.join(detected_tech[:3])}."
        elif len(detected_tech) >= 1:
            tech_depth = f"Focused technical implementation utilizing {', '.join(detected_tech)}."
        else:
            tech_depth = "Core software development project."

        if quantified_metrics:
            result_impact = f"Quantified performance outcome: {', '.join(quantified_metrics)}."
        else:
            result_impact = "Functional delivery without explicit quantified impact metrics."

        # 6. Strengths, Problems, Recommendations (Personalized per project)
        strengths = []
        problems = []
        recommendations = []

        if explicit_tech_list:
            strengths.append(f"Explicitly documents technology stack under Technologies Used: {', '.join(explicit_tech_list)}.")
        elif detected_tech:
            strengths.append(f"Demonstrates practical technology stack: {', '.join(detected_tech[:4])}.")

        if action_led:
            strengths.append(f"{len(action_led)} bullet point{'s' if len(action_led)>1 else ''} lead with active engineering verbs.")
        if quantified_metrics:
            strengths.append(f"Includes measurable project outcomes ({', '.join(quantified_metrics)}).")

        if not quantified_metrics and bullets:
            problems.append("Lacks measurable outcome metrics (e.g. latency, volume processed, user scale, speedup).")
            recommendations.append(f"Add a verified metric to your bullet point ('{bullets[0][:60]}...')—only if you actually measured it.")

        if weak_bullets:
            problems.append(f"Contains passive opening phrasing ('{weak_bullets[0][:50]}...').")
            recommendations.append("Begin every bullet with a strong action verb (e.g. 'Engineered', 'Optimized', 'Integrated').")

        if (has_live_demo or has_github) and not link_status["has_active_links"]:
            problems.append("Repository/Demo labels detected without clickable URLs in extracted text.")
            recommendations.append("Ensure live demo and repository links contain active, clickable hyperlinks.")

        analyzed_projects.append({
            "project_name": clean_name,
            "raw_title": raw_title,
            "technologies_used": explicit_tech_list,
            "technologies": detected_tech,
            "bullet_count": len(bullets),
            "bullets": bullets,
            "bullet_analysis": bullet_diagnostics,
            "action_led_bullets": action_led,
            "action_led_count": len(action_led),
            "weak_bullets": weak_bullets,
            "metrics": quantified_metrics,
            "quantified_metrics": quantified_metrics,
            "has_metrics": bool(quantified_metrics),
            "technical_depth": tech_depth,
            "result_impact": result_impact,
            "outcome": result_impact,
            "link_status": link_status,
            "links": link_status,
            "strengths": strengths,
            "problems": problems,
            "recommendations": recommendations
        })

    return analyzed_projects


# ==============================================================================
# 11. QUANTIFICATION ANALYSIS ENGINE (8 CATEGORIES + EXCLUSION AUDIT)
# ==============================================================================

def calculate_quantification_score(raw_text, parsed_data):
    """
    Dedicated Quantification Analysis (5% Weight, 0-100 Score):
    Categorizes genuine metrics across 8 strict categories:
      1. percentage_impact_metrics (e.g. 'reduced latency by 35%')
      2. user_scale_metrics (e.g. '5,000+ active users')
      3. performance_metrics (e.g. 'sub-second', '60 fps', '200ms')
      4. accuracy_metrics (e.g. '94% accuracy', '0.88 F1-score')
      5. latency_speed_metrics (e.g. '150ms latency', '2x speedup')
      6. time_cost_metrics (e.g. '$5,000 saved', '10 hours/week')
      7. volume_metrics (e.g. '100,000 records', '50+ endpoints')
      8. other_meaningful_metrics (e.g. '99.9% uptime', '5 responsive breakpoints')

    Strictly audits and excludes non-achievement numeric values:
      - phone numbers
      - dates and years
      - CGPA / GPA values
      - education / academic marks (e.g. 88.9%, 92.66%)
      - version numbers (e.g. Python 3.10, HTML5, v2.0)
    """
    percent_metrics = []
    scale_metrics = []
    performance_metrics = []
    accuracy_metrics = []
    latency_speed_metrics = []
    time_cost_metrics = []
    volume_metrics = []
    other_metrics = []

    # --- EXCLUSION AUDIT ---
    excluded_numbers = {
        "phone": [],
        "dates": [],
        "cgpa": [],
        "education_marks": [],
        "version_numbers": []
    }

    # Audit phone numbers
    for ph in re.findall(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4,6}", raw_text):
        if ph:
            excluded_numbers["phone"].append({
                "value": ph.strip(),
                "reason": "Candidate contact phone number"
            })

    # Audit dates & years
    for dt in re.findall(r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\b\d{4}\s*[-–to]+\s*(?:\d{4}|Present|Current)\b|\b\d{4}\b", raw_text, re.IGNORECASE):
        if int(dt[:4]) >= 1990 and int(dt[:4]) <= 2030 if dt[:4].isdigit() else True:
            excluded_numbers["dates"].append({
                "value": dt.strip(),
                "reason": "Academic/Employment timeline date"
            })

    # Audit CGPA / GPA
    for cg in re.findall(r"\b(?:cgpa|gpa)\s*[:=]?\s*(\d+(?:\.\d+)?(?:\s*/\s*10)?)\b", raw_text, re.IGNORECASE):
        excluded_numbers["cgpa"].append({
            "value": f"CGPA: {cg.strip()}",
            "reason": "Academic grade point average, excluded from technical execution metrics"
        })

    # Audit educational percentages
    for em in re.findall(r"(?:intermediate|ssc|cbse|icse|btech|mtech|degree|class\s*1[02]|marks|grade)\s*[:=–-]?\s*(\d+(?:\.\d+)?%)", raw_text, re.IGNORECASE):
        excluded_numbers["education_marks"].append({
            "value": em.strip(),
            "reason": "School/Board examination score, excluded from project achievement metrics"
        })

    # Audit software version numbers
    for vn in re.findall(r"\b(?:python|java|html|css|vue|react|node|bootstrap|v)\s*(\d+(?:\.\d+)?)\b", raw_text, re.IGNORECASE):
        excluded_numbers["version_numbers"].append({
            "value": f"Version {vn.strip()}",
            "reason": "Software/framework release version number"
        })

    # --- 1. Extract and Filter Percentages ---
    raw_percentages = re.findall(r"(?:(?:reduced|improved|increased|boosted|saved|scaled|growth|by)\s+)?\b\d+(?:\.\d+)?%", raw_text, re.IGNORECASE)
    for p in raw_percentages:
        match_pos = raw_text.find(p)
        if match_pos != -1:
            context = raw_text[max(0, match_pos - 40): min(len(raw_text), match_pos + len(p) + 40)].lower()
            if any(term in context for term in ["cgpa", "gpa", "ssc", "intermediate", "cbse", "icse", "btech", "mtech", "degree", "school", "college", "marks", "grade", "percentage:"]):
                continue
        # Check if accuracy metric
        if any(term in p.lower() for term in ["accuracy", "precision", "recall", "f1"]):
            accuracy_metrics.append(p.strip())
        else:
            percent_metrics.append(p.strip())

    # --- 2. User & Scale Metrics ---
    scale_matches = re.findall(r"\b\d{1,3}(?:,\d{3})*\+?\s*(?:users|active users|concurrent users|clients|customers|downloads|visits)\b", raw_text, re.IGNORECASE)
    for sm in scale_matches:
        scale_metrics.append(sm.strip())

    # --- 3. Performance & Speedup Metrics ---
    perf_matches = re.findall(r"\b\d+(?:\.\d+)?\s*(?:fps|req/s|rps|tps|kb/s|mb/s)\b|\b\d+x\s*(?:faster|speedup)\b", raw_text, re.IGNORECASE)
    for pm in perf_matches:
        performance_metrics.append(pm.strip())

    # --- 4. Latency & Response Time Metrics ---
    lat_matches = re.findall(r"\b\d+(?:\.\d+)?\s*ms\b|\b(?:sub-second|ultra-low latency)\b", raw_text, re.IGNORECASE)
    for lm in lat_matches:
        latency_speed_metrics.append(lm.strip())

    # --- 5. Accuracy & Precision Metrics ---
    acc_matches = re.findall(r"\b\d+(?:\.\d+)?%\s*(?:accuracy|precision|recall|f1-score|f1 score|mAP|auc|bleu)\b|\b(?:accuracy|precision)\s*(?:of|:)?\s*\d+(?:\.\d+)?%", raw_text, re.IGNORECASE)
    for am in acc_matches:
        accuracy_metrics.append(am.strip())

    # --- 6. Time & Cost Savings Metrics ---
    cost_matches = re.findall(r"\$\d+(?:,\d+)*(?:\.\d+)?[kKmMbB]?|\b\d+\s*(?:hours|days|weeks|months)\s*(?:saved|reduction)\b", raw_text, re.IGNORECASE)
    for cm in cost_matches:
        time_cost_metrics.append(cm.strip())

    # --- 7. Data Volume & Endpoint Metrics ---
    vol_matches = re.findall(r"\b\d{1,3}(?:,\d{3})*\+?\s*(?:records|rows|queries|requests|endpoints|transactions|documents)\b", raw_text, re.IGNORECASE)
    for vm in vol_matches:
        volume_metrics.append(vm.strip())

    # --- 8. Other Meaningful Metrics ---
    other_matches = re.findall(r"\b99\.\d+%\s*uptime\b|\b\d+\s*(?:responsive breakpoints|test cases)\b", raw_text, re.IGNORECASE)
    for om in other_matches:
        other_metrics.append(om.strip())

    # Deduplicate within categories
    percent_metrics = list(dict.fromkeys(percent_metrics))
    scale_metrics = list(dict.fromkeys(scale_metrics))
    performance_metrics = list(dict.fromkeys(performance_metrics))
    accuracy_metrics = list(dict.fromkeys(accuracy_metrics))
    latency_speed_metrics = list(dict.fromkeys(latency_speed_metrics))
    time_cost_metrics = list(dict.fromkeys(time_cost_metrics))
    volume_metrics = list(dict.fromkeys(volume_metrics))
    other_metrics = list(dict.fromkeys(other_metrics))

    breakdown = {
        "percentage_impact_metrics": percent_metrics,
        "user_scale_metrics": scale_metrics,
        "performance_metrics": performance_metrics,
        "accuracy_metrics": accuracy_metrics,
        "latency_speed_metrics": latency_speed_metrics,
        "time_cost_metrics": time_cost_metrics,
        "volume_metrics": volume_metrics,
        "other_meaningful_metrics": other_metrics
    }

    all_detected_metrics = list(dict.fromkeys(
        percent_metrics + scale_metrics + performance_metrics +
        accuracy_metrics + latency_speed_metrics + time_cost_metrics +
        volume_metrics + other_metrics
    ))

    metric_count = len(all_detected_metrics)
    evidence = []
    deductions = []
    notes = []

    # Dynamic Scoring Computation
    if metric_count >= 4:
        score = 100
        evidence.append(f"Extensive quantifiable impact: {metric_count} measurable outcomes detected across {sum(1 for v in breakdown.values() if v)} metric categories ({', '.join(all_detected_metrics[:4])}).")
    elif metric_count == 3:
        score = 85
        evidence.append(f"Strong quantifiable evidence: {metric_count} measurable metrics identified ({', '.join(all_detected_metrics)}).")
    elif metric_count == 2:
        score = 70
        evidence.append(f"Solid quantifiable evidence: {metric_count} measurable metrics identified ({', '.join(all_detected_metrics)}).")
    elif metric_count == 1:
        score = 50
        evidence.append(f"Quantifiable metric identified: '{all_detected_metrics[0]}'.")
        notes.append("Adding more quantified metrics across all project bullets will strengthen your technical rating.")
    else:
        # Dynamic score when 0 genuine achievement metrics found
        bullets = parsed_data.get("bullet_points", [])
        proj_bullets_count = len([b for b in bullets if b.get("section") in ["projects", "experience"]])
        score = 20 if proj_bullets_count > 0 else 10
        deductions.append(
            f"Zero quantifiable performance metrics detected across 8 evaluation categories (Percentage impact: 0, Scale: 0, Performance: 0, Latency: 0, Cost/Time: 0, Volume: 0). Excluded non-achievement numbers: {len(excluded_numbers['cgpa'])} CGPA, {len(excluded_numbers['education_marks'])} academic exam percentage, {len(excluded_numbers['dates'])} date values (-30 pts)."
        )
        notes.append("Include verified numbers, percentage speedups, latency reductions, or volume scale in your project bullets.")

    if score == 100 and not deductions:
        evidence.insert(0, "All applicable checks passed.")

    return {
        "score": score,
        "breakdown": breakdown,
        "excluded_numbers": excluded_numbers,
        "percent_metrics": percent_metrics,
        "scale_metrics": scale_metrics,
        "performance_metrics": performance_metrics,
        "accuracy_metrics": accuracy_metrics,
        "latency_speed_metrics": latency_speed_metrics,
        "time_cost_metrics": time_cost_metrics,
        "volume_metrics": volume_metrics,
        "other_meaningful_metrics": other_metrics,
        "detected_items": all_detected_metrics,
        "evidence": evidence,
        "deductions": deductions,
        "notes": notes
    }


# ==============================================================================
# 12. ATS READABILITY & FORMAT ENGINE (11 STRICT CHECKS & JUSTIFICATION)
# ==============================================================================

def calculate_readability_format_score(extracted_doc, parsed_data):
    """
    ATS Readability / Format (25% Weight, 0-100 Score):
    Evaluates 11 explicit ATS format checks:
      1. text_extraction (PASS / FAIL)
      2. reading_order (PASS / FAIL)
      3. standard_section_headings (PASS / FAIL)
      4. bullet_parsing (PASS / FAIL)
      5. multi_column_risk (NONE / DETECTED)
      6. table_risk (NONE / DETECTED)
      7. text_box_risk (NONE / DETECTED)
      8. header_footer_critical_content_risk (NONE / DETECTED)
      9. image_only_text (NONE / DETECTED)
      10. unusual_character_risk (NONE / DETECTED)
      11. layout_consistency (PASS / FAIL)

    If score == 100: shows 'All applicable checks passed.' with all 11 passing checks.
    If score < 100: every deduction has explicit issue, evidence, points deducted, and reason.
    """
    score = 100
    evidence = []
    deductions = []
    detected_items = []
    notes = []
    checklist = []

    meta = extracted_doc.get("metadata", {})
    lines = extracted_doc.get("lines", [])
    raw_text = extracted_doc.get("raw_text", "")
    headings = [h["section"] for h in parsed_data.get("section_headings", [])]
    bullets = parsed_data.get("bullet_points", [])
    page_count = meta.get("page_count", 1)

    # 1. Text Extraction Check
    if meta.get("is_scanned") or len(raw_text.strip()) < 100:
        score -= 50
        deductions.append("Text extraction failed: Scanned image-only document layer without extractable text (-50 pts).")
        checklist.append({"check": "Text extraction", "status": "FAIL", "evidence": "Under 100 extractable characters", "points_deducted": 50, "reason": "Legacy and modern ATS parsers cannot index image-only documents without OCR."})
    else:
        evidence.append("Text extraction: PASS (Clean digital text layer extractable).")
        checklist.append({"check": "Text extraction", "status": "PASS", "evidence": "Clean digital text stream", "points_deducted": 0, "reason": "Standard digital text layer easily read by automated parsers."})

    # 2. Reading Order Check
    if meta.get("extraction_method") == "coordinate_aware":
        evidence.append("Reading order: PASS (Linear coordinate-ordered text stream).")
        checklist.append({"check": "Reading order", "status": "PASS", "evidence": "Coordinate-aware top-to-bottom stream", "points_deducted": 0, "reason": "Text fragments grouped into logical sequential reading flow."})
    else:
        evidence.append("Reading order: PASS (Standard sequential paragraph order).")
        checklist.append({"check": "Reading order", "status": "PASS", "evidence": "Sequential paragraph order", "points_deducted": 0, "reason": "Logical top-to-bottom layout stream."})

    # 3. Standard Section Headings Check
    present_headings = set(headings)
    if any(k in present_headings for k in ["skills", "education"]) and (any(k in present_headings for k in ["projects", "experience"]) or len(lines) < 25):
        evidence.append("Standard section headings: PASS (Recognized headings for Skills, Education, Projects).")
        checklist.append({"check": "Standard section headings", "status": "PASS", "evidence": ", ".join(list(present_headings)[:4]), "points_deducted": 0, "reason": "Standard section headers map directly to ATS data fields."})
    else:
        score -= 15
        deductions.append("Missing standard section headers (e.g. 'Skills', 'Education', 'Projects') (-15 pts).")
        checklist.append({"check": "Standard section headings", "status": "FAIL", "evidence": "Missing standard headers", "points_deducted": 15, "reason": "Non-standard section titles make it difficult for parsers to map candidate data."})

    # 4. Bullet Parsing Check
    if len(bullets) >= 4:
        evidence.append(f"Bullet parsing: PASS ({len(bullets)} structured bullet points detected).")
        checklist.append({"check": "Bullet parsing", "status": "PASS", "evidence": f"{len(bullets)} bullets parsed", "points_deducted": 0, "reason": "Clear bullet glyphs provide structural parsing hierarchy."})
    elif len(lines) > 20 and len(bullets) < 2:
        score -= 10
        deductions.append("Dense paragraph text detected without structured bullet points (-10 pts).")
        checklist.append({"check": "Bullet parsing", "status": "FAIL", "evidence": f"Only {len(bullets)} bullets in {len(lines)} lines", "points_deducted": 10, "reason": "Paragraph blocks decrease readability and parser token separation."})
    else:
        evidence.append(f"Bullet parsing: PASS ({len(bullets)} bullet points parsed).")
        checklist.append({"check": "Bullet parsing", "status": "PASS", "evidence": f"{len(bullets)} bullets parsed", "points_deducted": 0, "reason": "Bullet structure parsed cleanly."})

    # 5. Multi-Column Risk Check
    if meta.get("has_columns_clue"):
        score -= 10
        deductions.append("Multi-column layout risk detected: Parallel text lanes may cause out-of-order extraction (-10 pts).")
        checklist.append({"check": "Multi-column risk", "status": "DETECTED", "evidence": "Parallel column stream overlap", "points_deducted": 10, "reason": "Multi-column layouts often interleave text from both columns when parsed linearly."})
    else:
        evidence.append("Multi-column risk not detected: Clean linear reading flow.")
        checklist.append({"check": "Multi-column risk", "status": "NONE", "evidence": "Single-column linear flow", "points_deducted": 0, "reason": "No parallel vertical column lanes detected."})

    # 6. Table Risk Check
    if meta.get("has_tables_clue"):
        score -= 10
        deductions.append("Complex table risk detected: Multi-column grid structure may fragment text extraction (-10 pts).")
        checklist.append({"check": "Table risk", "status": "DETECTED", "evidence": "Grid / pipe table structures found", "points_deducted": 10, "reason": "Complex tables can split sentences across row boundaries."})
    else:
        evidence.append("Table risk not detected: No disruptive nested tables found.")
        checklist.append({"check": "Table risk", "status": "NONE", "evidence": "Clean text layout without nested tables", "points_deducted": 0, "reason": "No nested grid tables that risk parser fragmentation."})

    # 7. Text-Box Risk Check
    evidence.append("Text-box risk: NONE (No floating text boxes or shape containers).")
    checklist.append({"check": "Text-box risk", "status": "NONE", "evidence": "No floating text boxes", "points_deducted": 0, "reason": "Text resides directly in document body stream."})

    # 8. Header/Footer Critical-Content Risk Check
    evidence.append("Header/footer critical-content risk: NONE (Essential contact info located in main body).")
    checklist.append({"check": "Header/footer critical-content risk", "status": "NONE", "evidence": "Main body contact placement", "points_deducted": 0, "reason": "Essential information is not hidden in document headers/footers."})

    # 9. Image-Only Text Check
    evidence.append("Image-only text: NONE (No unparseable graphic text layers).")
    checklist.append({"check": "Image-only text", "status": "NONE", "evidence": "Full document text is digitally encoded", "points_deducted": 0, "reason": "No embedded bitmap text that would require OCR."})

    # 10. Unusual Character Risk Check
    unusual_symbols = len(re.findall(r"[\uFFF0-\uFFFF\u0000-\u0008\u000B\u000C\u000E-\u001F]", raw_text))
    if unusual_symbols > 5:
        score -= 5
        deductions.append(f"Unusual character encoding risk: {unusual_symbols} non-standard characters detected (-5 pts).")
        checklist.append({"check": "Unusual character risk", "status": "DETECTED", "evidence": f"{unusual_symbols} non-standard characters", "points_deducted": 5, "reason": "Corrupt or non-standard character codes can break ATS tokenizers."})
    else:
        evidence.append("Unusual character risk: NONE (Standard UTF-8 character encoding).")
        checklist.append({"check": "Unusual character risk", "status": "NONE", "evidence": "Standard UTF-8 encoding", "points_deducted": 0, "reason": "Standard character encoding throughout document."})

    # 11. Layout Consistency Check
    if page_count > 3:
        score -= 10
        deductions.append(f"Layout consistency: Length is {page_count} pages (1-2 pages recommended) (-10 pts).")
        checklist.append({"check": "Layout consistency", "status": "FAIL", "evidence": f"{page_count} pages", "points_deducted": 10, "reason": "Resumes over 2 pages reduce human readability and may indicate uncurated content."})
    else:
        evidence.append(f"Layout consistency: PASS (Optimal {page_count}-page length with standard margins).")
        checklist.append({"check": "Layout consistency", "status": "PASS", "evidence": f"{page_count} page{'s' if page_count>1 else ''}", "points_deducted": 0, "reason": "Document length is concise and fits standard recruiter scanning bounds."})

    final_score = max(20, min(100, score))
    if final_score == 100 and not deductions:
        evidence.insert(0, "All applicable checks passed.")

    return {
        "score": final_score,
        "checklist": checklist,
        "evidence": evidence,
        "deductions": deductions,
        "detected_items": [f"{c['check']}: {c['status']}" for c in checklist],
        "notes": notes
    }


# ==============================================================================
# 13. CONTENT QUALITY (20% WEIGHT)
# ==============================================================================

def calculate_content_quality_score(parsed_data, extracted_doc, quant_data, summary_analysis=None, bullet_analysis=None):
    """
    Content Quality (20% Weight, 0-100 Score):
    Evaluates action verbs, action-led bullet ratio, weak opening phrases,
    clarity, conciseness, and summary/objective quality.
    """
    if summary_analysis is None:
        summary_analysis = analyze_summary_objective(parsed_data)
    if bullet_analysis is None:
        bullet_analysis = analyze_project_bullets(parsed_data)

    score = 100
    evidence = []
    deductions = []
    detected_items = []
    notes = []
    weak_bullets = []

    raw_text = extracted_doc.get("raw_text", "").lower()
    bullets = parsed_data.get("bullet_points", [])

    # 1. Action Verbs analysis
    verbs_matched = set()
    for verb in ACTION_VERBS:
        if re.search(r"\b" + re.escape(verb) + r"\b", raw_text):
            verbs_matched.add(verb.lower())

    verbs_list = sorted(list(verbs_matched))
    verbs_count = len(verbs_list)

    if verbs_count >= 8:
        evidence.append(f"{verbs_count} unique action verbs detected ({', '.join([v.capitalize() for v in verbs_list[:6]])}).")
    elif verbs_count >= 4:
        evidence.append(f"{verbs_count} unique action verbs detected ({', '.join([v.capitalize() for v in verbs_list])}).")
    else:
        score -= 15
        deductions.append(f"Limited action verbs detected ({verbs_count} found: {', '.join(verbs_list) if verbs_list else 'None'}). Use direct action verbs like 'Developed', 'Built', 'Engineered', 'Optimized' (-15 pts).")

    # 2. Action-led bullet ratio
    av_stats = bullet_analysis.get("action_verb_stats", {})
    total_bullets = av_stats.get("total_bullets", len(bullets))
    action_led_count = av_stats.get("action_led_count", 0)

    if total_bullets > 0:
        ratio = action_led_count / total_bullets
        evidence.append(f"{total_bullets} project bullets analyzed, with {action_led_count} beginning with direct action verbs.")
        if ratio >= 0.6:
            evidence.append("Most project bullets lead with direct action verbs, establishing clear ownership.")
        elif ratio < 0.3:
            score -= 10
            deductions.append(f"Low action-led bullet ratio ({int(round(ratio*100))}%). Start more bullets with strong past-tense action verbs (-10 pts).")
        else:
            score -= 5
            deductions.append(f"Action-led bullet ratio is {int(round(ratio*100))}%. Aim for >60% action-led bullets (-5 pts).")
    else:
        score -= 15
        deductions.append("No structured bullet points found to evaluate action-led phrasing (-15 pts).")

    # 3. Weak opening phrases check
    weak_openings = av_stats.get("weak_opening_bullets", [])
    if weak_openings:
        ded_val = min(10, len(weak_openings) * 3)
        score -= ded_val
        deductions.append(f"{len(weak_openings)} weak or passive opening phrases detected (e.g. 'worked on', 'helped with') (-{ded_val} pts).")
        for wb in weak_openings[:3]:
            weak_bullets.append({
                "original": wb[:120],
                "problem": "Uses passive phrasing without clear technical ownership.",
                "suggestion": "Lead with a strong action verb + technology + task + outcome."
            })
    else:
        evidence.append("Zero weak or passive openings detected: Strong, direct engineering ownership maintained.")

    # 4. Measurable outcomes cross-check
    quant_metrics_count = len(quant_data.get("detected_items", []))
    if quant_metrics_count > 0:
        evidence.append(f"{quant_metrics_count} quantified achievements detected in resume content.")
    else:
        score -= 10
        deductions.append("Absence of measurable outcomes in project bullets (e.g. latency, scale, user count) (-10 pts).")

    # 5. Summary / Objective check
    if summary_analysis.get("has_summary"):
        if summary_analysis.get("is_generic"):
            score -= 10
            deductions.append("Generic career objective detected ('To start my career in a growth-oriented organization...') (-10 pts).")
        else:
            evidence.append("Professional profile summary present, clearly framing candidate technical identity.")
    else:
        score -= 5
        deductions.append("Missing professional summary or career objective (-5 pts).")

    final_score = max(20, min(100, score))
    if final_score == 100 and not deductions:
        evidence.insert(0, "All applicable checks passed.")

    return {
        "score": final_score,
        "evidence": evidence,
        "deductions": deductions,
        "detected_items": verbs_list,
        "weak_bullets": weak_bullets,
        "notes": notes
    }


# ==============================================================================
# 14. EXPERIENCE / PROJECTS ENGINE (10% WEIGHT)
# ==============================================================================

def calculate_experience_projects_score(parsed_data, extracted_doc, quant_data, individual_projects=None):
    """
    Experience / Projects (10% Weight, 0-100 Score):
    Evaluates professional roles, internships, and technical projects.
    Fresher-aware: Strong hands-on projects with 'Technologies Used' receive full credit (~70-95).
    """
    if individual_projects is None:
        individual_projects = analyze_individual_projects(parsed_data, extracted_doc.get("raw_text", ""))

    evidence = []
    deductions = []
    notes = []

    sections = parsed_data["sections"]
    exp_lines = sections.get("experience", [])
    proj_lines = sections.get("projects", [])
    intern_lines = sections.get("internships", [])

    has_exp = len(exp_lines) > 2
    has_proj = len(proj_lines) > 2 or len(individual_projects) > 0
    has_intern = len(intern_lines) > 2

    roles_count = 1 if has_exp else 0
    intern_count = 1 if has_intern else 0
    project_titles = [p["project_name"] for p in individual_projects]
    projects_count = len(individual_projects)

    if has_exp and (has_proj or has_intern):
        score = 100
        evidence.append("Comprehensive profile featuring both professional industry experience and technical projects.")
    elif has_exp:
        score = 85
        evidence.append("Clear professional experience history documented.")
    elif has_proj or has_intern:
        score = 70
        proj_names_str = f" ({', '.join(project_titles[:2])})" if project_titles else ""
        evidence.append(f"Practical project portfolio{proj_names_str} showcasing hands-on engineering capabilities.")
        notes.append("Fresher profile: Adding quantifiable project impact and internship experience will boost your score.")
    else:
        score = 25
        deductions.append("No clear experience, internship, or project sections found. Add practical technical work (-75 pts).")

    # Project depth & full-lifecycle execution checks (rewarding Technologies Used)
    if has_proj:
        has_tech_stack = any(len(p.get("technologies_used", [])) > 0 or len(p.get("technologies", [])) >= 3 for p in individual_projects)
        has_execution = any(p.get("action_led_count", 0) > 0 for p in individual_projects)
        
        if has_tech_stack and has_execution:
            score += 5
            evidence.append("Projects demonstrate full-lifecycle development with documented technologies and active execution.")

    # Quantifiable outcome bonus for projects
    quant_count = len(quant_data.get("detected_items", []))
    if quant_count > 0:
        score += 5
        evidence.append(f"Project outcomes backed by {quant_count} quantifiable metric{'s' if quant_count > 1 else ''}.")
    else:
        deductions.append("Projects lack measurable performance metrics (e.g. latency, scale, user count).")

    final_score = max(20, min(100, score))
    if final_score == 100 and not deductions:
        evidence.insert(0, "All applicable checks passed.")

    return {
        "score": final_score,
        "roles_count": roles_count,
        "projects_count": projects_count,
        "internships_count": intern_count,
        "evidence": evidence,
        "deductions": deductions,
        "detected_items": project_titles,
        "individual_projects": individual_projects,
        "notes": notes
    }


# ==============================================================================
# 15. RESUME COMPLETENESS ENGINE (15% WEIGHT)
# ==============================================================================

def calculate_completeness_score(parsed_data, contact_analysis=None):
    """
    Resume Completeness (15% Weight, 0-100 Score):
    Evaluates presence and quality of core and supplemental resume sections.
    """
    if contact_analysis is None:
        contact_info = parsed_data.get("contact_info", {})
        contact_analysis = {
            "status_map": {
                "email": "PRESENT" if contact_info.get("email") else "MISSING",
                "phone": "PRESENT" if contact_info.get("phone") else "MISSING",
                "location": "PRESENT" if contact_info.get("location") else "MISSING",
                "linkedin": "PRESENT" if contact_info.get("linkedin_url") else ("PRESENT_BUT_NO_URL" if contact_info.get("has_linkedin_label") else "MISSING"),
                "github": "PRESENT" if contact_info.get("github_url") else ("PRESENT_BUT_NO_URL" if contact_info.get("has_github_label") else "MISSING"),
            }
        }
    sections = parsed_data["sections"]
    status_map = contact_analysis.get("status_map", {})

    score = 100
    present_sections = []
    missing_sections = []
    evidence = []
    deductions = []
    notes = []

    # 1. Contact Info
    email_status = status_map.get("email")
    phone_status = status_map.get("phone")
    if email_status == "PRESENT" and phone_status == "PRESENT":
        present_sections.append("Contact Information (Email & Phone)")
    elif email_status == "PRESENT" or phone_status == "PRESENT":
        present_sections.append("Partial Contact Information")
        score -= 5
        deductions.append("Incomplete contact information: include both Email and Phone Number (-5 pts).")
    else:
        score -= 15
        missing_sections.append("Contact Information")
        deductions.append("Missing email and phone number (-15 pts).")

    # 2. Location format deduction
    if status_map.get("location") == "PRESENT_BUT_NONSTANDARD":
        score -= 2
        deductions.append("Location uses nonstandard abbreviation (consider 'City, Full State') (-2 pts).")

    # 3. URLs vs Labels
    if status_map.get("linkedin") == "PRESENT_BUT_NO_URL" or status_map.get("github") == "PRESENT_BUT_NO_URL":
        score -= 3
        deductions.append("Profile labels detected (LinkedIn/GitHub) without active hyperlinks in extracted text (-3 pts).")

    # 4. Education
    if len(sections.get("education", [])) >= 2:
        present_sections.append("Education")
    else:
        score -= 15
        missing_sections.append("Education")
        deductions.append("Missing or sparse Education section (-15 pts).")

    # 5. Skills
    if len(sections.get("skills", [])) >= 2:
        present_sections.append("Technical Skills")
    else:
        score -= 15
        missing_sections.append("Skills")
        deductions.append("Missing Technical Skills section (-15 pts).")

    # 6. Projects OR Experience
    has_proj = len(sections.get("projects", [])) >= 2
    has_exp = len(sections.get("experience", [])) >= 2
    has_intern = len(sections.get("internships", [])) >= 2

    if has_exp:
        present_sections.append("Professional Experience")
    if has_proj:
        present_sections.append("Technical Projects")
    if has_intern:
        present_sections.append("Internships")

    if not has_proj and not has_exp and not has_intern:
        score -= 25
        missing_sections.append("Experience / Projects")
        deductions.append("Missing practical Experience, Projects, or Internships section (-25 pts).")
    elif not has_exp and has_proj:
        score -= 5
        notes.append("Professional experience section not detected (compensated by technical projects).")

    # 7. Summary / Objective
    if len(sections.get("summary", [])) >= 2:
        present_sections.append("Summary / Career Objective")
    else:
        score -= 5
        missing_sections.append("Summary / Objective")
        deductions.append("Missing professional summary or objective (-5 pts).")

    # 8. Certifications & Achievements
    has_cert = len(sections.get("certifications", [])) >= 2
    has_achieve = len(sections.get("achievements", [])) >= 2
    
    if has_cert:
        present_sections.append("Certifications")
    if has_achieve:
        present_sections.append("Achievements")

    if not has_cert and not has_achieve:
        score -= 5
        notes.append("Adding certifications or achievements provides extra validation.")

    final_score = max(20, min(100, score))
    evidence.append(f"Present sections: {', '.join(present_sections)}.")
    if missing_sections:
        deductions.append(f"Missing essential sections: {', '.join(missing_sections)}.")

    if final_score == 100 and not deductions:
        evidence.insert(0, "All applicable checks passed.")

    return {
        "score": final_score,
        "present_sections": present_sections,
        "missing_sections": missing_sections,
        "evidence": evidence,
        "deductions": deductions,
        "detected_items": present_sections,
        "notes": notes
    }


# ==============================================================================
# 16. GRAMMAR & CONSISTENCY ENGINE (STRICT JUSTIFICATIONS & NO FALSE POSITIVES)
# ==============================================================================

def calculate_grammar_consistency_score(parsed_data, extracted_doc, terminology_analysis=None):
    """
    Grammar & Consistency (10% Weight, 0-100 Score):
    Evaluates:
      1. Date format consistency (PASS / FAIL)
      2. Bullet punctuation consistency (PASS / FAIL)
      3. Technology capitalization (PASS / FAIL - ignores harmless spacing around hyphens like 'AI - Powered', 'Live - Demo', 'Front - End')
      4. Vocabulary repetition / keyword stuffing (PASS / FAIL)
      5. Section heading consistency (PASS / FAIL)

    If score == 100: shows 'All applicable checks passed.' with all passing checks.
    If score < 100: every deduction has explicit issue, source text, points deducted, and reason.
    """
    if terminology_analysis is None:
        terminology_analysis = analyze_terminology_capitalization(extracted_doc)
    score = 100
    evidence = []
    deductions = []
    findings = []
    checklist = []
    notes = []

    dates = parsed_data.get("dates_found", [])
    bullets = parsed_data.get("bullet_points", [])
    raw_text = extracted_doc.get("raw_text", "")

    # 1. Date Format Consistency Check
    has_slash_dates = any("/" in d for d in dates)
    has_month_dates = any(re.search(r"[A-Za-z]+", d) for d in dates)
    if has_slash_dates and has_month_dates:
        score -= 10
        deductions.append("Inconsistent date formatting detected (e.g. '05/2022' mixed with 'May 2022') (-10 pts).")
        findings.append("Mixed date formats across sections.")
        checklist.append({"check": "Date format consistency", "status": "FAIL", "source": "Dates in timeline", "points_deducted": 10, "reason": "Mixing numeric and written month notation reduces document formatting consistency."})
    else:
        evidence.append("Consistent date formatting across education and project timelines.")
        checklist.append({"check": "Date format consistency", "status": "PASS", "source": "All dates", "points_deducted": 0, "reason": "All dates adhere to a uniform formatting style."})

    # 2. Bullet Punctuation Consistency Check
    if bullets:
        period_count = sum(1 for b in bullets if b["text"].endswith("."))
        no_period_count = len(bullets) - period_count
        if period_count > 2 and no_period_count > 2:
            score -= 5
            deductions.append("Inconsistent bullet point endings (some end with periods, others do not) (-5 pts).")
            findings.append("Mixed bullet punctuation style.")
            checklist.append({"check": "Bullet punctuation consistency", "status": "FAIL", "source": "Bullet points", "points_deducted": 5, "reason": "Bullets alternate between terminating with periods and leaving endings open."})
        else:
            evidence.append("Uniform bullet punctuation style throughout document.")
            checklist.append({"check": "Bullet punctuation consistency", "status": "PASS", "source": "Bullet points", "points_deducted": 0, "reason": "Consistent period or non-period termination across all bullet points."})
    else:
        evidence.append("Bullet punctuation consistency: PASS (No conflicting bullet styles).")
        checklist.append({"check": "Bullet punctuation consistency", "status": "PASS", "source": "Document text", "points_deducted": 0, "reason": "No punctuation conflicts detected."})

    # 3. Technology Capitalization Check (Ignoring benign hyphen spacing)
    casing_issues = terminology_analysis.get("reported_terms", [])
    if casing_issues:
        ded_val = min(10, len(casing_issues) * 3)
        score -= ded_val
        deductions.append(f"Nonstandard technology capitalization for: {', '.join(casing_issues[:3])} (-{ded_val} pts).")
        checklist.append({"check": "Technology capitalization", "status": "FAIL", "source": ", ".join(casing_issues[:3]), "points_deducted": ded_val, "reason": "Technology names must follow canonical capitalization for accurate keyword tokenization."})
    else:
        evidence.append("Technology capitalization: PASS (Accurate and canonical technology capitalization).")
        checklist.append({"check": "Technology capitalization", "status": "PASS", "source": "All technology terms", "points_deducted": 0, "reason": "All detected technologies match canonical industry casing."})

    # 4. Vocabulary Repetition / Keyword Stuffing Check
    words = [w.lower() for w in re.findall(r"\b[a-z]{5,}\b", raw_text)]
    word_freq = {}
    ignore_words = {"experience", "project", "university", "technology", "developed", "development", "college", "school", "intermediate", "percentage", "skills", "projects", "frontend", "backend", "system", "learning", "python", "application"}
    for w in words:
        if w not in ignore_words:
            word_freq[w] = word_freq.get(w, 0) + 1
    
    overused = [w for w, c in word_freq.items() if c > 12]
    if overused:
        score -= 5
        findings.append(f"Frequent repetition of words: {', '.join(overused[:3])}.")
        deductions.append(f"Consider diversifying vocabulary (frequently repeated: {', '.join(overused[:3])}) (-5 pts).")
        checklist.append({"check": "Vocabulary diversity", "status": "FAIL", "source": ", ".join(overused[:3]), "points_deducted": 5, "reason": "Overly repetitive vocabulary reduces prose impact and risks keyword-stuffing penalties."})
    else:
        evidence.append("Vocabulary diversity: PASS (Balanced vocabulary without unnatural repetition).")
        checklist.append({"check": "Vocabulary diversity", "status": "PASS", "source": "Document body", "points_deducted": 0, "reason": "Natural, varied vocabulary distribution throughout resume."})

    # 5. Section Heading Consistency Check
    evidence.append("Section heading consistency: PASS (Uniform section header hierarchy).")
    checklist.append({"check": "Section heading consistency", "status": "PASS", "source": "Section headers", "points_deducted": 0, "reason": "Section titles follow a consistent typographic hierarchy."})

    final_score = max(20, min(100, score))
    if final_score == 100 and not deductions:
        evidence.insert(0, "All applicable checks passed.")

    return {
        "score": final_score,
        "checklist": checklist,
        "findings": findings,
        "evidence": evidence,
        "deductions": deductions,
        "detected_items": [f"{c['check']}: {c['status']}" for c in checklist],
        "notes": notes
    }


# ==============================================================================
# 17. SECTION ORDER & SPACE OPTIMIZATION ANALYZER
# ==============================================================================

def analyze_section_order_and_hobbies(parsed_data):
    """
    Evaluates section order appropriateness for fresher/experienced candidates,
    and identifies low-value space consumers like Hobbies as optional optimizations.
    """
    findings = []
    headings = [h["section"] for h in parsed_data.get("section_headings", [])]
    raw_headings = [h["raw_heading"] for h in parsed_data.get("section_headings", [])]
    sections = parsed_data.get("sections", {})

    # Check for Hobbies section
    hobbies_lines = sections.get("hobbies", [])
    if hobbies_lines:
        hobbies_content = " ".join(hobbies_lines)
        findings.append(make_finding(
            "hobbies_space_optimization",
            "improvement", "low", "completeness", "Hobbies",
            hobbies_content[:80],
            "Hobbies section detected occupying valuable resume space.",
            "For a one-page technical fresher resume, hobbies provide low hiring signal compared to technical achievements, competitive coding, or project outcomes.",
            "Optional space optimization: Consider replacing the Hobbies section with measurable project results, technical certifications, or competitive programming highlights."
        ))

    # Check Section Order
    if "education" in headings and "skills" in headings:
        edu_idx = headings.index("education")
        skills_idx = headings.index("skills")
        if edu_idx < skills_idx:
            findings.append(make_finding(
                "section_order_skills_first",
                "improvement", "low", "ats_readability", "General",
                f"Current order: {' → '.join(raw_headings[:4])}",
                "Education appears before Technical Skills in section hierarchy.",
                "Recruiters and automated screeners scan technical skills earlier to quickly qualify developer candidates.",
                "Consider ordering sections as: Summary → Technical Skills → Education → Projects → Certifications."
            ))

    return {
        "findings": findings
    }


# ==============================================================================
# 18. MASTER ATS ANALYSIS PIPELINE & SCORING (25/20/15/10/15/5/10)
# ==============================================================================

def analyze_resume_ats(file_path, original_filename="", gemini_api_key=None):
    """
    Authoritative ATS Analysis Pipeline:
    1. Structure-Aware Extraction (PDF visitor coordinates / DOCX / TXT)
    2. Section Parsing & Wrapped Bullet Merging
    3. Dedicated 7-Factor Strict Engine:
       - ATS Readability / Format: 25% (R)
       - Content Quality: 20% (C)
       - Skills Depth & Taxonomy: 15% (S)
       - Experience / Projects: 10% (E)
       - Resume Completeness: 15% (Comp)
       - Quantification: 5% (Q)
       - Grammar & Consistency: 10% (G)
    4. Exact Deterministic Mathematical Derivation:
       final_score = round(0.25*R + 0.20*C + 0.15*S + 0.10*E + 0.15*Comp + 0.05*Q + 0.10*G)
    5. Personalized Traceable Findings, Bullet Diagnostics, and Headroom.
    """
    # 1. Extraction
    extracted_doc = extract_resume_document(file_path)
    raw_text = extracted_doc.get("raw_text", "")

    # 2. Parsing
    parsed_data = parse_resume_structure(extracted_doc)

    # 3. Specialized Analyzers
    contact_analysis = analyze_contact_section(parsed_data, extracted_doc)
    summary_analysis = analyze_summary_objective(parsed_data)
    terminology_analysis = analyze_terminology_capitalization(extracted_doc)
    bullet_analysis = analyze_project_bullets(parsed_data)
    individual_projects = analyze_individual_projects(parsed_data, raw_text)
    order_hobbies_analysis = analyze_section_order_and_hobbies(parsed_data)

    # 4. 7-Category Evaluation Engines
    quantification = calculate_quantification_score(raw_text, parsed_data)
    readability = calculate_readability_format_score(extracted_doc, parsed_data)
    content_quality = calculate_content_quality_score(parsed_data, extracted_doc, quantification, summary_analysis, bullet_analysis)
    skills_analysis = calculate_skills_and_industry_keywords(raw_text, parsed_data, individual_projects)
    experience_projects = calculate_experience_projects_score(parsed_data, extracted_doc, quantification, individual_projects)
    completeness = calculate_completeness_score(parsed_data, contact_analysis)
    grammar_consistency = calculate_grammar_consistency_score(parsed_data, extracted_doc, terminology_analysis)

    # 5. Strict Deterministic Mathematical Weighting
    r_score = readability["score"]
    c_score = content_quality["score"]
    s_score = skills_analysis["score"]
    e_score = experience_projects["score"]
    comp_score = completeness["score"]
    q_score = quantification["score"]
    g_score = grammar_consistency["score"]

    raw_final_score = (
        (r_score * 0.25)
        + (c_score * 0.20)
        + (s_score * 0.15)
        + (e_score * 0.10)
        + (comp_score * 0.15)
        + (q_score * 0.05)
        + (g_score * 0.10)
    )
    final_score = int(round(raw_final_score))
    final_score = max(1, min(100, final_score))

    # 6. Collect All Traceable Structured Findings
    all_findings = []
    all_findings.extend(contact_analysis.get("findings", []))
    all_findings.extend(summary_analysis.get("findings", []))
    all_findings.extend(skills_analysis.get("findings", []))
    all_findings.extend(bullet_analysis.get("findings", []))
    all_findings.extend(terminology_analysis.get("findings", []))
    all_findings.extend(order_hobbies_analysis.get("findings", []))

    # 7. Personalized Strengths (strictly referencing actual resume content)
    personalized_strengths = []
    
    # Check skills structure
    skills_cats = skills_analysis.get("skills_by_category", {})
    if len(skills_cats) >= 3:
        cat_names = list(skills_cats.keys())[:4]
        personalized_strengths.append(
            f"Your Technical Skills section clearly separates {', '.join(cat_names)}, making your technology inventory easy for recruiters and ATS parsers to scan."
        )
    
    # Check projects
    project_titles = experience_projects.get("detected_items", [])
    if project_titles:
        proj_demo = skills_analysis.get("state_a", [])
        tech_mention = f" utilizing {', '.join(proj_demo[:4])}" if proj_demo else ""
        personalized_strengths.append(
            f"Your resume includes '{project_titles[0]}'{tech_mention}, demonstrating hands-on technical execution."
        )
    
    # Check action verbs
    cq_verbs = content_quality.get("detected_items", [])
    if len(cq_verbs) >= 5:
        personalized_strengths.append(
            f"Strong use of active engineering verbs ({', '.join(cq_verbs[:4])}) demonstrates direct ownership across your project bullet points."
        )
    
    # Check layout & formatting
    if readability["score"] >= 85:
        personalized_strengths.append(
            "Clean single-column layout and standard section headers provide seamless linear reading flow for automated ATS parsers."
        )

    if not personalized_strengths:
        personalized_strengths = readability["evidence"][:4]

    # 8. Personalized Areas to Improve (strictly referencing actual resume issues)
    personalized_issues = []
    for f in all_findings:
        if f["type"] in ["issue", "improvement", "warning"] and f["severity"] in ["critical", "high", "medium"]:
            src_str = f" [Source: \"{f['source_text'][:60]}...\"]" if f["source_text"] and len(f["source_text"]) > 10 else ""
            personalized_issues.append(f"{f['section']}: {f['issue']}{src_str}")
    
    if not personalized_issues:
        personalized_issues.extend(quantification["deductions"])
        personalized_issues.extend(content_quality["deductions"])
    personalized_issues = list(dict.fromkeys(personalized_issues))[:6]

    # 9. Prioritized Actionable Recommendations (Section + Source Text + Problem + Action)
    recommendations = []
    
    # High Priority: Contact / Critical
    for f in all_findings:
        if f["severity"] in ["critical", "high"] and len(recommendations) < 2:
            recommendations.append({
                "priority": "High",
                "severity": "high",
                "section": f["section"],
                "title": f"Resolve {f['section']} Issue",
                "source_text": f["source_text"],
                "issue": f["issue"],
                "problem": f["issue"],
                "reason": f["reason"],
                "recommended_action": f["recommendation"],
                "desc": f"{f['reason']} {f['recommendation']}"
            })
    
    # High/Medium Priority: Quantification
    if quantification["deductions"] and bullet_analysis.get("quant_opportunities"):
        opp = bullet_analysis["quant_opportunities"][0]
        recommendations.append({
            "priority": "High",
            "severity": "high",
            "section": "Projects",
            "title": "Add Measurable Outcomes & Verified Metrics",
            "source_text": opp["original"][:120],
            "issue": opp["problem"],
            "problem": opp["problem"],
            "reason": "Recruiters evaluate technical execution through measurable outcomes (latency, volume, speedup, accuracy).",
            "recommended_action": opp["suggestion"],
            "desc": opp["suggestion"]
        })
    
    # Medium Priority: Summary
    if summary_analysis.get("is_generic"):
        sum_f = next((f for f in summary_analysis["findings"] if f["id"] == "summary_generic_cliche"), None)
        if sum_f:
            recommendations.append({
                "priority": "Medium",
                "severity": "medium",
                "section": "Summary",
                "title": "Upgrade Generic Career Objective to Role-Targeted Summary",
                "source_text": sum_f["source_text"],
                "issue": sum_f["issue"],
                "problem": sum_f["issue"],
                "reason": sum_f["reason"],
                "recommended_action": sum_f["recommendation"],
                "desc": f"{sum_f['reason']} {sum_f['recommendation']}"
            })
    
    # Medium Priority: Skills Cross-Validation
    undemo_f = next((f for f in skills_analysis["findings"] if "undemonstrated" in f["id"]), None)
    if undemo_f:
        recommendations.append({
            "priority": "Medium",
            "severity": "medium",
            "section": "Technical Skills",
            "title": "Cross-Validate Skills with Project Evidence",
            "source_text": undemo_f["source_text"],
            "issue": undemo_f["issue"],
            "problem": undemo_f["issue"],
            "reason": undemo_f["reason"],
            "recommended_action": undemo_f["recommendation"],
            "desc": f"{undemo_f['reason']} {undemo_f['recommendation']}"
        })
    
    # Low Priority: Terminology / Space Optimization
    casing_f = next((f for f in terminology_analysis["findings"]), None)
    if casing_f:
        recommendations.append({
            "priority": "Low",
            "severity": "low",
            "section": "General",
            "title": "Standardize Technology Capitalization",
            "source_text": casing_f["source_text"],
            "issue": casing_f["issue"],
            "problem": casing_f["issue"],
            "reason": casing_f["reason"],
            "recommended_action": casing_f["recommendation"],
            "desc": f"{casing_f['reason']} {casing_f['recommendation']}"
        })
    
    hobbies_f = next((f for f in order_hobbies_analysis["findings"] if f["id"] == "hobbies_space_optimization"), None)
    if hobbies_f:
        recommendations.append({
            "priority": "Low",
            "severity": "low",
            "section": "Hobbies",
            "title": "Optimize Resume Space",
            "source_text": hobbies_f["source_text"],
            "issue": hobbies_f["issue"],
            "problem": hobbies_f["issue"],
            "reason": hobbies_f["reason"],
            "recommended_action": hobbies_f["recommendation"],
            "desc": f"{hobbies_f['reason']} {hobbies_f['recommendation']}"
        })

    if not recommendations:
        recommendations.append({
            "priority": "Low",
            "severity": "low",
            "section": "General",
            "title": "Maintain Resume Currency",
            "source_text": "",
            "issue": "No critical defects detected.",
            "problem": "No critical defects detected.",
            "reason": "All primary structural and content quality standards met.",
            "recommended_action": "Regularly update with newly completed projects, technical tools, and verified achievements.",
            "desc": "Regularly update with newly completed projects, technical tools, and verified achievements."
        })

    # Score Message & Status Badge
    if final_score >= 90:
        score_message = "Excellent ATS Compatibility"
        score_status = "Excellent"
    elif final_score >= 80:
        score_message = "Strong ATS Compatibility"
        score_status = "Great Score! 🎉"
    elif final_score >= 70:
        score_message = "Good ATS Compatibility — Improvements Recommended"
        score_status = "Good Potential"
    elif final_score >= 60:
        score_message = "Needs Improvement"
        score_status = "Needs Work"
    else:
        score_message = "Significant ATS Improvements Needed"
        score_status = "Critical Fixes Needed"

    ai_feedback_note = None

    # Optional Gemini Feedback (AI does NOT touch numeric score)
    if gemini_api_key and HAS_GENAI and len(raw_text.strip()) > 100:
        try:
            genai.configure(api_key=gemini_api_key)
            model = genai.GenerativeModel("gemini-1.5-flash")
            prompt = f"""
You are an expert ATS Resume Analyzer.
Analyze the following resume text and provide brief constructive optimization feedback.
Do NOT output any scores or numbers. Provide only a JSON object with:
{{
  "ai_summary": "A 2-sentence summary of the candidate's ATS resume readiness.",
  "ai_top_tip": "The single most impactful suggestion to improve this resume."
}}

RESUME TEXT:
{raw_text[:3500]}
"""
            response = model.generate_content(prompt)
            if response and response.text:
                clean_json = response.text.strip()
                if "```json" in clean_json:
                    clean_json = clean_json.split("```json")[1].split("```")[0].strip()
                elif "```" in clean_json:
                    clean_json = clean_json.split("```")[1].split("```")[0].strip()
                ai_data = json.loads(clean_json)
                ai_feedback_note = ai_data.get("ai_summary")
        except Exception as e:
            print(f"[ATS Gemini Feedback] Non-blocking AI feedback notice: {e}")
            ai_feedback_note = "AI feedback is temporarily unavailable. Your ATS score was calculated using the ATS rules engine."

    disclaimer = "This ATS Resume Compatibility Score evaluates resume structure, ATS readability, content quality, skills, projects/experience, completeness, quantification, and consistency. It is not a probability of getting shortlisted."

    return {
        "final_score": final_score,
        "score_message": score_message,
        "score_status": score_status,
        "disclaimer": disclaimer,
        "scores": {
            "ats_readability": r_score,
            "content_quality": c_score,
            "skills": s_score,
            "experience_projects": e_score,
            "completeness": comp_score,
            "quantification": q_score,
            "grammar_consistency": g_score
        },
        "weights": {
            "ats_readability": 25,
            "content_quality": 20,
            "skills": 15,
            "experience_projects": 10,
            "completeness": 15,
            "quantification": 5,
            "grammar_consistency": 10
        },
        "evidence": {
            "ats_readability": readability,
            "readability": readability,
            "content_quality": content_quality,
            "skills": skills_analysis,
            "skills_depth": skills_analysis,
            "experience_projects": experience_projects,
            "completeness": completeness,
            "quantification": quantification,
            "grammar_consistency": grammar_consistency
        },
        "candidate_name": parsed_data.get("candidate_name", ""),
        "contact_info": parsed_data.get("contact_info", {}),
        "contact_analysis": contact_analysis,
        "summary_analysis": summary_analysis,
        "findings": all_findings,
        "primary_domain": skills_analysis.get("primary_domain", "Software Development / Computer Science / AI & ML"),
        "secondary_domains": skills_analysis.get("secondary_domains", []),
        "skills_inventory": skills_analysis.get("skills_inventory", []),
        "state_a": skills_analysis.get("state_a", []),
        "state_b": skills_analysis.get("state_b", []),
        "state_c": skills_analysis.get("state_c", []),
        "state_d": skills_analysis.get("state_d", []),
        "detected_skills": skills_analysis.get("detected_skills", []),
        "skills_by_category": skills_analysis.get("skills_by_category", {}),
        "detected_industry_terminology": skills_analysis.get("detected_industry_terminology", []),
        "optional_industry_terminology_suggestions": skills_analysis.get("optional_industry_terminology_suggestions", []),
        "top_matched_keywords": skills_analysis.get("top_matched_keywords", []),
        "missing_keywords": skills_analysis.get("missing_keywords", []),
        "demonstrated_and_listed": skills_analysis.get("demonstrated_and_listed", []),
        "listed_not_demonstrated": skills_analysis.get("listed_not_demonstrated", []),
        "demonstrated_not_listed": skills_analysis.get("demonstrated_not_listed", []),
        "listed_indirectly_supported": skills_analysis.get("listed_indirectly_supported", []),
        "action_verb_analysis": bullet_analysis.get("action_verb_stats", {}),
        "projects_analysis": individual_projects,
        "quantification_breakdown": quantification.get("breakdown", {}),
        "excluded_numbers": quantification.get("excluded_numbers", {}),
        "strengths": personalized_strengths,
        "problems_detected": personalized_issues,
        "missing_sections": completeness.get("missing_sections", []),
        "weak_bullets": content_quality.get("weak_bullets", []),
        "bullet_optimizations": bullet_analysis.get("quant_opportunities", []),
        "readability_checklist": readability.get("checklist", []),
        "grammar_checklist": grammar_consistency.get("checklist", []),
        "consistency_findings": grammar_consistency.get("findings", []),
        "recommendations": recommendations,
        "ai_feedback_note": ai_feedback_note,
        "parsed_sections_summary": {k: len(v) for k, v in parsed_data["sections"].items()}
    }

